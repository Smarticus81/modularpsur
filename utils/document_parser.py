"""
State-of-the-art document parsing utilities for PSUR generation.
Uses modern NLP and parsing libraries for robust text extraction.

ENHANCED: Now includes semantic parsing integration via Anthropic Claude API.
Use get_semantic_parser() for advanced semantic understanding.

SUPPORTS: .docx (Word) and .pdf files
"""

import docx2txt
import dateparser
from docx import Document
from datetime import datetime, timedelta
from typing import List, Tuple, Optional, Dict, Any
import re
import os
from pathlib import Path

# PDF support
try:
    import pdfplumber
    import pypdf
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False


class DocumentParser:
    """
    Advanced document parser using state-of-the-art parsing libraries.
    
    Legacy Methods:
    - extract_text_with_structure(): Basic text extraction
    - extract_context_by_keywords(): Keyword-based extraction
    - extract_dates(), find_date_ranges(): Date extraction
    
    Semantic Methods (NEW):
    - get_semantic_parser(): Get semantic parser instance
    - extract_with_semantic_understanding(): Advanced parsing with Claude API
    """
    
    @staticmethod
    def extract_text_with_structure(doc_path: str) -> Dict[str, any]:
        """
        Extract text from document (Word or PDF) with structure preservation.
        Automatically detects file type and uses appropriate parser.
        
        Args:
            doc_path: Path to .docx or .pdf file
        
        Returns:
            Dict with 'full_text', 'paragraphs', and 'tables'
        """
        file_ext = Path(doc_path).suffix.lower()
        
        if file_ext == '.pdf':
            return DocumentParser._extract_pdf_with_structure(doc_path)
        elif file_ext == '.docx':
            return DocumentParser._extract_docx_with_structure(doc_path)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported: .docx, .pdf")
    
    @staticmethod
    def _extract_docx_with_structure(docx_path: str) -> Dict[str, any]:
        """Extract text from Word document with structure preservation."""
        try:
            # Extract plain text using docx2txt (handles complex documents better)
            full_text = docx2txt.process(docx_path)
            
            # Also use python-docx for structured access
            doc = Document(docx_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                'full_text': full_text,
                'paragraphs': paragraphs,
                'tables': tables,
                'file_type': 'docx'
            }
        except Exception as e:
            print(f"Error extracting Word document: {e}")
            return {'full_text': '', 'paragraphs': [], 'tables': [], 'file_type': 'docx'}
    
    @staticmethod
    def _extract_pdf_with_structure(pdf_path: str) -> Dict[str, any]:
        """
        Extract text from PDF with structure preservation.
        Uses pdfplumber for robust table extraction and pypdf as fallback.
        """
        if not PDF_SUPPORT:
            raise ImportError(
                "PDF support requires pypdf and pdfplumber. "
                "Install with: pip install pypdf pdfplumber"
            )
        
        try:
            paragraphs = []
            tables = []
            full_text_parts = []
            
            print(f"  Extracting PDF: {Path(pdf_path).name}")
            
            # Use pdfplumber for better table and text extraction
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
                print(f"  Processing {total_pages} pages...")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    if page_num % 10 == 0:
                        print(f"    Page {page_num}/{total_pages}...")
                    
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        full_text_parts.append(page_text)
                        
                        # Split into paragraphs (blank line separated)
                        page_paragraphs = [
                            p.strip() 
                            for p in page_text.split('\n\n') 
                            if p.strip()
                        ]
                        paragraphs.extend(page_paragraphs)
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    if page_tables:
                        for table in page_tables:
                            # Clean table data
                            clean_table = [
                                [cell.strip() if cell else '' for cell in row]
                                for row in table
                            ]
                            tables.append(clean_table)
            
            full_text = '\n\n'.join(full_text_parts)
            
            print(f"  Extracted {len(paragraphs)} paragraphs, {len(tables)} tables from PDF")
            
            return {
                'full_text': full_text,
                'paragraphs': paragraphs,
                'tables': tables,
                'file_type': 'pdf',
                'page_count': total_pages
            }
            
        except Exception as e:
            print(f"Error extracting PDF: {e}")
            print(f"  Attempting fallback extraction with pypdf...")
            
            # Fallback to pypdf for text-only extraction
            try:
                from pypdf import PdfReader
                
                reader = PdfReader(pdf_path)
                full_text_parts = []
                
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        full_text_parts.append(text)
                
                full_text = '\n\n'.join(full_text_parts)
                paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]
                
                print(f"  Fallback extraction: {len(paragraphs)} paragraphs")
                
                return {
                    'full_text': full_text,
                    'paragraphs': paragraphs,
                    'tables': [],
                    'file_type': 'pdf',
                    'extraction_method': 'fallback'
                }
            except Exception as fallback_error:
                print(f"  Fallback extraction also failed: {fallback_error}")
                return {
                    'full_text': '', 
                    'paragraphs': [], 
                    'tables': [], 
                    'file_type': 'pdf',
                    'error': str(fallback_error)
                }
    
    @staticmethod
    def extract_context_by_keywords(text: str, keywords: List[str], 
                                    context_size: int = 500) -> List[str]:
        """
        Extract text chunks around keywords using advanced context extraction.
        
        Args:
            text: Full text to search
            keywords: List of keywords to find
            context_size: Characters of context around each keyword
            
        Returns:
            List of relevant text chunks
        """
        chunks = []
        text_lower = text.lower()
        
        for keyword in keywords:
            keyword_lower = keyword.lower()
            pos = 0
            
            while True:
                pos = text_lower.find(keyword_lower, pos)
                if pos == -1:
                    break
                
                # Extract context around keyword
                start = max(0, pos - context_size)
                end = min(len(text), pos + len(keyword) + context_size)
                chunk = text[start:end]
                
                # Clean up chunk boundaries (try to end on sentence/word boundaries)
                if start > 0:
                    # Find first sentence start
                    sentence_start = chunk.find('. ')
                    if sentence_start != -1 and sentence_start < 100:
                        chunk = chunk[sentence_start + 2:]
                
                if end < len(text):
                    # Find last sentence end
                    sentence_end = chunk.rfind('. ')
                    if sentence_end != -1 and len(chunk) - sentence_end < 100:
                        chunk = chunk[:sentence_end + 1]
                
                chunks.append(chunk)
                pos += len(keyword)
        
        # Remove duplicates while preserving order
        seen = set()
        unique_chunks = []
        for chunk in chunks:
            if chunk not in seen:
                seen.add(chunk)
                unique_chunks.append(chunk)
        
        return unique_chunks
    
    @staticmethod
    def extract_dates(text: str, search_future: bool = False) -> List[Tuple[datetime, str]]:
        """
        Extract dates from text using state-of-the-art dateparser.
        
        Args:
            text: Text to extract dates from
            search_future: Whether to include future dates
            
        Returns:
            List of (datetime, original_text) tuples
        """
        found_dates = []
        
        # Common date patterns to look for
        date_patterns = [
            r'\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}',
            r'(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}',
            r'\d{4}-\d{2}-\d{2}',
            r'\d{1,2}[/-]\d{1,2}[/-]\d{4}',
            r'\d{4}',  # Just year
        ]
        
        for pattern in date_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                date_str = match.group()
                
                # Use dateparser for robust parsing
                try:
                    parsed_date = dateparser.parse(
                        date_str,
                        settings={
                            'PREFER_DATES_FROM': 'past',
                            'RETURN_AS_TIMEZONE_AWARE': False,
                            'DATE_ORDER': 'DMY'  # European date order
                        }
                    )
                    
                    if parsed_date:
                        # Filter out unreasonable dates
                        current_year = datetime.now().year
                        if 1990 <= parsed_date.year <= (current_year + 5 if search_future else current_year):
                            found_dates.append((parsed_date, date_str))
                except:
                    continue
        
        # Remove duplicates and sort
        found_dates = list(set(found_dates))
        found_dates.sort(key=lambda x: x[0])
        
        return found_dates
    
    @staticmethod
    def find_date_ranges(text: str) -> List[Tuple[datetime, datetime, str]]:
        """
        Find date ranges in text (e.g., "from X to Y", "X through Y").
        
        Returns:
            List of (start_date, end_date, original_text) tuples
        """
        date_ranges = []
        
        # Look for range patterns
        range_patterns = [
            r'from\s+(.+?)\s+(?:to|through|until)\s+(.+?)(?:\.|,|\s+and|\s+include)',
            r'between\s+(.+?)\s+and\s+(.+?)(?:\.|,)',
            r'period[:\s]+(.+?)\s+(?:to|-)\s+(.+?)(?:\.|,)',
            r'(\d{1,2}\s+\w+\s+\d{4})\s+(?:to|through|-)\s+(\d{1,2}\s+\w+\s+\d{4})',
        ]
        
        for pattern in range_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    date1_str = match.group(1).strip()
                    date2_str = match.group(2).strip()
                    
                    date1 = dateparser.parse(date1_str, settings={'PREFER_DATES_FROM': 'past'})
                    date2 = dateparser.parse(date2_str, settings={'PREFER_DATES_FROM': 'past'})
                    
                    if date1 and date2:
                        # Ensure start is before end
                        if date1 > date2:
                            date1, date2 = date2, date1
                        
                        date_ranges.append((date1, date2, match.group(0)))
                except:
                    continue
        
        return date_ranges
    
    @staticmethod
    def extract_numbers_with_context(text: str, keywords: List[str]) -> Dict[str, List[int]]:
        """
        Extract numbers near specific keywords (e.g., "articles", "publications").
        
        Returns:
            Dict mapping keywords to lists of associated numbers
        """
        results = {}
        
        for keyword in keywords:
            numbers = []
            keyword_lower = keyword.lower()
            text_lower = text.lower()
            
            # Find keyword positions
            pos = 0
            while True:
                pos = text_lower.find(keyword_lower, pos)
                if pos == -1:
                    break
                
                # Look for numbers within 50 characters before or after
                context_start = max(0, pos - 50)
                context_end = min(len(text), pos + len(keyword) + 50)
                context = text[context_start:context_end]
                
                # Find all numbers in context
                number_matches = re.findall(r'\b(\d+(?:,\d+)*)\b', context)
                for num_str in number_matches:
                    try:
                        # Remove commas and convert
                        num = int(num_str.replace(',', ''))
                        if 0 < num < 100000:  # Reasonable range for article counts
                            numbers.append(num)
                    except:
                        continue
                
                pos += len(keyword)
            
            if numbers:
                results[keyword] = numbers
        
        return results
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean text by removing excessive whitespace and special characters"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special Unicode characters that cause encoding issues
        text = text.encode('ascii', 'ignore').decode('ascii')
        return text.strip()
    
    @staticmethod
    def get_semantic_parser():
        """
        Get semantic parser instance for advanced document understanding.
        
        Returns:
            SemanticDocumentParser instance
            
        Raises:
            ImportError: If semantic parser dependencies not available
            ValueError: If ANTHROPIC_API_KEY not set
        """
        try:
            from .semantic_document_parser import SemanticDocumentParser
            return SemanticDocumentParser()
        except ImportError as e:
            raise ImportError(
                "Semantic parser not available. Ensure anthropic package is installed: "
                "pip install anthropic"
            ) from e
    
    @staticmethod
    def extract_with_semantic_understanding(
        docx_path: str,
        use_semantic: bool = True,
        use_cache: bool = True
    ) -> Any:
        """
        Extract document with optional semantic processing.
        
        Args:
            docx_path: Path to .docx file
            use_semantic: Use Anthropic Claude for semantic understanding
            use_cache: Use cached semantic data if available
            
        Returns:
            If use_semantic=True: CERData object with full semantic understanding
            If use_semantic=False: Dict with basic extraction (legacy mode)
        """
        if use_semantic:
            try:
                from .semantic_document_parser import get_semantic_cer_data
                return get_semantic_cer_data(docx_path, force_refresh=not use_cache)
            except Exception as e:
                print(f"  Warning: Semantic parsing failed, falling back to basic extraction: {e}")
                return DocumentParser.extract_text_with_structure(docx_path)
        else:
            return DocumentParser.extract_text_with_structure(docx_path)
    
    @staticmethod
    def upgrade_to_semantic_parser() -> bool:
        """
        Check if semantic parser upgrade is available.
        
        Returns:
            True if semantic parser can be used, False otherwise
        """
        try:
            # Check for Anthropic API key
            if not os.environ.get("ANTHROPIC_API_KEY"):
                print("  Semantic parser requires ANTHROPIC_API_KEY environment variable")
                return False
            
            # Try importing semantic parser
            from .semantic_document_parser import SemanticDocumentParser
            print("  ✓ Semantic parser available")
            return True
            
        except ImportError:
            print("  Semantic parser not installed. Run: pip install anthropic")
            return False

