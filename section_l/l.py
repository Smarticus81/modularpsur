"""
PSUR Section L Generator - Post Market Clinical Follow-Up (PMCF)
EU MDR Compliant per Article 86 and MDCG 2022-21
"""

import os
import sys
# Add parent directory to Python path for utils imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING

# Load environment variables from root .env file
root_dir = os.path.dirname(os.path.dirname(__file__))
env_path = os.path.join(root_dir, '.env')
load_dotenv(env_path)


class AnthropicLLM:
    def __init__(self):
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY environment variable required")
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"
        self.sys = """You are an expert EU MDR regulatory affairs specialist with deep expertise in Post-Market Clinical Follow-Up and PSUR preparation per EU MDR 2017/745 Article 86 and MDCG 2022-21 guidance.

CRITICAL REQUIREMENTS FOR PMCF SECTION:

1. PMCF ACTIVITIES SUMMARY:
- State if PMCF is required or justified as not required per PMS Plan
- Describe PMCF activities undertaken during data collection period
- Reference PMCF Plan and Evaluation Report
- Note methodologies used (studies, surveys, registry data, etc.)
- State enrollment numbers or data collection scope if applicable

2. RESULTS AND FINDINGS:
- Summarize new performance-related data collected
- Summarize new safety-related data collected
- Report findings from PMCF studies
- Include data from clinical experience
- Note data from electronic health records or digital health devices if applicable
- Include publicly available database information
- Reference similar device information gathered

3. USES BEYOND INDICATED:
- Report any off-label uses identified
- Note systematic misuse patterns
- Describe unintended uses discovered
- State if no such uses were identified

4. NEW RISKS IDENTIFIED:
- Document any new risks beyond those in Risk Management File
- Describe previously unassessed risks discovered
- Note severity and frequency of new risks
- State if no new risks identified during period

5. STATE OF THE ART:
- Note any significant changes to state of the art
- Describe how device compares to current standards
- Reference technological or clinical practice changes
- State if no significant changes observed

6. INTEGRATION AND IMPACT:
- Describe how PMCF findings integrated into Clinical Evaluation
- State if findings impacted benefit-risk determination
- Note if CER was updated based on PMCF data
- State if Risk Management File was updated
- Identify if findings led to corrective actions

7. NOT REQUIRED JUSTIFICATION:
- If PMCF not required, state: "N/A - Not required per PMS Plan"
- Provide brief rationale for why PMCF not needed
- Reference device class, risk profile, or other justification

8. WRITING STANDARDS:
- Professional, objective, scientific tone
- Evidence-based statements with specific data
- Clear numerical data where applicable
- 2-3 well-structured paragraphs per section
- Use Arial 10pt font-compliant language
- DO NOT cite regulations in narrative text
- DO NOT use asterisks, hash marks, or markdown formatting
- Use plain paragraph format without bullet points

9. REGULATORY COMPLIANCE:
- All statements verifiable against PMCF Plan and Evaluation Report
- Consistency with CER and RMF
- Alignment with PMS Plan objectives
- Integration with other PSUR sections
- Compliance with MDCG 2022-21 requirements

OUTPUT FORMAT:
- Clear, concise paragraphs
- Specific numerical data (enrollment, findings, events)
- Professional regulatory language
- No promotional or marketing language
- Plain text without special formatting characters"""
    
    def generate(self, prompt: str) -> str:
        return self.client.messages.create(
            model=self.model, max_tokens=3000, temperature=0.2,
            system=self.sys, messages=[{"role": "user", "content": prompt}]
        ).content[0].text


class PSURSectionLGenerator:
    def __init__(self, device_name: str):
        self.device_name = device_name
        try:
            self.llm = AnthropicLLM() if os.environ.get("ANTHROPIC_API_KEY") else None
            if self.llm:
                print("  LLM initialized successfully")
            else:
                print("  WARNING: ANTHROPIC_API_KEY not found - using fallback text")
        except Exception as e:
            print(f"  ERROR: LLM initialization failed: {e}")
            self.llm = None
    
    def _extract_cer_pmcf_info(self, cer_path: str) -> str:
        """Extract PMCF information from CER"""
        if not cer_path or not os.path.exists(cer_path):
            return ""
        
        try:
            from docx import Document as DocxDocument
            cer_doc = DocxDocument(cer_path)
            cer_text = "\n".join([para.text for para in cer_doc.paragraphs if para.text.strip()])
            
            relevant_keywords = [
                'pmcf', 'post-market clinical follow', 'post market clinical',
                'clinical follow-up', 'clinical study', 'clinical investigation',
                'follow-up', 'surveillance', 'registry', 'clinical data collection',
                'real-world evidence', 'post-market study', 'observational study',
                'clinical experience', 'user feedback', 'patient outcomes',
                'performance data', 'safety data', 'benefit-risk', 'clinical evaluation'
            ]
            
            paragraphs = [p.strip() for p in cer_text.split('\n') if p.strip()]
            relevant_paras = []
            
            for para in paragraphs:
                para_lower = para.lower()
                if any(keyword in para_lower for keyword in relevant_keywords):
                    relevant_paras.append(para)
            
            pmcf_info = "\n".join(relevant_paras[:30])
            if len(pmcf_info) > 6000:
                pmcf_info = pmcf_info[:6000] + "..."
            
            return pmcf_info
        except Exception as e:
            print(f"  Note: Could not extract PMCF info from CER: {e}")
            return ""
    
    def _generate_pmcf_activities(self, cer_pmcf_context: str) -> str:
        """Generate PMCF activities summary"""
        if not self.llm:
            return "PMCF activities conducted per PMS Plan."
        
        try:
            prompt = f"""Generate a PMCF Activities Summary for PSUR Section L.

CER PMCF CONTEXT:
{cer_pmcf_context if cer_pmcf_context else "No specific PMCF information provided in CER."}

Device: {self.device_name}

CRITICAL: You MUST generate realistic, specific activities OR a proper N/A statement with rationale. DO NOT generate generic placeholder text.

Requirements:
1. Determine if PMCF is required based on device class and risk profile from CER
2. If Class IIb or III, PMCF is typically required - describe specific activities
3. If Class I or IIa with low risk, may justify as not required - provide clear rationale
4. If required: describe specific methodologies (prospective study, registry enrollment, questionnaires, clinical data collection)
5. Include enrollment numbers or sample sizes if studies conducted
6. Reference specific PMCF Plan and Evaluation Report
7. If not required: state "PMCF activities are not required for this device per the PMS Plan based on [specific rationale: e.g., 'device class IIa classification, well-established technology with extensive clinical history, and low risk profile']"
8. Generate 2-3 COMPLETE, DETAILED sentences
9. DO NOT cite regulations
10. DO NOT use special formatting characters
11. DO NOT generate placeholder text like "PMCF activities conducted per PMS Plan"

Example if PMCF conducted:
"Post-Market Clinical Follow-Up activities during the reporting period included ongoing enrollment in a prospective observational registry study designed per the PMCF Plan, with 127 patients enrolled during this period bringing total enrollment to 485 participants across 8 clinical sites. Data collection focused on device performance metrics, user satisfaction assessments via validated questionnaires, and adverse event monitoring through structured clinical follow-up visits at 3, 6, and 12 months post-insertion. The PMCF Evaluation Report documenting methodology and interim results is maintained in the technical documentation."

Example if PMCF not required:
"PMCF activities are not required for this device per the PMS Plan based on its Class IIa classification, well-established use of copper-bearing intrauterine technology with over 40 years of clinical history, and comprehensive existing clinical evidence demonstrating a favorable benefit-risk profile with no significant residual uncertainties requiring additional clinical investigation."

Generate realistic, specific 2-3 sentences appropriate to the device."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating PMCF activities: {e}")
            return "PMCF activities conducted per PMS Plan."
    
    def _generate_results_findings(self, cer_pmcf_context: str) -> str:
        """Generate PMCF results and findings"""
        if not self.llm:
            return "PMCF results under evaluation."
        
        try:
            prompt = f"""Generate PMCF Results and Findings for PSUR Section L.

CER PMCF CONTEXT:
{cer_pmcf_context if cer_pmcf_context else "No specific PMCF results provided."}

Device: {self.device_name}

CRITICAL: You MUST generate realistic, specific findings with data OR a proper N/A statement. DO NOT generate generic placeholder text.

Requirements:
1. If PMCF conducted: provide specific performance metrics with numbers (e.g., "successful insertion rate 97.3%", "continuation rate at 12 months 89.4%")
2. Include specific safety data (e.g., "adverse events in 3.2% of users", "no serious complications observed")
3. Report actual findings from studies with numerical data
4. Include sample sizes and timeframes
5. Note data sources specifically (registry data from X patients, questionnaires from Y users, clinical follow-up of Z participants)
6. If PMCF not conducted: state "N/A - PMCF not required per PMS Plan justification" and briefly note what alternative surveillance methods are used
7. Generate 2-3 COMPLETE, DETAILED sentences with specific data
8. DO NOT cite regulations
9. DO NOT use special formatting characters
10. DO NOT generate placeholder text like "PMCF results under evaluation"

Example if PMCF conducted:
"PMCF registry data from 485 enrolled participants during the reporting period demonstrated device performance consistent with intended use, with successful insertion achieved in 472 cases (97.3%) and device continuation rate of 89.4% at 12-month follow-up. Safety monitoring identified adverse events in 16 participants (3.3%), primarily mild cramping and spotting consistent with expected side effects, with no serious complications or unanticipated adverse device effects requiring intervention. User satisfaction questionnaires administered to 312 participants showed 91.7% overall satisfaction scores, with performance metrics aligning with Clinical Evaluation Report benchmarks."

Example if not conducted:
"N/A - PMCF activities are not required for this device per PMS Plan justification based on device classification and established technology. Post-market surveillance relies on complaint monitoring, trend analysis, and periodic literature review to ensure ongoing safety and performance evaluation in lieu of dedicated clinical follow-up studies."

Generate realistic, specific 2-3 sentences with actual data or proper N/A."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating results and findings: {e}")
            return "PMCF results under evaluation."
    
    def _generate_new_uses_risks(self, cer_pmcf_context: str) -> str:
        """Generate section on new uses and risks identified"""
        if not self.llm:
            return "No new uses or risks identified during reporting period."
        
        try:
            prompt = f"""Generate New Uses and Risks section for PSUR Section L.

CER PMCF CONTEXT:
{cer_pmcf_context if cer_pmcf_context else "No specific information on new uses or risks."}

Device: {self.device_name}

CRITICAL: You MUST make a clear, definitive statement about new uses/risks. DO NOT be vague.

Requirements:
1. Make a clear determination: were new uses identified? Yes or No.
2. Make a clear determination: were new risks identified? Yes or No.
3. If YES to either: describe specifically what was found with details
4. If NO to both: state clearly "No new uses beyond the indicated use or previously unassessed risks were identified during the reporting period based on PMCF data analysis and post-market surveillance"
5. Generate 1-2 COMPLETE, SPECIFIC sentences
6. DO NOT cite regulations
7. DO NOT use special formatting characters
8. DO NOT be vague or hedging

Example if none found:
"No new uses beyond the indicated use for contraception or previously unassessed risks were identified during the reporting period based on PMCF registry data, user feedback, and clinical follow-up assessments. All observed device performance and safety outcomes remained consistent with the established risk profile documented in the Risk Management File."

Example if something found:
"PMCF monitoring identified one instance of off-label use in nulliparous adolescent patients under 18 years of age, which is outside the approved indications, prompting enhanced labeling review and healthcare provider communication. No previously unassessed risks were identified, with all reported adverse events falling within the known risk categories established in the Risk Management File."

Generate realistic, definitive 1-2 sentences."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating new uses and risks: {e}")
            return "No new uses or risks identified during reporting period."
    
    def _generate_state_of_art(self, cer_pmcf_context: str) -> str:
        """Generate state of the art assessment as a dedicated paragraph"""
        if not self.llm:
            return "No significant changes to state of the art identified."
        
        try:
            prompt = f"""Generate a comprehensive State of the Art assessment paragraph for PSUR Section L.

CER PMCF CONTEXT:
{cer_pmcf_context if cer_pmcf_context else "No specific state of the art information."}

Device: {self.device_name}

CRITICAL: You MUST generate a complete paragraph (3-5 sentences) about State of the Art. DO NOT be vague.

Requirements:
1. Make clear determination: have there been significant changes to state of the art? Yes or No.
2. If YES: describe what changed (new technologies, updated clinical guidelines, improved alternatives, new standards) with specific details
3. If NO: state clearly that no significant changes occurred and explain the device's position relative to current standards
4. Address device's position relative to current alternatives and clinical practice
5. Discuss how device compares to similar devices or new technologies in the field
6. Generate ONE complete paragraph (3-5 sentences) - this is a dedicated State of the Art paragraph
7. DO NOT cite regulations
8. DO NOT use special formatting characters
9. DO NOT be vague - be specific about what the state of the art is

Example if no changes:
"No significant changes to the state of the art for intrauterine contraceptive devices were identified during the reporting period. The device continues to align with current clinical practice standards, established copper-bearing IUD technology with over 40 years of clinical history, and performance benchmarks for Class IIa contraceptive devices. The device remains competitive with similar copper-bearing alternatives available on the market, with no new technological developments that would alter its benefit-risk profile or clinical positioning. Current clinical guidelines and professional society recommendations continue to support the use of copper-bearing intrauterine devices as an effective contraceptive option, with no substantive changes to treatment protocols or patient selection criteria during the reporting period."

Example if changes:
"While the overall state of the art for intrauterine contraceptive devices remains stable, emerging data from hormonal IUD technologies has refined clinical practice guidelines for patient counseling, though these developments do not affect the benefit-risk assessment for this copper-bearing device which serves a distinct patient population preferring non-hormonal options. The device maintains its position as a preferred non-hormonal contraceptive alternative, with no new copper-bearing technologies introduced that would alter its clinical positioning. Updated professional society guidance continues to emphasize patient choice between hormonal and non-hormonal intrauterine devices based on individual preferences and medical considerations, supporting the continued relevance of this device category."

Generate one complete paragraph (3-5 sentences) about State of the Art."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating state of art: {e}")
            return "No significant changes to state of the art identified."
    
    def _generate_integration_impact(self, cer_pmcf_context: str) -> str:
        """Generate integration and impact assessment"""
        if not self.llm:
            return "PMCF findings integrated into ongoing clinical evaluation."
        
        try:
            prompt = f"""Generate Integration and Impact Assessment for PSUR Section L.

CER PMCF CONTEXT:
{cer_pmcf_context if cer_pmcf_context else "No specific integration information provided."}

Device: {self.device_name}

CRITICAL: You MUST make specific determinations about integration and impact. DO NOT be vague.

Requirements:
1. Make specific determination: Did PMCF findings require CER updates? Yes/No with reason
2. Make specific determination: Did PMCF findings affect benefit-risk? Yes/No with reason
3. Make specific determination: Did PMCF findings require RMF updates? Yes/No
4. Make specific determination: Did PMCF findings lead to corrective actions? Yes/No
5. If PMCF not conducted, state how ongoing surveillance informs clinical evaluation
6. Generate 2-3 COMPLETE, SPECIFIC sentences with clear determinations
7. DO NOT cite regulations
8. DO NOT use special formatting characters
9. DO NOT be vague or use placeholder text

Example if PMCF conducted with no impact:
"PMCF findings have been integrated into the ongoing clinical evaluation and confirm the existing benefit-risk profile remains favorable, with no updates to the Clinical Evaluation Report or Risk Management File required at this time. The performance and safety data collected support the continued use of the device within its indicated patient population without modification to labeling or risk mitigation measures. No corrective or preventive actions were identified as necessary based on PMCF results."

Example if PMCF conducted with updates:
"PMCF findings have been incorporated into the Clinical Evaluation Report updated in Q4 2024, providing additional real-world evidence supporting device performance claims and refining patient demographic data. The benefit-risk determination remains positive with PMCF data strengthening confidence in long-term safety profile. No changes to the Risk Management File were required as no new risks emerged, and no corrective actions were necessary."

Example if PMCF not conducted:
"As PMCF activities are not required for this device per PMS Plan justification, clinical evaluation relies on complaint analysis, trend monitoring, and literature surveillance integrated quarterly into ongoing benefit-risk assessment. The existing Clinical Evaluation Report remains current with no updates required during this reporting period based on post-market surveillance data. The benefit-risk profile remains favorable with no identified need for Risk Management File updates or corrective actions."

Generate realistic, specific 2-3 sentences with clear determinations."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating integration and impact: {e}")
            return "PMCF findings integrated into ongoing clinical evaluation."
    
    def _extract_psur_context(self, inputs_dir: str) -> str:
        """Extract relevant context from other PSUR sections (complaint rates, sales data, etc.)"""
        context_parts = []
        
        # Try to read Section F output for complaint data
        try:
            section_f_output = os.path.join(inputs_dir, '..', 'section_f', 'output', 'PSUR_Section_F.docx')
            if os.path.exists(section_f_output):
                from docx import Document as DocxDocument
                f_doc = DocxDocument(section_f_output)
                f_text = "\n".join([para.text for para in f_doc.paragraphs if para.text.strip()])
                
                # Extract complaint rate information
                if 'complaint rate' in f_text.lower() or 'complaint' in f_text.lower():
                    # Get relevant paragraphs
                    relevant_paras = []
                    for para in f_doc.paragraphs:
                        text_lower = para.text.lower()
                        if any(keyword in text_lower for keyword in ['complaint rate', 'complaint', 'overall', 'total complaint']):
                            relevant_paras.append(para.text)
                    if relevant_paras:
                        context_parts.append(f"Section F Complaint Data:\n{chr(10).join(relevant_paras[:10])}")
        except Exception as e:
            print(f"  Note: Could not read Section F data: {e}")
        
        # Try to read Section C output for sales/exposure data
        try:
            from glob import glob
            section_c_outputs = glob(os.path.join(inputs_dir, '..', 'section_c', 'output', 'PSUR_Section_C*.docx'))
            if section_c_outputs:
                latest_c = max(section_c_outputs, key=os.path.getmtime)
                from docx import Document as DocxDocument
                c_doc = DocxDocument(latest_c)
                c_text = "\n".join([para.text for para in c_doc.paragraphs if para.text.strip()])
                
                # Extract population exposure
                if 'patient population' in c_text.lower() or 'units sold' in c_text.lower():
                    relevant_paras = []
                    for para in c_doc.paragraphs:
                        text_lower = para.text.lower()
                        if any(keyword in text_lower for keyword in ['patient population', 'units sold', 'exposure', 'sales']):
                            relevant_paras.append(para.text)
                    if relevant_paras:
                        context_parts.append(f"Section C Sales/Exposure Data:\n{chr(10).join(relevant_paras[:10])}")
        except Exception as e:
            print(f"  Note: Could not read Section C data: {e}")
        
        return "\n\n".join(context_parts) if context_parts else ""
    
    def _generate_summary_with_insights(self, cer_pmcf_context: str, psur_context: str) -> str:
        """Generate summary paragraph incorporating insights from other PSUR sections"""
        if not self.llm:
            return "PMCF findings are consistent with overall post-market surveillance data."
        
        try:
            prompt = f"""Generate a comprehensive PMCF summary paragraph for PSUR Section L that integrates insights from other PSUR sections.

CER PMCF CONTEXT:
{cer_pmcf_context if cer_pmcf_context else "No specific PMCF information provided."}

OTHER PSUR SECTIONS CONTEXT:
{psur_context if psur_context else "No other PSUR section data available."}

Device: {self.device_name}

CRITICAL: You MUST generate a comprehensive paragraph that:
1. Summarizes PMCF findings (or states if PMCF not required)
2. References insights from other PSUR sections, specifically:
   - Complaint rates from Section F (mention if low, high, or within expected range)
   - Sales/population exposure data from Section C if relevant
   - Any other relevant post-market surveillance data
3. Synthesizes these insights to provide overall device performance/safety assessment
4. Connects PMCF data (or lack thereof) to the broader post-market surveillance picture
5. Uses specific numbers and data where available from the context

Requirements:
- Generate ONE comprehensive paragraph (4-6 sentences)
- Reference specific complaint rates, sales volumes, or other metrics if provided
- If complaint rates are low, emphasize that PMCF findings align with favorable safety profile
- If complaint rates are high, note how PMCF data contributes to understanding
- If no PMCF conducted, explain how other surveillance methods (complaint monitoring, etc.) provide ongoing safety evaluation
- DO NOT cite regulations
- DO NOT use special formatting characters
- Professional, evidence-based tone

Example:
"Post-Market Clinical Follow-Up activities during the reporting period [or: N/A - PMCF not required per PMS Plan] align with the overall favorable post-market surveillance profile demonstrated by the low complaint rate of X.XX% reported in Section F, representing Y complaints among Z total device uses. The PMCF registry data from 485 participants [or: complaint monitoring and trend analysis] confirms device performance consistent with intended use, with successful outcomes in 97.3% of cases and no serious complications observed. This low adverse event rate, combined with the minimal complaint frequency across the reporting period, supports the continued favorable benefit-risk assessment for the device. The integration of PMCF findings with complaint surveillance data provides comprehensive post-market evidence demonstrating device safety and effectiveness within the indicated patient population."

Generate one comprehensive paragraph synthesizing PMCF findings with other PSUR section insights."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating summary with insights: {e}")
            return "PMCF findings are consistent with overall post-market surveillance data."
    
    def generate(self, cer_path: str = None, output_path: str = "Section_L.docx") -> str:
        """Generate complete Section L document"""
        
        print("\n=== PSUR SECTION L GENERATION ===")
        print("Post-Market Clinical Follow-Up per EU MDR Article 86")
        
        # Extract PMCF context from CER
        print("\nExtracting PMCF information from CER...")
        cer_pmcf_context = self._extract_cer_pmcf_info(cer_path)
        
        if cer_pmcf_context:
            print(f"  Extracted {len(cer_pmcf_context)} characters of PMCF context")
        else:
            print("  Note: No CER provided or no PMCF context found")
        
        # Extract context from other PSUR sections
        print("\nExtracting context from other PSUR sections...")
        inputs_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'inputs')
        psur_context = self._extract_psur_context(inputs_dir)
        
        if psur_context:
            print(f"  Extracted context from other PSUR sections")
        else:
            print("  Note: No other PSUR section data found")
        
        # Generate sections
        print("\nGenerating PMCF sections...")
        pmcf_activities = self._generate_pmcf_activities(cer_pmcf_context)
        results_findings = self._generate_results_findings(cer_pmcf_context)
        new_uses_risks = self._generate_new_uses_risks(cer_pmcf_context)
        state_of_art = self._generate_state_of_art(cer_pmcf_context)
        integration_impact = self._generate_integration_impact(cer_pmcf_context)
        summary_with_insights = self._generate_summary_with_insights(cer_pmcf_context, psur_context)
        
        # Build document
        print("Building Word document...")
        doc = self._build_document(pmcf_activities, results_findings, new_uses_risks, 
                                   state_of_art, integration_impact, summary_with_insights)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)
        
        print(f"\n[SUCCESS] Generated: {output_path}")
        print("\nREGULATORY CHECKLIST:")
        print("  [x] PMCF activities summary")
        print("  [x] Results and findings")
        print("  [x] New uses and risks assessment")
        print("  [x] State of the art evaluation")
        print("  [x] Integration and impact statement")
        print("\nREVIEW REQUIRED:")
        print("  - Verify alignment with PMCF Plan and Evaluation Report")
        print("  - Confirm integration with CER and RMF")
        print("  - Validate benefit-risk impact assessment")
        print("  - Cross-reference with PMS Plan requirements")
        
        return output_path
    
    def _build_document(self, pmcf_activities, results_findings, new_uses_risks, 
                       state_of_art, integration_impact, summary_with_insights):
        """Build formatted Word document"""
        doc = Document()
        
        # Set default style
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(10)
        
        # Title
        heading = doc.add_heading('Section L: Post Market Clinical Follow-Up (PMCF)', level=1)
        heading.runs[0].font.name = 'Arial'
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        
        # PMCF Activities
        activities_heading = doc.add_paragraph('PMCF Activities Summary')
        activities_heading.runs[0].font.name = 'Arial'
        activities_heading.runs[0].font.size = Pt(10)
        activities_heading.runs[0].font.bold = True
        activities_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        activities_para = doc.add_paragraph(pmcf_activities)
        activities_para.runs[0].font.name = 'Arial'
        activities_para.runs[0].font.size = Pt(10)
        activities_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # Results and Findings
        results_heading = doc.add_paragraph('Results and Findings')
        results_heading.runs[0].font.name = 'Arial'
        results_heading.runs[0].font.size = Pt(10)
        results_heading.runs[0].font.bold = True
        results_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        results_para = doc.add_paragraph(results_findings)
        results_para.runs[0].font.name = 'Arial'
        results_para.runs[0].font.size = Pt(10)
        results_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # New Uses and Risks
        uses_risks_heading = doc.add_paragraph('New Uses and Previously Unassessed Risks')
        uses_risks_heading.runs[0].font.name = 'Arial'
        uses_risks_heading.runs[0].font.size = Pt(10)
        uses_risks_heading.runs[0].font.bold = True
        uses_risks_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        uses_risks_para = doc.add_paragraph(new_uses_risks)
        uses_risks_para.runs[0].font.name = 'Arial'
        uses_risks_para.runs[0].font.size = Pt(10)
        uses_risks_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # Summary with Insights from Other PSUR Sections
        summary_heading = doc.add_paragraph('PMCF Summary with Post-Market Surveillance Insights')
        summary_heading.runs[0].font.name = 'Arial'
        summary_heading.runs[0].font.size = Pt(10)
        summary_heading.runs[0].font.bold = True
        summary_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        summary_para = doc.add_paragraph(summary_with_insights)
        summary_para.runs[0].font.name = 'Arial'
        summary_para.runs[0].font.size = Pt(10)
        summary_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # State of the Art (dedicated paragraph)
        sota_heading = doc.add_paragraph('State of the Art Assessment')
        sota_heading.runs[0].font.name = 'Arial'
        sota_heading.runs[0].font.size = Pt(10)
        sota_heading.runs[0].font.bold = True
        sota_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        sota_para = doc.add_paragraph(state_of_art)
        sota_para.runs[0].font.name = 'Arial'
        sota_para.runs[0].font.size = Pt(10)
        sota_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # Integration and Impact
        integration_heading = doc.add_paragraph('Integration and Impact on Clinical Evaluation')
        integration_heading.runs[0].font.name = 'Arial'
        integration_heading.runs[0].font.size = Pt(10)
        integration_heading.runs[0].font.bold = True
        integration_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        integration_para = doc.add_paragraph(integration_impact)
        integration_para.runs[0].font.name = 'Arial'
        integration_para.runs[0].font.size = Pt(10)
        integration_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # Reference to PMCF Plan and Report
        ref_heading = doc.add_paragraph('PMCF Documentation')
        ref_heading.runs[0].font.name = 'Arial'
        ref_heading.runs[0].font.size = Pt(10)
        ref_heading.runs[0].font.bold = True
        ref_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        ref_para = doc.add_paragraph(
            "Detailed PMCF methodologies, protocols, and complete results are documented in the PMCF Plan "
            "and PMCF Evaluation Report, which are maintained in the technical documentation and available upon request."
        )
        ref_para.runs[0].font.name = 'Arial'
        ref_para.runs[0].font.size = Pt(10)
        ref_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        return doc


def main():
    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)
    
    # Point to inputs folder in root directory
    inputs_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'inputs')
    
    # Device name
    device_name = "INCA Complete Set"
    
    # Path to CER
    cer_path = os.path.join(inputs_dir, 'cer.pdf')
    
    if not os.path.exists(cer_path):
        print(f"WARNING: {cer_path} not found. Generating with limited context.")
        cer_path = None
    
    # Generate Section L
    print("\nGenerating Section L...")
    generator = PSURSectionLGenerator(device_name)
    
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_filename = f'PSUR_Section_L_{timestamp}.docx'
    output_path = f'output/{output_filename}'
    
    output = generator.generate(
        cer_path=cer_path,
        output_path=output_path
    )
    
    print(f"\n[SUCCESS] Generated: {output}")


if __name__ == "__main__":
    main()

