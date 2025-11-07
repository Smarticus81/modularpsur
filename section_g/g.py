#!/usr/bin/env python3
"""
Section G Generator for PSUR
Automatically generates Section G (Trend Reporting) from sales and complaints data
with LLM-powered analysis and narrative generation.

Requirements:
- pandas, openpyxl, matplotlib
- anthropic (pip install anthropic)
- python-docx
"""

import os
import sys
# Add parent directory to Python path for utils imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pandas as pd
import numpy as np
from datetime import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.dates import MonthLocator, DateFormatter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import anthropic
from dotenv import load_dotenv

# Load environment variables from root .env file
root_dir = os.path.dirname(os.path.dirname(__file__))
env_path = os.path.join(root_dir, '.env')
load_dotenv(env_path)

# Configuration
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')  # Set your API key here or as environment variable
MODEL = "claude-sonnet-4-20250514"

def load_and_process_data(sales_file, complaints_file, start_date, end_date):
    """Load sales and complaints data, process and merge."""
    
    # Read sales data
    sales_df = pd.read_excel(sales_file)
    
    print(f"  Sales columns: {list(sales_df.columns)[:8]}...")
    
    # Flexible column detection
    month_col = year_col = quantity_col = None
    for col in sales_df.columns:
        col_lower = col.lower()
        if 'month' in col_lower and not month_col:
            month_col = col
        if 'year' in col_lower and not year_col:
            year_col = col
        if ('quantity' in col_lower or 'units' in col_lower) and not quantity_col:
            quantity_col = col
    
    if not all([month_col, year_col, quantity_col]):
        raise ValueError(f"Missing columns: Month={month_col}, Year={year_col}, Qty={quantity_col}")
    
    print(f"  Using: {month_col}, {year_col}, {quantity_col}")
    
    # Create datetime from Month name and year column
    month_map = {
        'January': 1, 'February': 2, 'March': 3, 'April': 4, 
        'May': 5, 'June': 6, 'July': 7, 'August': 8,
        'September': 9, 'October': 10, 'November': 11, 'December': 12,
        'Jan': 1, 'Feb': 2, 'Mar': 3, 'Apr': 4, 'May': 5, 'Jun': 6,
        'Jul': 7, 'Aug': 8, 'Sep': 9, 'Oct': 10, 'Nov': 11, 'Dec': 12
    }
    
    # Handle data structure: Month (name) + year columns
    sales_df['Month_Num'] = sales_df[month_col].map(month_map)
    if sales_df['Month_Num'].isna().any():
        try:
            sales_df['Month_Num'] = sales_df['Month_Num'].fillna(pd.to_numeric(sales_df[month_col], errors='coerce'))
        except:
            pass
    
    sales_df['DateTime'] = pd.to_datetime(
        sales_df[[year_col, 'Month_Num']].rename(
            columns={year_col: 'year', 'Month_Num': 'month'}
        ).assign(day=1)
    )
    
    # Filter sales data by date range
    sales_filtered = sales_df[
        (sales_df['DateTime'] >= start_date) & 
        (sales_df['DateTime'] <= end_date)
    ].copy()
    
    # Aggregate monthly sales
    monthly_sales = sales_filtered.groupby('DateTime')[quantity_col].sum().reset_index()
    monthly_sales = monthly_sales.rename(columns={'DateTime': 'Month', quantity_col: 'Quantity'}).sort_values('Month')
    
    # Read and process complaints - flexible sheet detection
    try:
        complaints_df = pd.read_excel(complaints_file, sheet_name='CSI Complaints')
    except:
        complaints_df = pd.read_excel(complaints_file, sheet_name=0)
    
    print(f"  Complaints columns: {list(complaints_df.columns)[:8]}...")
    
    # Flexible date column detection
    date_col = None
    for col in complaints_df.columns:
        col_lower = col.lower()
        if 'date' in col_lower and 'entered' in col_lower:
            date_col = col
            break
        elif 'date' in col_lower and not date_col:
            date_col = col
    
    if not date_col:
        raise ValueError("No date column found in complaints")
    
    print(f"  Using date column: {date_col}")
    
    # Convert date to datetime
    complaints_df['Date_Entered'] = pd.to_datetime(complaints_df[date_col], errors='coerce')
    
    # Filter complaints by date range
    complaints_filtered = complaints_df[
        (complaints_df['Date_Entered'] >= start_date) & 
        (complaints_df['Date_Entered'] <= end_date)
    ].copy()
    
    # Extract year-month for grouping
    complaints_filtered['Month'] = complaints_filtered['Date_Entered'].dt.to_period('M').dt.to_timestamp()
    
    # Count complaints per month
    monthly_complaints = complaints_filtered.groupby('Month').size().reset_index(name='Complaint_Count')
    
    # Merge sales and complaints data
    merged_data = pd.merge(monthly_sales, monthly_complaints, on='Month', how='left')
    merged_data['Complaint_Count'] = merged_data['Complaint_Count'].fillna(0).astype(int)
    
    # Calculate complaint rate as percentage (complaints/units sold * 100)
    merged_data['Complaint_Rate'] = merged_data.apply(
        lambda row: (row['Complaint_Count'] / row['Quantity']) * 100 if row['Quantity'] > 0 else 0,
        axis=1
    )
    
    return merged_data

def calculate_control_limits(merged_data):
    """Calculate statistical control limits."""
    mean_rate = merged_data['Complaint_Rate'].mean()
    std_rate = merged_data['Complaint_Rate'].std()
    ucl = mean_rate + (3 * std_rate)
    lcl = max(0, mean_rate - (3 * std_rate))
    
    return mean_rate, std_rate, ucl, lcl

def identify_breaches(merged_data, ucl):
    """Identify UCL breaches."""
    breaches = merged_data[merged_data['Complaint_Rate'] > ucl].copy()
    return breaches

def create_trend_chart(merged_data, mean_rate, ucl, lcl, breaches, product_name, start_date, end_date, output_path):
    """Generate the complaint rate trend chart."""
    fig, ax = plt.subplots(figsize=(15, 7))
    
    # Plot complaint rate
    ax.plot(merged_data['Month'], merged_data['Complaint_Rate'], 
            marker='o', linewidth=2.5, markersize=6, color='#2E75B6', label='Complaint Rate', zorder=3)
    
    # Plot mean line
    ax.axhline(y=mean_rate, color='green', linestyle='--', linewidth=2, label=f'Mean ({mean_rate:.4f}%)', alpha=0.7)
    
    # Plot UCL
    ax.axhline(y=ucl, color='red', linestyle='--', linewidth=2.5, label=f'UCL ({ucl:.4f}%)', alpha=0.8)
    
    # Plot LCL if > 0
    if lcl > 0:
        ax.axhline(y=lcl, color='orange', linestyle='--', linewidth=2, label=f'LCL ({lcl:.4f}%)', alpha=0.7)
    
    # Highlight breaches if any
    if len(breaches) > 0:
        ax.scatter(breaches['Month'], breaches['Complaint_Rate'], 
                  color='red', s=200, zorder=5, marker='X', edgecolors='darkred', linewidths=2, label='UCL Breach')
    
    # Formatting
    ax.set_xlabel('Month', fontsize=12, fontweight='bold')
    ax.set_ylabel('Complaint Rate (%)', fontsize=12, fontweight='bold')
    
    start_str = pd.to_datetime(start_date).strftime('%B %Y')
    end_str = pd.to_datetime(end_date).strftime('%B %Y')
    ax.set_title(f'Monthly Complaint Rate Trend with Control Limits\n{product_name} ({start_str} - {end_str})', 
                 fontsize=14, fontweight='bold', pad=20)
    ax.legend(loc='upper left', fontsize=10, framealpha=0.9)
    ax.grid(True, alpha=0.3, linestyle=':')
    
    # Format x-axis
    ax.xaxis.set_major_locator(MonthLocator(interval=3))
    ax.xaxis.set_major_formatter(DateFormatter('%b %Y'))
    plt.xticks(rotation=45, ha='right')
    
    # Set y-axis to start at 0
    ax.set_ylim(bottom=0)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()

def generate_llm_analysis(merged_data, mean_rate, ucl, breaches, start_date, end_date, product_name):
    """Use Claude API to generate narrative analysis."""
    
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    
    # Prepare breach details
    breach_details = ""
    if len(breaches) > 0:
        breach_details = "UCL BREACHES IDENTIFIED:\n"
        for idx, row in breaches.iterrows():
            breach_details += f"- {row['Month'].strftime('%B %Y')}: Rate = {row['Complaint_Rate']:.4f}% (UCL = {ucl:.4f}%)\n"
            breach_details += f"  Units sold: {row['Quantity']:,}, Complaints: {int(row['Complaint_Count'])}\n"
            pct_over = ((row['Complaint_Rate'] - ucl) / ucl) * 100
            breach_details += f"  Exceeded UCL by {pct_over:.1f}%\n"
    else:
        breach_details = "NO UCL BREACHES: No months exceeded the Upper Control Limit during the reporting period."
    
    # Create system instruction
    system_instruction = """You are a Post-Market Surveillance (PMS) regulatory expert specializing in medical device PSUR (Periodic Safety Update Report) preparation. Your task is to generate professional, factual narratives for Section G (Trend Reporting) based on statistical complaint data.

Your analysis must:
1. Be strictly factual and data-driven - NO assumptions about investigations, root causes, or actions taken
2. Report only observable statistical patterns from the data
3. For UCL breaches: state the facts (date, rate, units, complaints) without inventing explanations
4. Be concise and suitable for regulatory submission
5. Use appropriate technical terminology
6. If actual investigation results or actions are unknown, state "Further investigation documented separately" or similar

Generate a SINGLE narrative paragraph that:
- Describes the overall trend pattern observed
- Reports any UCL breaches factually (month, rate, circumstances)
- Notes the overall complaint rate and control limits
- Does NOT invent corrective actions, root causes, or investigation findings
- Maintains professional regulatory tone"""

    # Create user prompt
    user_prompt = f"""Generate Section G narrative analysis for the following complaint trending data:

PRODUCT: {product_name}
REPORTING PERIOD: {pd.to_datetime(start_date).strftime('%B %Y')} to {pd.to_datetime(end_date).strftime('%B %Y')}
TOTAL MONTHS: {len(merged_data)}
TOTAL UNITS SOLD: {merged_data['Quantity'].sum():,}
TOTAL COMPLAINTS: {merged_data['Complaint_Count'].sum():,}
OVERALL COMPLAINT RATE: {(merged_data['Complaint_Count'].sum() / merged_data['Quantity'].sum() * 100):.4f}%
MEAN MONTHLY COMPLAINT RATE: {mean_rate:.4f}%
UPPER CONTROL LIMIT (UCL): {ucl:.4f}%

{breach_details}

Generate a single factual paragraph describing the trend data. Include UCL breach information if applicable, but do NOT invent explanations, investigations, or corrective actions that are not explicitly provided in the data above."""

    # Call Claude API
    message = client.messages.create(
        model=MODEL,
        max_tokens=2000,
        system=system_instruction,
        messages=[
            {"role": "user", "content": user_prompt}
        ]
    )
    
    return message.content[0].text

def create_section_g_document(merged_data, mean_rate, ucl, breaches, chart_path, 
                               llm_analysis, product_name, output_path):
    """Create the Section G Word document."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Set default font to Arial 10
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # Set line spacing to 1.5
    style.paragraph_format.line_spacing = 1.5
    
    # Add Section Heading - Arial 10 Bold Underlined
    heading = doc.add_paragraph()
    run = heading.add_run('Section G: Information From Trend Reporting')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.underline = True
    heading.paragraph_format.space_after = Pt(12)
    
    # Overall Monthly Complaint Rate Trending - Arial 10 Bold Underlined
    subheading1 = doc.add_paragraph()
    run = subheading1.add_run('Overall Monthly Complaint Rate Trending')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.underline = True
    subheading1.paragraph_format.space_after = Pt(6)
    
    # Add graph
    para_graph = doc.add_paragraph()
    run_caption = para_graph.add_run('Figure 1: Monthly Complaint Rate Trend Chart')
    run_caption.font.name = 'Arial'
    run_caption.font.size = Pt(10)
    run_caption.italic = True
    run_img = para_graph.add_run()
    run_img.add_picture(chart_path, width=Inches(6.5))
    para_graph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_graph.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    
    # Add context/commentary section for UCL breaches
    if len(breaches) > 0:
        # Add breach details factually
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = 1.5
        run_label = para.add_run('Context/Commentary: ')
        run_label.font.name = 'Arial'
        run_label.font.size = Pt(10)
        run_label.bold = True
        
        breach_text = ""
        for idx, row in breaches.iterrows():
            breach_date = row['Month'].strftime('%B %Y')
            breach_rate = row['Complaint_Rate']
            breach_count = int(row['Complaint_Count'])
            breach_units = int(row['Quantity'])
            
            breach_text += f"In {breach_date}, the complaint rate of {breach_rate:.4f}% ({breach_count} complaints from {breach_units:,} units sold) exceeded the Upper Control Limit of {ucl:.4f}%. "
        
        breach_text += "Further investigation and any actions taken are documented in the applicable CAPA records. "
        run_text = para.add_run(breach_text)
        run_text.font.name = 'Arial'
        run_text.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(12)
    
    # Add LLM-generated analysis
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(llm_analysis.strip())
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(12)
    
    # Trend Reporting Section - Arial 10 Bold Underlined
    subheading2 = doc.add_paragraph()
    run = subheading2.add_run('Trend Reporting')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.underline = True
    subheading2.paragraph_format.space_after = Pt(6)
    
    # Add default text if no trend reports (no guidance text in output)
    trend_default = doc.add_paragraph()
    trend_default.paragraph_format.line_spacing = 1.5
    run = trend_default.add_run('N/A – no trend reports were submitted regarding the data collection period.')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    trend_default.paragraph_format.space_after = Pt(12)
    
    # Save
    doc.save(output_path)

def export_data_to_excel(merged_data, mean_rate, ucl, lcl, breaches, chart_path, product_name, start_date, end_date, output_path):
    """Export trend data and chart to Excel file."""
    
    # Create Excel writer
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        # Prepare main data with formatted columns
        export_df = merged_data.copy()
        export_df['Month'] = export_df['Month'].dt.strftime('%Y-%m-%d')
        export_df['Complaint_Rate'] = export_df['Complaint_Rate'].round(4)
        export_df['Mean_Rate'] = mean_rate
        export_df['UCL'] = ucl
        export_df['LCL'] = lcl
        export_df['Is_Breach'] = export_df['Complaint_Rate'] > ucl
        
        # Rename columns for clarity
        export_df = export_df.rename(columns={
            'Month': 'Month',
            'Quantity': 'Units_Sold',
            'Complaint_Count': 'Complaints',
            'Complaint_Rate': 'Complaint_Rate_%',
            'Mean_Rate': 'Mean_Rate_%',
            'UCL': 'Upper_Control_Limit_%',
            'LCL': 'Lower_Control_Limit_%',
            'Is_Breach': 'UCL_Breach'
        })
        
        # Write data to 'Trend Data' sheet
        export_df.to_excel(writer, sheet_name='Trend Data', index=False)
        
        # Create summary sheet
        summary_data = {
            'Parameter': [
                'Product Name',
                'Start Date',
                'End Date',
                'Total Months',
                'Total Units Sold',
                'Total Complaints',
                'Overall Complaint Rate (%)',
                'Mean Monthly Complaint Rate (%)',
                'Standard Deviation (%)',
                'Upper Control Limit (UCL) (%)',
                'Lower Control Limit (LCL) (%)',
                'Number of UCL Breaches'
            ],
            'Value': [
                product_name,
                pd.to_datetime(start_date).strftime('%B %Y'),
                pd.to_datetime(end_date).strftime('%B %Y'),
                len(merged_data),
                merged_data['Quantity'].sum(),
                merged_data['Complaint_Count'].sum(),
                round((merged_data['Complaint_Count'].sum() / merged_data['Quantity'].sum() * 100), 4),
                round(mean_rate, 4),
                round(merged_data['Complaint_Rate'].std(), 4),
                round(ucl, 4),
                round(lcl, 4),
                len(breaches)
            ]
        }
        summary_df = pd.DataFrame(summary_data)
        summary_df.to_excel(writer, sheet_name='Summary', index=False)
        
        # Create breaches sheet if any exist
        if len(breaches) > 0:
            breach_export = breaches.copy()
            breach_export['Month'] = breach_export['Month'].dt.strftime('%B %Y')
            breach_export['Complaint_Rate'] = breach_export['Complaint_Rate'].round(4)
            breach_export['UCL'] = ucl
            breach_export['Exceeded_By_%'] = ((breach_export['Complaint_Rate'] - ucl) / ucl * 100).round(2)
            
            breach_export = breach_export.rename(columns={
                'Month': 'Month',
                'Quantity': 'Units_Sold',
                'Complaint_Count': 'Complaints',
                'Complaint_Rate': 'Complaint_Rate_%',
                'UCL': 'Upper_Control_Limit_%',
                'Exceeded_By_%': 'Exceeded_UCL_By_%'
            })
            
            breach_export[['Month', 'Units_Sold', 'Complaints', 'Complaint_Rate_%', 'Upper_Control_Limit_%', 'Exceeded_UCL_By_%']].to_excel(
                writer, sheet_name='UCL Breaches', index=False
            )
        
        # Get workbook and chart sheet
        workbook = writer.book
        chart_sheet = workbook.create_sheet('Chart')
        
        # Add chart image to Excel
        from openpyxl.drawing.image import Image
        img = Image(chart_path)
        img.width = 800
        img.height = 400
        chart_sheet.add_image(img, 'A1')
        
        # Format the data sheet
        worksheet = writer.sheets['Trend Data']
        
        # Auto-adjust column widths
        for column in worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            worksheet.column_dimensions[column_letter].width = adjusted_width
        
        # Format summary sheet
        summary_worksheet = writer.sheets['Summary']
        for column in summary_worksheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            summary_worksheet.column_dimensions[column_letter].width = adjusted_width

def main():
    """Main execution function."""
    
    # Configuration - MODIFY THESE
    # Point to inputs folder in root directory
    inputs_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'inputs')
    SALES_FILE = os.path.join(inputs_dir, '33_sales.xlsx')
    COMPLAINTS_FILE = os.path.join(inputs_dir, '33_complaints.xlsx')
    START_DATE = '2020-04-01'
    END_DATE = '2024-12-31'
    PRODUCT_NAME = 'INCA Complete Set'
    OUTPUT_FILE = 'output/Section_G_Trend_Reporting.docx'
    CHART_FILE = 'output/complaint_rate_trend.png'
    EXCEL_FILE = 'output/Section_G_Trend_Data.xlsx'
    
    # Check API key
    if not ANTHROPIC_API_KEY:
        print("ERROR: ANTHROPIC_API_KEY not set. Please set environment variable or update script.")
        sys.exit(1)
    
    print(f"Processing Section G for {PRODUCT_NAME}...")
    print(f"Period: {START_DATE} to {END_DATE}")
    
    # Load and process data
    print("\n1. Loading and processing data...")
    merged_data = load_and_process_data(SALES_FILE, COMPLAINTS_FILE, START_DATE, END_DATE)
    print(f"   - {len(merged_data)} months analyzed")
    print(f"   - {merged_data['Quantity'].sum():,} units sold")
    print(f"   - {merged_data['Complaint_Count'].sum():,} complaints")
    
    # Calculate control limits
    print("\n2. Calculating control limits...")
    mean_rate, std_rate, ucl, lcl = calculate_control_limits(merged_data)
    print(f"   - Mean: {mean_rate:.4f}%")
    print(f"   - UCL: {ucl:.4f}%")
    
    # Identify breaches
    print("\n3. Identifying UCL breaches...")
    breaches = identify_breaches(merged_data, ucl)
    print(f"   - {len(breaches)} breach(es) found")
    
    # Create chart
    print("\n4. Generating trend chart...")
    create_trend_chart(merged_data, mean_rate, ucl, lcl, breaches, PRODUCT_NAME, 
                      START_DATE, END_DATE, CHART_FILE)
    print(f"   - Chart saved: {CHART_FILE}")
    
    # Generate LLM analysis
    print("\n5. Generating LLM-powered narrative analysis...")
    llm_analysis = generate_llm_analysis(merged_data, mean_rate, ucl, breaches, 
                                         START_DATE, END_DATE, PRODUCT_NAME)
    print("   - Analysis generated")
    
    # Create document
    print("\n6. Creating Section G document...")
    create_section_g_document(merged_data, mean_rate, ucl, breaches, CHART_FILE,
                              llm_analysis, PRODUCT_NAME, OUTPUT_FILE)
    print(f"   - Document saved: {OUTPUT_FILE}")
    
    # Export data to Excel
    print("\n7. Exporting data to Excel...")
    export_data_to_excel(merged_data, mean_rate, ucl, lcl, breaches, CHART_FILE,
                        PRODUCT_NAME, START_DATE, END_DATE, EXCEL_FILE)
    print(f"   - Excel file saved: {EXCEL_FILE}")
    
    print("\n[SUCCESS] Section G generation complete!")
    print(f"\nOutput file: {OUTPUT_FILE}")
    print(f"Chart file: {CHART_FILE}")
    print(f"Excel data file: {EXCEL_FILE}")

if __name__ == "__main__":
    main()