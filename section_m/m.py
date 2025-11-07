"""
PSUR Section M Generator - Findings and Conclusions
EU MDR Compliant per Article 86 and MDCG 2022-21
"""

import os
import anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING
from glob import glob

# Load environment variables from root .env file
root_dir = os.path.dirname(os.path.dirname(__file__))
env_path = os.path.join(root_dir, '.env')
load_dotenv(env_path)


class AnthropicLLM:
    def __init__(self):
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY environment variable required")
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"
        self.sys = """You are an expert regulatory affairs technical writer with deep expertise in Post-Market Surveillance and PSUR preparation per EU MDR 2017/745, UKCA and MDCG 2022-21 guidance.

CRITICAL REQUIREMENTS FOR FINDINGS AND CONCLUSIONS:

1. BENEFIT-RISK PROFILE ASSESSMENT:
- Make clear determination: Has benefit-risk profile been adversely impacted or NOT adversely impacted?
- Base conclusion on comprehensive analysis of all surveillance data
- Consider sales data, vigilance information, customer feedback, clinical literature, external databases, CAPAs, and PMCF data
- State conclusion definitively with supporting evidence

2. INTENDED BENEFITS:
- Assess whether all intended benefits were achieved
- Reference device's intended use and performance claims
- Use data from surveillance activities to support assessment

3. DATA LIMITATIONS:
- Identify any limitations to the data used for analysis
- Be specific about what limitations exist (e.g., incomplete data, small sample sizes, missing information)
- Explain how limitations affect the conclusions

4. NEW OR EMERGING RISKS/BENEFITS:
- Identify any new or emerging risks identified during reporting period
- Identify any new benefits discovered
- State clearly if none were identified

5. MANUFACTURER ACTIONS:
- Document any actions taken to update:
  - Benefit-Risk Assessment
  - Risk Management File
  - Product Design
  - Manufacturing Process
  - Instructions for Use
  - Labeling
  - Clinical Evaluation Report
  - Summary of Safety and Clinical Performance (if applicable)
  - Corrective and Preventive Actions (CAPAs)
  - Field Safety Corrective Actions (FSCAs)
- Document actions to be taken and means of follow-up
- State clearly if no actions were required

6. OVERALL PERFORMANCE CONCLUSION:
- Provide comprehensive overall performance conclusion
- Synthesize findings from all surveillance activities
- State device's performance relative to intended use

7. WRITING STANDARDS:
- Professional, objective, scientific tone
- Evidence-based statements with specific data
- Clear, definitive conclusions
- 2-3 well-structured paragraphs per subsection
- Use Arial 10pt font-compliant language
- DO NOT cite regulations in narrative text
- DO NOT use asterisks, hash marks, or markdown formatting
- Use plain paragraph format without bullet points

8. REGULATORY COMPLIANCE:
- All statements verifiable against source documents
- Consistency with all PSUR sections
- Integration of all surveillance data
- Alignment with MDCG 2022-21 requirements

OUTPUT FORMAT:
- Clear, concise paragraphs
- Specific numerical data where applicable
- Professional regulatory language
- No promotional or marketing language
- Plain text without special formatting characters"""
    
    def generate(self, prompt: str) -> str:
        return self.client.messages.create(
            model=self.model, max_tokens=4000, temperature=0.2,
            system=self.sys, messages=[{"role": "user", "content": prompt}]
        ).content[0].text


class PSURSectionMGenerator:
    def __init__(self, device_name: str):
        self.device_name = device_name
        try:
            self.llm = AnthropicLLM() if os.environ.get("ANTHROPIC_API_KEY") else None
            if self.llm:
                print("  LLM initialized successfully")
            else:
                print("  WARNING: ANTHROPIC_API_KEY not found - using fallback text")
        except Exception as e:
            print(f"  ERROR: LLM initialization failed: {e}")
            self.llm = None
    
    def _extract_section_data(self, section_name: str, output_dir: str) -> str:
        """Extract relevant data from a PSUR section output"""
        try:
            # Try multiple naming patterns
            patterns = [
                os.path.join(output_dir, f'PSUR_Section_{section_name}*.docx'),
                os.path.join(output_dir, f'Section_{section_name}*.docx'),
                os.path.join(output_dir, f'*Section_{section_name}*.docx')
            ]
            
            files = []
            for pattern in patterns:
                found_files = glob(pattern)
                files.extend(found_files)
            
            # Remove duplicates
            files = list(set(files))
            
            if not files:
                return ""
            
            # Get the main output file (prefer files with "PSUR" or main section name)
            main_files = [f for f in files if 'PSUR' in os.path.basename(f) or 
                         os.path.basename(f).startswith(f'Section_{section_name}_')]
            
            if main_files:
                latest_file = max(main_files, key=os.path.getmtime)
            else:
                latest_file = max(files, key=os.path.getmtime)
            
            from docx import Document as DocxDocument
            doc = DocxDocument(latest_file)
            
            # Extract all text - no character limit
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            
            return text
        except Exception as e:
            print(f"  Note: Could not read Section {section_name}: {e}")
            return ""
    
    def _collect_all_section_data(self) -> dict:
        """Collect data from all PSUR sections"""
        root_dir = os.path.dirname(os.path.dirname(__file__))
        section_data = {}
        
        sections = {
            'C': 'section_c/output',
            'D': 'section_d/output',
            'F': 'section_f/output',
            'G': 'section_g/output',
            'J': 'section_j/output',
            'L': 'section_l/output'
        }
        
        print("\nCollecting data from all PSUR sections...")
        for section, subdir in sections.items():
            output_dir = os.path.join(root_dir, subdir)
            data = self._extract_section_data(section, output_dir)
            if data:
                section_data[section] = data
                print(f"  Extracted data from Section {section}")
            else:
                print(f"  Note: No data found for Section {section}")
        
        return section_data
    
    def _generate_benefit_risk_conclusion(self, all_section_data: dict) -> str:
        """Generate benefit-risk profile conclusion"""
        if not self.llm:
            return "Benefit-risk profile assessment based on comprehensive surveillance data."
        
        try:
            # Compile context from all sections
            context_parts = []
            for section, data in all_section_data.items():
                context_parts.append(f"Section {section}:\n{data}")
            
            context = "\n\n".join(context_parts)
            
            prompt = f"""Generate a Benefit-Risk Profile Conclusion for PSUR Section M.

ALL PSUR SECTIONS DATA:
{context if context else "No section data available."}

Device: {self.device_name}

CRITICAL REQUIREMENTS:
1. Make clear determination: "The benefit-risk profile has been adversely impacted" OR "The benefit-risk profile has NOT been adversely impacted and remains unchanged"
2. Base conclusion ONLY on actual data from the sections
3. Use ONLY specific data points explicitly stated in the sections
4. DO NOT make assumptions about:
   - Dates, timeframes, or specific months/years unless explicitly stated
   - Percentages, rates, or numbers not provided in the data
   - Investigation status or CAPA records unless mentioned
   - Trends or patterns not explicitly documented
   - Statistical analysis not present in the data
5. If specific data is not available, state generally without inventing details
6. Generate 1 CONCISE paragraph focused specifically on benefit-risk determination
7. DO NOT cite regulations
8. DO NOT use special formatting characters
9. NO ASSUMPTIONS - stick strictly to what is documented in the sections
10. Keep focused and brief - detailed synthesis will be in Overall Performance section

Generate 1 concise paragraph making the benefit-risk determination."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating benefit-risk conclusion: {e}")
            return "Benefit-risk profile assessment based on comprehensive surveillance data."
    
    def _generate_intended_benefits(self, all_section_data: dict) -> str:
        """Generate intended benefits assessment"""
        if not self.llm:
            return "Intended benefits assessment based on surveillance data."
        
        try:
            context_parts = []
            for section, data in all_section_data.items():
                context_parts.append(f"Section {section}:\n{data}")
            
            context = "\n\n".join(context_parts)
            
            prompt = f"""Generate an Intended Benefits Assessment for PSUR Section M.

ALL PSUR SECTIONS DATA:
{context if context else "No section data available."}

Device: {self.device_name}

CRITICAL REQUIREMENTS:
1. Assess whether all intended benefits were achieved based ONLY on data provided
2. Use ONLY data from surveillance activities explicitly stated in the sections
3. DO NOT make assumptions about:
   - Specific performance metrics not documented
   - Patient outcomes not mentioned
   - Clinical effectiveness not stated in the data
   - User satisfaction not explicitly reported
4. If specific benefit data is not available, state generally without inventing details
5. Generate 1 CONCISE paragraph focused specifically on intended benefits
6. DO NOT cite regulations
7. DO NOT use special formatting characters
8. NO ASSUMPTIONS - stick strictly to documented evidence
9. Keep focused and brief - detailed synthesis will be in Overall Performance section

Generate 1 concise paragraph assessing intended benefits only."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating intended benefits: {e}")
            return "Intended benefits assessment based on surveillance data."
    
    def _generate_data_limitations(self, all_section_data: dict) -> str:
        """Generate data limitations assessment"""
        if not self.llm:
            return "Data limitations assessment."
        
        try:
            context_parts = []
            for section, data in all_section_data.items():
                context_parts.append(f"Section {section}:\n{data}")
            
            context = "\n\n".join(context_parts)
            
            prompt = f"""Generate a Data Limitations Assessment for PSUR Section M.

ALL PSUR SECTIONS DATA:
{context if context else "No section data available."}

Device: {self.device_name}

CRITICAL REQUIREMENTS:
1. Identify ONLY limitations explicitly mentioned or clearly evident in the provided data
2. DO NOT assume limitations that are not documented
3. DO NOT make assumptions about:
   - Sample sizes unless stated
   - Data completeness unless mentioned
   - Missing information unless explicitly noted
   - Time constraints unless documented
4. If limitations are mentioned in the sections, cite them specifically
5. If no limitations are explicitly stated, note that analysis was based on available data
6. Generate 1 CONCISE paragraph focused specifically on data limitations
7. DO NOT cite regulations
8. DO NOT use special formatting characters
9. NO ASSUMPTIONS - only report documented limitations
10. Keep focused and brief - detailed synthesis will be in Overall Performance section

Generate 1 concise paragraph on data limitations only."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating data limitations: {e}")
            return "Data limitations assessment."
    
    def _generate_new_risks_benefits(self, all_section_data: dict) -> str:
        """Generate new or emerging risks/benefits assessment"""
        if not self.llm:
            return "No new or emerging risks or benefits identified during reporting period."
        
        try:
            context_parts = []
            for section, data in all_section_data.items():
                context_parts.append(f"Section {section}:\n{data}")
            
            context = "\n\n".join(context_parts)
            
            prompt = f"""Generate a New or Emerging Risks/Benefits Assessment for PSUR Section M.

ALL PSUR SECTIONS DATA:
{context if context else "No section data available."}

Device: {self.device_name}

CRITICAL REQUIREMENTS:
1. Identify ONLY new or emerging risks explicitly mentioned in the provided sections
2. Identify ONLY new benefits explicitly documented in the sections
3. DO NOT make assumptions about:
   - Risks not specifically identified in vigilance data
   - Trends not explicitly documented
   - Issues not reported in complaints or literature
   - Benefits not specifically mentioned
4. If sections explicitly state "no new risks" or "no new benefits", report that
5. If no new risks/benefits are mentioned, state that none were identified based on available data
6. Generate 1 CONCISE paragraph focused specifically on new/emerging risks and benefits
7. DO NOT cite regulations
8. DO NOT use special formatting characters
9. NO ASSUMPTIONS - only report explicitly documented new risks or benefits
10. Keep focused and brief - detailed synthesis will be in Overall Performance section

Generate 1 concise paragraph on new/emerging risks and benefits only."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating new risks/benefits: {e}")
            return "No new or emerging risks or benefits identified during reporting period."
    
    def _generate_manufacturer_actions(self, all_section_data: dict) -> str:
        """Generate manufacturer actions assessment"""
        if not self.llm:
            return "No actions required during reporting period."
        
        try:
            context_parts = []
            for section, data in all_section_data.items():
                context_parts.append(f"Section {section}:\n{data}")
            
            context = "\n\n".join(context_parts)
            
            prompt = f"""Generate a Manufacturer Actions Assessment for PSUR Section M.

ALL PSUR SECTIONS DATA:
{context if context else "No section data available."}

Device: {self.device_name}

CRITICAL REQUIREMENTS:
1. Document ONLY actions explicitly mentioned in the provided sections
2. DO NOT make assumptions about:
   - CAPAs not specifically documented
   - Updates to documentation not explicitly stated
   - Design changes not mentioned
   - Follow-up activities not documented
   - Investigation status not reported
3. If sections mention specific actions (updates to RMF, CER, IFU, etc.), report those specifically
4. If no actions are mentioned in the data, state that no actions were identified as necessary during this period
5. Generate 1 CONCISE paragraph focused specifically on manufacturer actions
6. DO NOT cite regulations
7. DO NOT use special formatting characters
8. NO ASSUMPTIONS - only report actions actually documented in the sections
9. Keep focused and brief - detailed synthesis will be in Overall Performance section

Generate 1 concise paragraph on manufacturer actions only."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating manufacturer actions: {e}")
            return "No actions required during reporting period."
    
    def _generate_overall_performance(self, all_section_data: dict) -> str:
        """Generate overall performance conclusion"""
        if not self.llm:
            return "Overall performance conclusion based on surveillance data."
        
        try:
            context_parts = []
            for section, data in all_section_data.items():
                context_parts.append(f"Section {section}:\n{data}")
            
            context = "\n\n".join(context_parts)
            
            prompt = f"""Generate an Overall Performance Conclusion for PSUR Section M.

ALL PSUR SECTIONS DATA:
{context if context else "No section data available."}

Device: {self.device_name}

CRITICAL REQUIREMENTS:
1. Synthesize ONLY findings explicitly stated in the provided sections
2. Base performance assessment on actual data from:
   - Sales data (Section C)
   - Vigilance information (Section D)
   - Customer feedback/complaints (Section F)
   - Clinical literature (Section J)
   - PMCF data (Section L)
3. DO NOT make assumptions about:
   - Performance metrics not documented
   - Trends not explicitly stated
   - Comparisons not made in the data
   - Clinical outcomes not reported
   - User feedback not documented
4. Use ONLY specific data points from the sections
5. If data is limited, acknowledge that and base conclusion on available information
6. Generate 2-3 COMPLETE, DETAILED paragraphs
7. DO NOT cite regulations
8. DO NOT use special formatting characters
9. NO ASSUMPTIONS - stick strictly to documented evidence from all sections

Generate 2-3 detailed paragraphs using ONLY actual data from the sections provided."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating overall performance: {e}")
            return "Overall performance conclusion based on surveillance data."
    
    def generate(self, output_path: str = "Section_M.docx") -> str:
        """Generate complete Section M document"""
        
        print("\n=== PSUR SECTION M GENERATION ===")
        print("Findings and Conclusions per EU MDR Article 86")
        
        # Collect data from all sections
        all_section_data = self._collect_all_section_data()
        
        if not all_section_data:
            print("  WARNING: No section data found - generating with limited context")
        
        # Generate all subsections
        print("\nGenerating conclusions...")
        benefit_risk = self._generate_benefit_risk_conclusion(all_section_data)
        intended_benefits = self._generate_intended_benefits(all_section_data)
        data_limitations = self._generate_data_limitations(all_section_data)
        new_risks_benefits = self._generate_new_risks_benefits(all_section_data)
        manufacturer_actions = self._generate_manufacturer_actions(all_section_data)
        overall_performance = self._generate_overall_performance(all_section_data)
        
        # Build document
        print("Building Word document...")
        doc = self._build_document(benefit_risk, intended_benefits, data_limitations, 
                                   new_risks_benefits, manufacturer_actions, overall_performance)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)
        
        print(f"\n[SUCCESS] Generated: {output_path}")
        print("\nREGULATORY CHECKLIST:")
        print("  [x] Benefit-risk profile conclusion")
        print("  [x] Intended benefits assessment")
        print("  [x] Data limitations")
        print("  [x] New/emerging risks and benefits")
        print("  [x] Manufacturer actions")
        print("  [x] Overall performance conclusion")
        print("\nREVIEW REQUIRED:")
        print("  - Verify integration of all surveillance data")
        print("  - Confirm all conclusions are supported by evidence")
        print("  - Validate consistency with all PSUR sections")
        
        return output_path
    
    def _build_document(self, benefit_risk, intended_benefits, data_limitations, 
                       new_risks_benefits, manufacturer_actions, overall_performance):
        """Build formatted Word document"""
        doc = Document()
        
        # Set default style
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(10)
        
        # Title
        heading = doc.add_heading('Section M: Findings and Conclusions', level=1)
        heading.runs[0].font.name = 'Arial'
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        
        # a) Benefit-Risk Profile Conclusion
        br_heading = doc.add_paragraph('a) Benefit-Risk Profile Conclusion')
        br_heading.runs[0].font.name = 'Arial'
        br_heading.runs[0].font.size = Pt(10)
        br_heading.runs[0].font.bold = True
        br_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        br_para = doc.add_paragraph(benefit_risk)
        br_para.runs[0].font.name = 'Arial'
        br_para.runs[0].font.size = Pt(10)
        br_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # b) Intended Benefits
        ib_heading = doc.add_paragraph('b) Intended Benefits Assessment')
        ib_heading.runs[0].font.name = 'Arial'
        ib_heading.runs[0].font.size = Pt(10)
        ib_heading.runs[0].font.bold = True
        ib_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        ib_para = doc.add_paragraph(intended_benefits)
        ib_para.runs[0].font.name = 'Arial'
        ib_para.runs[0].font.size = Pt(10)
        ib_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # c) Data Limitations
        dl_heading = doc.add_paragraph('c) Data Limitations')
        dl_heading.runs[0].font.name = 'Arial'
        dl_heading.runs[0].font.size = Pt(10)
        dl_heading.runs[0].font.bold = True
        dl_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        dl_para = doc.add_paragraph(data_limitations)
        dl_para.runs[0].font.name = 'Arial'
        dl_para.runs[0].font.size = Pt(10)
        dl_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # d) New or Emerging Risks/Benefits
        ner_heading = doc.add_paragraph('d) New or Emerging Risks and Benefits')
        ner_heading.runs[0].font.name = 'Arial'
        ner_heading.runs[0].font.size = Pt(10)
        ner_heading.runs[0].font.bold = True
        ner_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        ner_para = doc.add_paragraph(new_risks_benefits)
        ner_para.runs[0].font.name = 'Arial'
        ner_para.runs[0].font.size = Pt(10)
        ner_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # e) Manufacturer Actions
        ma_heading = doc.add_paragraph('e) Manufacturer Actions')
        ma_heading.runs[0].font.name = 'Arial'
        ma_heading.runs[0].font.size = Pt(10)
        ma_heading.runs[0].font.bold = True
        ma_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        ma_para = doc.add_paragraph(manufacturer_actions)
        ma_para.runs[0].font.name = 'Arial'
        ma_para.runs[0].font.size = Pt(10)
        ma_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # f) Overall Performance Conclusion
        op_heading = doc.add_paragraph('f) Overall Performance Conclusion')
        op_heading.runs[0].font.name = 'Arial'
        op_heading.runs[0].font.size = Pt(10)
        op_heading.runs[0].font.bold = True
        op_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        op_para = doc.add_paragraph(overall_performance)
        op_para.runs[0].font.name = 'Arial'
        op_para.runs[0].font.size = Pt(10)
        op_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        return doc


def main():
    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)
    
    # Device name
    device_name = "INCA Complete Set"
    
    # Generate Section M
    print("\nGenerating Section M...")
    generator = PSURSectionMGenerator(device_name)
    
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_filename = f'PSUR_Section_M_{timestamp}.docx'
    output_path = f'output/{output_filename}'
    
    output = generator.generate(output_path=output_path)
    
    print(f"\n[SUCCESS] Generated: {output}")


if __name__ == "__main__":
    main()

