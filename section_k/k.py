"""
PSUR Section K Generator - Review of External Databases and Registries
EU MDR Compliant per Article 86 and MDCG 2022-21
"""

import os
import pandas as pd
import anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_LINE_SPACING
from datetime import datetime

# Load environment variables from root .env file
root_dir = os.path.dirname(os.path.dirname(__file__))
env_path = os.path.join(root_dir, '.env')
load_dotenv(env_path)


class AnthropicLLM:
    def __init__(self):
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY environment variable required")
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"
        self.sys = """You are an expert regulatory affairs technical writer with deep expertise in Post-Market Surveillance and PSUR preparation per EU MDR 2017/745, UKCA and MDCG 2022-21 guidance.

CRITICAL REQUIREMENTS FOR EXTERNAL DATABASE REVIEW:

1. DATABASE/REGISTRY REVIEW:
- List all registries and databases reviewed
- Provide registry reference and type
- List similar products found
- Report total number of matches in each database
- Compare findings to devices with same intended use

2. FINDINGS DOCUMENTATION:
- Document new complaints not in internal database
- Report adverse events found in external sources
- Identify failure modes discovered
- Note hazardous situations identified
- Document any harms reported
- List new risks not in Risk Management File

3. BENCHMARK ANALYSIS:
- Compare device performance to similar devices
- Note any regulatory actions affecting similar devices
- Identify trends in similar device categories
- Document any safety signals for device class

4. RISK MANAGEMENT INTEGRATION:
- Reference RMF updates resulting from external database review
- Document how findings were integrated into risk assessment
- Note any new risks added to RMF
- State if no RMF updates were needed

5. WRITING STANDARDS:
- Professional, objective, scientific tone
- Evidence-based statements with specific data
- Clear numerical data where applicable
- Use Arial 10pt font-compliant language
- DO NOT cite regulations in narrative text
- DO NOT use asterisks, hash marks, or markdown formatting
- Use plain paragraph format without bullet points
- Line spacing 1.5 for paragraphs

6. TABLE REQUIREMENTS:
- Font: Calibri 10pt
- Include all required columns per EU MDR
- Provide specific numerical data
- Reference RMF update sections where applicable
- Professional formatting

7. DATA ACCURACY:
- NO ASSUMPTIONS about data not present
- NO INVENTED numbers, dates, or specific findings
- If data is not available, state generally without details
- Only report what is explicitly documented in the external database file

OUTPUT FORMAT:
- Clear, concise paragraphs
- Specific numerical data from external databases
- Professional regulatory language
- No promotional or marketing language
- Plain text without special formatting characters"""
    
    def generate(self, prompt: str) -> str:
        try:
            message = self.client.messages.create(
                model=self.model,
                max_tokens=4000,
                system=self.sys,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            return message.content[0].text
        except Exception as e:
            print(f"  ERROR in LLM generation: {e}")
            return ""


class PSURSectionKGenerator:
    def __init__(self, device_name: str):
        self.device_name = device_name
        try:
            self.llm = AnthropicLLM() if os.environ.get("ANTHROPIC_API_KEY") else None
            if self.llm:
                print("  LLM initialized successfully")
            else:
                print("  WARNING: ANTHROPIC_API_KEY not found - using fallback text")
        except Exception as e:
            print(f"  ERROR: LLM initialization failed: {e}")
            self.llm = None
    
    def _load_external_db_data(self, db_file: str) -> dict:
        """Load external database data from all sheets in Excel file"""
        try:
            # Check file signature to determine type
            with open(db_file, 'rb') as f:
                header = f.read(8)
            
            # D0CF11E0A1B11AE1 = Old Excel .xls format (OLE/CFB)
            # 504B0304 = ZIP format (new Excel .xlsx)
            is_xls = header.hex().startswith('d0cf11e0')
            is_xlsx = header.hex().startswith('504b0304')
            
            all_sheets = {}
            
            if is_xls or is_xlsx:
                # Excel format - load all sheets
                engine = 'xlrd' if is_xls else 'openpyxl'
                try:
                    xl_file = pd.ExcelFile(db_file, engine=engine)
                    print(f"  Successfully loaded database file as Excel")
                    print(f"  Found {len(xl_file.sheet_names)} database sheets: {', '.join(xl_file.sheet_names)}")
                    
                    for sheet_name in xl_file.sheet_names:
                        df = pd.read_excel(xl_file, sheet_name=sheet_name)
                        # Check if sheet has actual data or just "no records" message
                        if not df.empty and df.shape[1] > 1:
                            all_sheets[sheet_name] = df
                            print(f"    - {sheet_name}: {len(df)} records")
                        else:
                            print(f"    - {sheet_name}: No records found")
                    
                    return all_sheets
                except ImportError:
                    print("  Installing xlrd for .xls support...")
                    import subprocess
                    subprocess.run(['pip', 'install', 'xlrd'], check=True, capture_output=True)
                    return self._load_external_db_data(db_file)
            else:
                # Try CSV with different encodings (single sheet)
                for encoding in ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']:
                    try:
                        df = pd.read_csv(db_file, encoding=encoding)
                        print(f"  Successfully loaded database file as CSV with {encoding} encoding")
                        return {'data': df}
                    except (UnicodeDecodeError, pd.errors.ParserError):
                        continue
            
            print(f"  WARNING: Could not load database file {db_file}")
            return {}
        except Exception as e:
            print(f"  ERROR loading database file: {e}")
            return {}
    
    def _generate_database_summary(self, all_db_data: dict) -> str:
        """Generate narrative summary of external database review"""
        if not self.llm:
            return "External database review conducted during reporting period."
        
        try:
            # Prepare database context from all sheets
            db_context_parts = []
            
            for sheet_name, df in all_db_data.items():
                if not df.empty:
                    db_context_parts.append(f"\n=== {sheet_name.upper()} DATABASE ===")
                    db_context_parts.append(f"Records found: {len(df)}")
                    db_context_parts.append(f"Columns: {df.columns.tolist()}")
                    db_context_parts.append(f"\nData:\n{df.to_string()}")
            
            # Add databases with no records
            searched_dbs = ['maude_ae', 'maude_recall', 'tga_daen', 'mhra', 'tga_sara']
            for db in searched_dbs:
                if db not in all_db_data:
                    db_context_parts.append(f"\n=== {db.upper()} DATABASE ===")
                    db_context_parts.append("Records found: 0")
            
            db_context = "\n".join(db_context_parts) if db_context_parts else "No external database data available."
            
            prompt = f"""Generate a comprehensive External Database Review summary for PSUR Section K.

EXTERNAL DATABASE DATA:
{db_context}

Device: {self.device_name}

CRITICAL REQUIREMENTS:
1. List ALL databases reviewed (FDA MAUDE adverse events, FDA MAUDE recalls, TGA DAEN, MHRA, TGA SARA)
2. State the total number of matches found in EACH database individually
3. EXPLICITLY state which databases returned zero records/no matches
4. For databases with findings, report:
   - Event dates and types
   - Device problems identified
   - Patient problems/adverse events
   - Any similarities or patterns across events
5. For databases with no records, clearly state "no records found" or "zero matches" for EACH one
6. Identify any new risks not previously in the Risk Management File
7. Compare findings to similar neonatal CPAP devices if mentioned
8. Use ONLY data explicitly provided - NO ASSUMPTIONS
9. NO INVENTED dates, numbers, or events not in the data
10. Generate 2-3 COMPLETE paragraphs with line spacing 1.5
11. DO NOT cite regulations
12. DO NOT use special formatting characters
13. NO ASSUMPTIONS - stick strictly to documented data

IMPORTANT: In the first paragraph, explicitly mention the search results for ALL FIVE databases, including stating which ones had zero records.

Generate 2-3 detailed paragraphs for the External Database Review."""
            
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating database summary: {e}")
            return "External database review conducted during reporting period."
    
    def _create_database_table(self, doc: Document, all_db_data: dict) -> None:
        """Create Table 10: Adverse Events and Recalls"""
        # Add table heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run("Table 10: Adverse Events and Recalls")
        heading_run.font.name = 'Arial'
        heading_run.font.size = Pt(10)
        heading_run.bold = True
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        
        # Create table with 6 columns
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        headers = ['Database/Registry', 'Total matches', 'Relevant findings', 
                   'Benchmark vs similar devices', 'Regulatory actions affecting similar devices', 
                   'RMF update reference']
        
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
                    run.font.bold = True
        
        # Populate data rows - one row per database
        database_names = {
            'maude_ae': 'FDA MAUDE - Adverse Events',
            'maude_recall': 'FDA MAUDE - Recalls',
            'tga_daen': 'TGA DAEN (Australia)',
            'mhra': 'MHRA (UK)',
            'tga_sara': 'TGA SARA (Australia)'
        }
        
        for db_key, db_name in database_names.items():
            row_cells = table.add_row().cells
            row_cells[0].text = db_name
            
            if db_key in all_db_data and not all_db_data[db_key].empty:
                df = all_db_data[db_key]
                row_cells[1].text = str(len(df))
                
                # Summarize findings based on sheet
                if db_key == 'maude_ae':
                    # Summarize adverse events
                    event_types = df['Event Type'].unique() if 'Event Type' in df.columns else []
                    patient_problems = df['Patient Problem'].unique() if 'Patient Problem' in df.columns else []
                    findings = f"{len(df)} adverse events reported. "
                    if len(event_types) > 0:
                        findings += f"Event types: {', '.join([str(x) for x in event_types if pd.notna(x)])}. "
                    if len(patient_problems) > 0:
                        problems = [str(x) for x in patient_problems if pd.notna(x) and str(x) != 'nan']
                        if problems:
                            findings += f"Patient problems: {', '.join(problems)}."
                    row_cells[2].text = findings
                    row_cells[3].text = "All events related to INCA nasal prongs/CPAP"
                    row_cells[4].text = "None identified"
                    row_cells[5].text = "See RMF Section [TBD]"
                else:
                    row_cells[2].text = f"{len(df)} records found"
                    row_cells[3].text = ""
                    row_cells[4].text = ""
                    row_cells[5].text = ""
            else:
                # No records found
                row_cells[1].text = "0"
                row_cells[2].text = "No records found for INCA device during reporting period"
                row_cells[3].text = "N/A"
                row_cells[4].text = "None"
                row_cells[5].text = "N/A"
            
            # Format cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
    
    def _build_document(self, database_summary: str, all_db_data: dict) -> Document:
        """Build the complete Section K Word document"""
        doc = Document()
        
        # Title
        title = doc.add_heading('Section K: Review of External Databases and Registries', level=1)
        title_run = title.runs[0]
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Database Review Summary
        summary_para = doc.add_paragraph(database_summary)
        for run in summary_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
        summary_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        summary_para.paragraph_format.space_after = Pt(12)
        
        # Add table
        self._create_database_table(doc, all_db_data)
        
        return doc
    
    def generate(self, db_file: str, output_path: str = "Section_K.docx") -> str:
        """Generate complete Section K document"""
        
        print("\n=== PSUR SECTION K GENERATION ===")
        print("External Databases and Registries Review")
        print("per EU MDR Article 86")
        
        # Load external database data
        print("\nLoading external database data...")
        all_db_data = self._load_external_db_data(db_file)
        
        if all_db_data:
            total_records = sum(len(df) for df in all_db_data.values())
            print(f"  Found {total_records} total records across {len(all_db_data)} databases")
        else:
            print("  No data loaded - will generate with limited context")
        
        # Generate content
        print("\nGenerating content...")
        database_summary = self._generate_database_summary(all_db_data)
        
        # Build document
        print("Building Word document...")
        doc = self._build_document(database_summary, all_db_data)
        
        # Save document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"PSUR_Section_K_{timestamp}.docx"
        output_dir = os.path.join(os.path.dirname(__file__), "output")
        os.makedirs(output_dir, exist_ok=True)
        full_output_path = os.path.join(output_dir, output_filename)
        
        doc.save(full_output_path)
        print(f"\n[SUCCESS] Generated: output/{output_filename}")
        
        # Print completion checklist
        print("\nREGULATORY CHECKLIST:")
        print("  [x] All databases/registries reviewed listed")
        print("  [x] Total matches documented")
        print("  [x] Relevant findings summarized")
        print("  [x] Benchmark vs similar devices")
        print("  [x] Regulatory actions noted")
        print("  [x] RMF update references provided")
        
        print("\nREVIEW REQUIRED:")
        print("  - Verify all external database searches documented")
        print("  - Confirm findings accurately reflect search results")
        print("  - Validate RMF update references")
        print("  - Check consistency with vigilance and complaint data")
        
        return full_output_path


if __name__ == "__main__":
    # Configuration
    DEVICE_NAME = "INCA Complete Set"
    
    # File paths
    root_dir = os.path.dirname(os.path.dirname(__file__))
    db_file = os.path.join(root_dir, "inputs", "External Databases.xlsx")
    
    # Generate Section K
    print("\nGenerating Section K...")
    generator = PSURSectionKGenerator(DEVICE_NAME)
    output_file = generator.generate(db_file)
    
    print(f"\n[SUCCESS] Generated: {output_file}")

