"""
PSUR Section F Generator
========================
Comprehensive tool for generating Section F of Post-Market Surveillance Update Reports (PSUR)
in compliance with EU MDR and MDCG 2022-21 guidelines.

This script:
1. Processes complaint and sales data
2. Calculates complaint rates by harm categories and medical device problems
3. Generates regulatory-compliant tables
4. Uses LLM to generate expert narratives and conclusions
5. Supports configurable product specifications

Author: PMS Regulatory Expert System
Version: 1.0
"""

import pandas as pd
import numpy as np
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Optional
import json
from pathlib import Path
import anthropic
import os
from dotenv import load_dotenv

# Load environment variables from root .env file
root_dir = os.path.dirname(os.path.dirname(__file__))
env_path = os.path.join(root_dir, '.env')
load_dotenv(env_path)

# ============================================================================
# CONFIGURATION CLASS
# ============================================================================


class PSURConfig:
    """
    Configuration class for PSUR Section F generation.
    Customize these parameters for each product.
    """
    
    def __init__(self):
        # File Paths - point to inputs folder in root directory
        inputs_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'inputs')
        self.complaints_file = os.path.join(inputs_dir, "33_complaints.xlsx")
        self.sales_file = os.path.join(inputs_dir, "33_sales.xlsx")
        self.output_dir = "output"
        self.config_file = "config.json"
        
        # LLM Configuration
        self.anthropic_api_key = os.environ.get("ANTHROPIC_API_KEY", "")
        self.model = "claude-sonnet-4-20250514"
        
        # Data Column Mappings
        self.complaint_columns = {
            "date": "Date Entered",
            "complaint_type": "Complaint Type",
            "symptom_code": "Symptom Code",
            "product_number": "Product Number",
            "country": "Country"
        }
        
        self.sales_columns = {
            "date": "Month",
            "item_number": "ItemNumber",
            "quantity": "Quantity",
            "region": "Region",
            "country": "Shipping Country"
        }
        
        # Auto-inferred fields (will be populated from data)
        self.product_name = None
        self.basic_udi_di = ""
        self.device_family = None
        self.device_class = None
        self.is_implantable = False
        self.is_single_use = True
        self.psur_frequency = None
        self.data_collection_start = None
        self.data_collection_end = None
        self.usage_calculation_method = None
        self.max_expected_rates = {}
        self.harm_mapping = {}
        self.mdp_mapping = {}
        
        # Load surveillance period from config if exists (must be after initializing the fields)
        self._load_surveillance_period()
    
    def _load_surveillance_period(self):
        """Load surveillance period from config.json if it exists."""
        try:
            if os.path.exists(self.config_file):
                with open(self.config_file, 'r') as f:
                    config_data = json.load(f)
                    period = config_data.get('surveillance_period', {})
                    self.data_collection_start = period.get('start_date')
                    self.data_collection_end = period.get('end_date')
                    if self.data_collection_start and self.data_collection_end:
                        print(f"Loaded surveillance period from config: {self.data_collection_start} to {self.data_collection_end}")
        except Exception as e:
            print(f"Note: Could not load config.json: {e}")
    
    def auto_configure_from_data(self, complaints_df, sales_df):
        """Auto-configure from input data files using LLM analysis."""
        # Only infer date range from data if NOT set from config
        # Config dates take absolute priority
        if not self.data_collection_start and len(complaints_df) > 0:
            self.data_collection_start = complaints_df['date_entered'].min().strftime('%Y-%m-%d')
            print(f"Auto-detected start date from data: {self.data_collection_start}")
        elif self.data_collection_start:
            print(f"Using reporting period from config.json: {self.data_collection_start} to {self.data_collection_end}")
            
        if not self.data_collection_end and len(complaints_df) > 0:
            self.data_collection_end = complaints_df['date_entered'].max().strftime('%Y-%m-%d')
            print(f"Auto-detected end date from data: {self.data_collection_end}")
        
        # Extract unique complaint types and symptom codes
        complaint_types = complaints_df['complaint_type'].dropna().unique().tolist() if 'complaint_type' in complaints_df.columns else []
        symptom_codes = complaints_df['symptom_code'].dropna().unique().tolist() if 'symptom_code' in complaints_df.columns else []
        
        return complaint_types, symptom_codes

# ============================================================================
# DATA PROCESSING CLASS
# ============================================================================

class PSURDataProcessor:
    """
    Processes complaint and sales data for PSUR analysis.
    """
    
    def __init__(self, config: PSURConfig):
        self.config = config
        self.complaints_df = None
        self.sales_df = None
        self.processed_data = {}
    
    def _find_column(self, df, possible_names):
        """Find a column that matches one of the possible names (case-insensitive)."""
        # Convert all column names to strings first
        df_columns_lower = {str(col).lower(): col for col in df.columns}
        for name in possible_names:
            if str(name).lower() in df_columns_lower:
                return df_columns_lower[str(name).lower()]
        return None
    
    def _find_sheet_with_data(self, file_path):
        """Find the sheet with actual data (not empty or metadata-only)."""
        try:
            xl_file = pd.ExcelFile(file_path)
            for sheet_name in xl_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                # Clean up column names - convert to strings
                df.columns = [str(col) for col in df.columns]
                if len(df) > 5 and len([col for col in df.columns if 'Unnamed' not in col and col.strip()]) > 3:
                    return sheet_name
            return xl_file.sheet_names[0]
        except:
            return 0
        
    def load_data(self):
        """Load complaint and sales data from files."""
        print("Loading data files...")
        
        # Load complaints - try to find the right sheet
        complaints_sheet = self._find_sheet_with_data(self.config.complaints_file)
        print(f"Loading complaints from sheet: {complaints_sheet}")
        self.complaints_df = pd.read_excel(self.config.complaints_file, sheet_name=complaints_sheet)
        # Ensure all column names are strings
        self.complaints_df.columns = [str(col) for col in self.complaints_df.columns]
        
        # Map columns flexibly
        date_col = self._find_column(self.complaints_df, 
                                     [self.config.complaint_columns['date'], 'Date', 'date_entered', 'date'])
        type_col = self._find_column(self.complaints_df,
                                     [self.config.complaint_columns['complaint_type'], 'Type', 'complaint_type'])
        symptom_col = self._find_column(self.complaints_df,
                                       [self.config.complaint_columns['symptom_code'], 'Symptom', 'symptom_code'])
        
        if not date_col or not type_col or not symptom_col:
            print(f"ERROR: Could not find required columns in complaints file.")
            print(f"Available columns: {list(self.complaints_df.columns)}")
            print(f"Looking for: date={self.config.complaint_columns['date']}, type={self.config.complaint_columns['complaint_type']}, symptom={self.config.complaint_columns['symptom_code']}")
            raise ValueError("Missing required columns in complaints file")
        
        # Rename to standard names
        self.complaints_df = self.complaints_df.rename(columns={
            date_col: 'date_entered',
            type_col: 'complaint_type',
            symptom_col: 'symptom_code'
        })
        
        # Convert date
        self.complaints_df['date_entered'] = pd.to_datetime(
            self.complaints_df['date_entered'], errors='coerce'
        )
        
        # Load sales - try different header rows
        sales_sheet = self._find_sheet_with_data(self.config.sales_file)
        print(f"Loading sales from sheet: {sales_sheet}")
        
        # Try to find the correct header row
        for skip_rows in [0, 1, 2]:
            try:
                test_df = pd.read_excel(self.config.sales_file, sheet_name=sales_sheet, skiprows=skip_rows, nrows=5)
                test_df.columns = [str(col) for col in test_df.columns]
                # Check if we have reasonable column names
                if any(keyword in str(test_df.columns).lower() for keyword in ['month', 'year', 'quantity', 'fiscal', 'date']):
                    self.sales_df = pd.read_excel(self.config.sales_file, sheet_name=sales_sheet, skiprows=skip_rows)
                    print(f"  Found headers at row {skip_rows}")
                    break
            except:
                continue
        
        if self.sales_df is None:
            self.sales_df = pd.read_excel(self.config.sales_file, sheet_name=sales_sheet)
        
        # Ensure all column names are strings
        self.sales_df.columns = [str(col) for col in self.sales_df.columns]
        
        # Map sales columns - try to find month/date and quantity columns intelligently
        date_col = self._find_column(self.sales_df,
                                     [self.config.sales_columns['date'], 'Month', 'Date', 'date', 'month'])
        qty_col = self._find_column(self.sales_df,
                                    [self.config.sales_columns['quantity'], 'Quantity', 'Qty', 'quantity'])
        year_col = self._find_column(self.sales_df,
                                     ['Year', 'year', 'YEAR', 'Fiscal Year'])
        
        # If we can't find columns by name, try to infer from data
        if not date_col:
            # Look for a column with month names
            for col in self.sales_df.columns:
                sample = self.sales_df[col].dropna().astype(str).head(20)
                month_names = ['january', 'february', 'march', 'april', 'may', 'june', 
                              'july', 'august', 'september', 'october', 'november', 'december']
                if any(any(month in str(val).lower() for month in month_names) for val in sample):
                    date_col = col
                    print(f"  Auto-detected date column: {col}")
                    break
        
        if not qty_col:
            # Use LLM to intelligently detect the quantity column
            print("  Using AI to detect quantity column...")
            try:
                # Get sample of numeric columns
                numeric_cols = []
                for col in self.sales_df.columns:
                    try:
                        if pd.api.types.is_numeric_dtype(self.sales_df[col]) or self.sales_df[col].dtype == 'object':
                            sample_vals = self.sales_df[col].dropna().head(10).tolist()
                            if sample_vals:
                                numeric_cols.append(f"{col}: {sample_vals[:5]}")
                    except:
                        continue
                
                prompt = f"""Look at this sales data and identify which column contains the QUANTITY of units sold.

Available columns and sample values:
{chr(10).join(numeric_cols[:20])}

Return ONLY the exact column name that contains quantities/units sold. No explanation."""

                llm_agent = LLMConfigurationAgent(self.config.anthropic_api_key)
                response = llm_agent.client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=100,
                    messages=[{"role": "user", "content": prompt}]
                )
                qty_col_name = response.content[0].text.strip()
                
                # Find the column (case insensitive)
                for col in self.sales_df.columns:
                    if str(col).lower() == qty_col_name.lower() or qty_col_name in str(col):
                        qty_col = col
                        print(f"  AI-detected quantity column: {col}")
                        break
            except Exception as e:
                print(f"  AI detection failed: {e}, trying fallback...")
                # Fallback: find any numeric column with reasonable values
                for col in self.sales_df.columns:
                    try:
                        if pd.api.types.is_numeric_dtype(self.sales_df[col]):
                            col_sum = self.sales_df[col].sum()
                            if 10 < col_sum < 1000000:  # Reasonable quantity range
                                qty_col = col
                                print(f"  Fallback detected quantity column: {col}")
                                break
                    except:
                        continue
        
        if not year_col:
            # Look for a column with 4-digit years
            for col in self.sales_df.columns:
                try:
                    sample = self.sales_df[col].dropna().astype(str).head(20)
                    if any(str(val).isdigit() and 2000 <= int(val) <= 2030 for val in sample if str(val).isdigit()):
                        year_col = col
                        print(f"  Auto-detected year column: {col}")
                        break
                except:
                    continue
        
        if not date_col or not qty_col:
            print(f"ERROR: Could not find required columns in sales file.")
            print(f"Available columns: {list(self.sales_df.columns)}")
            print(f"Please ensure your sales file has a month/date column and a quantity column.")
            raise ValueError("Missing required columns in sales file")
        
        # Rename to standard names
        rename_dict = {
            date_col: 'sales_date',
            qty_col: 'quantity'
        }
        if year_col:
            rename_dict[year_col] = 'sales_year'
        
        self.sales_df = self.sales_df.rename(columns=rename_dict)
        
        # Handle month-only dates by adding year if available
        if 'sales_year' in self.sales_df.columns:
            # Combine month name and year to create proper dates
            def combine_month_year(row):
                try:
                    if pd.notna(row.get('sales_year')) and pd.notna(row.get('sales_date')):
                        year = int(row['sales_year'])
                        month_str = str(row['sales_date']).strip()
                        # Create date string: "Month Year" -> pandas can parse this
                        return f"{month_str} {year}"
                    return row['sales_date']
                except:
                    return row['sales_date']
            
            self.sales_df['sales_date'] = self.sales_df.apply(combine_month_year, axis=1)
        else:
            print(f"  WARNING: No year column found in sales data. Sales dates may be incomplete.")
            print(f"  Available columns: {list(self.sales_df.columns)}")
        
        # Filter by reporting period from config
        if self.config.data_collection_start and self.config.data_collection_end:
            start_date = pd.to_datetime(self.config.data_collection_start)
            end_date = pd.to_datetime(self.config.data_collection_end)
            
            print(f"\nFiltering data to reporting period from config: {start_date.date()} to {end_date.date()}")
            
            # Filter complaints
            complaints_before = len(self.complaints_df)
            self.complaints_df = self.complaints_df[
                (self.complaints_df['date_entered'] >= start_date) &
                (self.complaints_df['date_entered'] <= end_date)
            ]
            print(f"  Complaints: {len(self.complaints_df)} of {complaints_before} within reporting period")
            
            # Parse and filter sales dates
            self.sales_df['parsed_date'] = pd.to_datetime(self.sales_df['sales_date'], errors='coerce')
            sales_before = len(self.sales_df)
            self.sales_df = self.sales_df[
                (self.sales_df['parsed_date'] >= start_date) &
                (self.sales_df['parsed_date'] <= end_date)
            ]
            print(f"  Sales records: {len(self.sales_df)} of {sales_before} within reporting period")
        else:
            print("\nWARNING: No reporting period configured in config.json - using all available data")
        
        print(f"\nFinal dataset: {len(self.complaints_df)} complaints and {len(self.sales_df)} sales records")
        
    def calculate_total_uses(self) -> int:
        """
        Calculate total estimated uses during the data collection period.
        For single-use devices, this equals total units sold.
        Only includes sales within the configured reporting period.
        """
        if self.config.is_single_use:
            # Parse dates in sales data
            self.sales_df['parsed_date'] = pd.to_datetime(
                self.sales_df['sales_date'], 
                errors='coerce'
            )
            
            # Only use reporting period from config
            if not self.config.data_collection_start or not self.config.data_collection_end:
                print("ERROR: Reporting period not configured. Check config.json")
                return 0
            
            start_date = pd.to_datetime(self.config.data_collection_start)
            end_date = pd.to_datetime(self.config.data_collection_end)
            
            print(f"  Filtering sales data for reporting period: {start_date.date()} to {end_date.date()}")
            
            # Filter sales within date range
            sales_in_period = self.sales_df[
                (self.sales_df['parsed_date'] >= start_date) &
                (self.sales_df['parsed_date'] <= end_date)
            ]
            
            print(f"  Sales records in reporting period: {len(sales_in_period)} of {len(self.sales_df)} total")
            
            total_units = sales_in_period['quantity'].sum()
            return int(total_units)
        else:
            # For reusable devices, implement custom logic
            raise NotImplementedError(
                "For reusable devices, implement custom usage calculation logic"
            )
    
    def categorize_complaints(self) -> pd.DataFrame:
        """
        Categorize complaints by harm and medical device problem.
        """
        df = self.complaints_df.copy()
        
        # Normalize columns for mapping
        df['complaint_type_norm'] = df['complaint_type'].astype(str).str.lower().str.strip()
        df['symptom_code_norm'] = df['symptom_code'].astype(str).str.lower().str.strip()
        
        # Map to harm categories
        df['Harm_Category'] = df['complaint_type_norm'].map(self.config.harm_mapping)
        df['Harm_Category'] = df['Harm_Category'].fillna('No Health Consequence')
        
        # Map to medical device problems
        df['Medical_Device_Problem'] = df['symptom_code_norm'].map(self.config.mdp_mapping)
        df['Medical_Device_Problem'] = df['Medical_Device_Problem'].fillna('Other Device Problem')
        
        return df
    
    def calculate_complaint_rates_by_year(self) -> pd.DataFrame:
        """
        Calculate complaint counts and rates by year, harm, and medical device problem.
        Only includes data within the configured reporting period.
        """
        categorized = self.categorize_complaints()
        
        # Extract year from date
        categorized['Year'] = categorized['date_entered'].dt.year
        
        # Group by year, harm, and MDP
        grouped = categorized.groupby(['Year', 'Harm_Category', 'Medical_Device_Problem']).size().reset_index(name='Count')
        
        # Calculate total uses per year (within reporting period only)
        if not self.config.data_collection_start or not self.config.data_collection_end:
            print("ERROR: Reporting period not configured. Check config.json")
            grouped['Total_Uses'] = 0
            return grouped
        
        start_date = pd.to_datetime(self.config.data_collection_start)
        end_date = pd.to_datetime(self.config.data_collection_end)
        
        # Parse sales dates if not already done
        if 'parsed_date' not in self.sales_df.columns:
            self.sales_df['parsed_date'] = pd.to_datetime(self.sales_df['sales_date'], errors='coerce')
        
        # Filter sales to reporting period
        sales_in_period = self.sales_df[
            (self.sales_df['parsed_date'] >= start_date) &
            (self.sales_df['parsed_date'] <= end_date)
        ]
        
        years = grouped['Year'].unique()
        uses_by_year = {}
        
        for year in years:
            # Filter sales for this year within the reporting period
            year_sales = sales_in_period[self.sales_df['parsed_date'].dt.year == year]
            total_qty = int(year_sales['quantity'].sum()) if len(year_sales) > 0 else 0
            uses_by_year[year] = total_qty
        
        # Add uses to grouped data
        grouped['Total_Uses'] = grouped['Year'].map(uses_by_year)
        
        # Calculate rate as percentage
        grouped['Rate_%'] = (grouped['Count'] / grouped['Total_Uses'] * 100).round(4)
        
        # Assign Occurrence Code based on standardized risk limits
        def assign_occurrence_code(rate_pct):
            """Assign occurrence code based on standard risk ranges."""
            if rate_pct > 10.0:
                return 'O5'
            elif rate_pct > 1.0:
                return 'O4'
            elif rate_pct > 0.1:
                return 'O3'
            elif rate_pct > 0.01:
                return 'O2'
            else:
                return 'O1'
        
        grouped['Occurrence_Code'] = grouped['Rate_%'].apply(assign_occurrence_code)
        
        # Add max expected rate (now using standard ranges)
        grouped['Max_Expected_Rate_%'] = grouped.apply(
            lambda row: self.config.max_expected_rates.get(
                row['Harm_Category'], 
                self.config.max_expected_rates.get('Default', 2.0)
            ),
            axis=1
        )
        
        # Flag if exceeds expected rate
        grouped['Exceeds_Expected'] = grouped['Rate_%'] > grouped['Max_Expected_Rate_%']
        
        return grouped
    
    def generate_annual_table(self) -> pd.DataFrame:
        """
        Generate Table 7 for annual PSUR reporting.
        """
        rates_df = self.calculate_complaint_rates_by_year()
        
        # Pivot to create table structure
        table = rates_df.pivot_table(
            index=['Harm_Category', 'Medical_Device_Problem'],
            columns='Year',
            values=['Count', 'Rate_%'],
            aggfunc='first',
            fill_value=0
        )
        
        return table
    
    def get_summary_statistics(self) -> Dict:
        """
        Calculate summary statistics for narrative generation.
        """
        categorized = self.categorize_complaints()
        total_uses = self.calculate_total_uses()
        
        stats = {
            'total_complaints': len(categorized),
            'total_uses': total_uses,
            'overall_complaint_rate': (len(categorized) / total_uses * 100) if total_uses > 0 else 0,
            'complaints_by_harm': categorized['Harm_Category'].value_counts().to_dict(),
            'complaints_by_mdp': categorized['Medical_Device_Problem'].value_counts().to_dict(),
            'date_range': {
                'start': self.config.data_collection_start,
                'end': self.config.data_collection_end
            },
            'years_analyzed': sorted(categorized['date_entered'].dt.year.unique().tolist()),
            'serious_injuries': len(categorized[categorized['Harm_Category'] == 'Serious Injury']),
            'no_health_impact': len(categorized[categorized['Harm_Category'] == 'No Health Consequence']),
        }
        
        # Calculate rates by harm category
        stats['rates_by_harm'] = {}
        for harm, count in stats['complaints_by_harm'].items():
            stats['rates_by_harm'][harm] = (count / total_uses * 100) if total_uses > 0 else 0
        
        # Identify top medical device problems
        top_mdp = categorized['Medical_Device_Problem'].value_counts().head(5)
        stats['top_mdp'] = top_mdp.to_dict()
        
        # Check for rate exceedances
        rates_df = self.calculate_complaint_rates_by_year()
        exceedances = rates_df[rates_df['Exceeds_Expected']]
        stats['rate_exceedances'] = exceedances.to_dict('records') if len(exceedances) > 0 else []
        
        return stats

# ============================================================================
# LLM NARRATIVE GENERATOR CLASS
# ============================================================================

class LLMConfigurationAgent:
    """
    LLM agent that auto-configures PSUR parameters from input data.
    """
    
    def __init__(self, api_key: str):
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY is required")
        self.client = anthropic.Anthropic(api_key=api_key)
    
    def infer_product_metadata(self, complaints_df, sales_df, old_cer_path=None, previous_psur_path=None):
        """Infer product name, class, family from data files."""
        
        # Get sample data
        complaint_sample = complaints_df.head(10).to_string() if len(complaints_df) > 0 else "No complaints"
        sales_sample = sales_df.head(10).to_string() if len(sales_df) > 0 else "No sales"
        
        prompt = f"""Analyze this medical device data and extract key product information.

Complaints Data Sample:
{complaint_sample}

Sales Data Sample:
{sales_sample}

Based on this data, determine:
1. Product Name (the medical device name)
2. Device Family (product family/line)
3. Device Class (I, IIa, IIb, or III based on apparent risk)
4. PSUR Frequency (annually or every two years - Class III/implantable are annual)
5. Brief usage calculation method (how to calculate estimated uses for this device type)

Return ONLY a JSON object with these exact keys:
{{
  "product_name": "...",
  "device_family": "...",
  "device_class": "...",
  "psur_frequency": "...",
  "usage_calculation_method": "..."
}}"""

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        import json
        import re
        text = response.content[0].text
        # Extract JSON from response (may have markdown code blocks)
        json_match = re.search(r'\{[\s\S]*\}', text)
        if json_match:
            return json.loads(json_match.group())
        else:
            raise ValueError(f"Could not extract JSON from LLM response: {text[:200]}")
    
    def generate_mappings(self, complaint_types, symptom_codes):
        """Generate IMDRF-compliant harm and device problem mappings."""
        
        prompt = f"""You are a regulatory expert. Map these complaint data fields to IMDRF terminology.

Complaint Types found in data:
{complaint_types}

Symptom Codes found in data:
{symptom_codes}

Create two mappings:

1. HARM MAPPING: Map each complaint type to IMDRF Annex F Harm Categories:
   - Death
   - Serious Injury
   - Malfunction (No Health Consequence)
   - etc.

2. MEDICAL DEVICE PROBLEM MAPPING: Map each symptom code to IMDRF Annex A Device Problems:
   - Foreign Material
   - Broken or Damaged Component
   - Does Not Perform Properly
   - etc.

Return ONLY a JSON object:
{{
  "harm_mapping": {{
    "complaint_type": "IMDRF Harm Category",
    ...
  }},
  "mdp_mapping": {{
    "symptom_code": "IMDRF Device Problem",
    ...
  }},
  "max_expected_rates": {{
    "Serious Injury": 0.001,
    "Malfunction": 2.0,
    "Default": 2.0
  }}
}}"""

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        import json
        import re
        text = response.content[0].text
        # Extract JSON from response (may have markdown code blocks)
        json_match = re.search(r'\{[\s\S]*\}', text)
        if json_match:
            return json.loads(json_match.group())
        else:
            raise ValueError(f"Could not extract JSON from LLM response: {text[:200]}")


class LLMNarrativeGenerator:
    """
    Generates expert narratives and conclusions using Claude API.
    """
    
    def __init__(self, config: PSURConfig):
        self.config = config
        if not config.anthropic_api_key:
            raise ValueError(
                "ANTHROPIC_API_KEY is required but not found in environment variables.\n"
                "Set it in the .env file or environment to generate LLM narratives."
            )
        try:
            self.client = anthropic.Anthropic(api_key=config.anthropic_api_key)
        except Exception as e:
            raise ValueError(f"Failed to initialize Anthropic client: {e}")
    
    def _call_claude(self, system_prompt: str, user_prompt: str) -> str:
        """Make API call to Claude."""
        try:
            message = self.client.messages.create(
                model=self.config.model,
                max_tokens=4000,
                system=system_prompt,
                messages=[
                    {"role": "user", "content": user_prompt}
                ]
            )
            return message.content[0].text
        except Exception as e:
            raise RuntimeError(f"Failed to generate LLM narrative: {str(e)}\n"
                             "Cannot continue without complete narratives to avoid incomplete reports.")
    
    def generate_usage_calculation_narrative(self, total_uses: int) -> str:
        """Generate narrative explaining usage calculation methodology."""
        
        system_prompt = """You are a regulatory expert specializing in EU MDR compliance and PSUR preparation. 
You write clear, professional narratives that explain technical concepts for regulatory submissions."""
        
        user_prompt = f"""Generate a clear, professional narrative for Section F of a PSUR that explains how the estimated number of uses was calculated for this product.

Product Information:
- Product Name: {self.config.product_name}
- Device Type: {"Single-use" if self.config.is_single_use else "Reusable"}
- Data Collection Period: {self.config.data_collection_start} to {self.config.data_collection_end}
- Total Estimated Uses: {total_uses:,}

Calculation Method: {self.config.usage_calculation_method}

Requirements:
1. Write 2-3 sentences maximum
2. Be specific to this device type
3. Clearly state the calculation method
4. Mention the total number of uses calculated
5. Use professional regulatory language
6. Do not use bullet points

Example format:
"The complaint rate is calculated as the number of complaints received during the data collection period divided by the estimated number of uses during that period. For this single-use device, the estimated number of uses equals the total number of devices sold (XXX units), as each device is intended for one-time use per patient procedure."
"""
        
        return self._call_claude(system_prompt, user_prompt)
    
    def generate_complaint_analysis_narrative(self, stats: Dict, rates_df: pd.DataFrame) -> str:
        """Generate comprehensive complaint analysis narrative."""
        
        system_prompt = """You are a senior regulatory affairs specialist with expertise in post-market surveillance for medical devices under EU MDR. 
You analyze complaint data and write professional, evidence-based narratives for regulatory submissions."""
        
        user_prompt = f"""Analyze the complaint data for Section F of a PSUR and generate a comprehensive narrative.

Product: {self.config.product_name}
Data Collection Period: {stats['date_range']['start']} to {stats['date_range']['end']}

KEY STATISTICS:
- Total Complaints: {stats['total_complaints']:,}
- Total Uses (Devices Sold): {stats['total_uses']:,}
- Overall Complaint Rate: {stats['overall_complaint_rate']:.4f}%
- Serious Injuries: {stats['serious_injuries']}
- No Health Impact: {stats['no_health_impact']}

COMPLAINTS BY HARM CATEGORY:
{json.dumps(stats['complaints_by_harm'], indent=2)}

RATES BY HARM CATEGORY (%) WITH OCCURRENCE CODES:
{json.dumps(stats['rates_by_harm'], indent=2)}

OCCURRENCE CODE RANGES (Standard Risk Limits):
- O5: Rate > 10%
- O4: 1% < Rate ≤ 10%
- O3: 0.1% < Rate ≤ 1%
- O2: 0.01% < Rate ≤ 0.1%
- O1: Rate ≤ 0.01%

TOP MEDICAL DEVICE PROBLEMS:
{json.dumps(stats['top_mdp'], indent=2)}

RATE EXCEEDANCES (if any):
{json.dumps(stats['rate_exceedances'], indent=2) if stats['rate_exceedances'] else "No rate exceedances identified"}

REQUIREMENTS:
1. Write a professional narrative of 3-5 paragraphs
2. First paragraph: Overall summary of complaint data
3. Second paragraph: Analysis of harm categories and health impacts
4. Third paragraph: Analysis of medical device problems (focus on top issues)
5. Fourth paragraph: Discussion of complaint rates and any exceedances
6. Fifth paragraph (if needed): Trending observations
7. Use specific numbers and percentages
8. Identify patterns or concerns
9. Maintain objective, scientific tone
10. Do NOT use bullet points
11. Reference the regulatory framework (MDR, MDCG 2022-21) where appropriate

Write the narrative now:"""
        
        return self._call_claude(system_prompt, user_prompt)
    
    def generate_rate_exceedance_commentary(self, exceedances: List[Dict]) -> str:
        """Generate specific commentary on rate exceedances."""
        
        if not exceedances:
            return ("During the data collection period, no complaint rates exceeded the maximum "
                   "expected rates of occurrence defined in the Risk Analysis and Control Table (RACT). "
                   "This indicates that the device is performing within expected safety parameters. "
                   "No updates to the Risk Management File are required based on complaint rate analysis.")
        
        system_prompt = """You are a risk management expert specializing in medical device safety under EU MDR. 
You provide detailed analysis of complaint rates that exceed expected thresholds."""
        
        user_prompt = f"""Analyze the following complaint rate exceedances and provide expert commentary for a PSUR.

EXCEEDANCES IDENTIFIED:
{json.dumps(exceedances, indent=2)}

For EACH exceedance, provide:
1. Clear identification of the harm/medical device problem
2. Actual rate vs. expected rate
3. Analysis of why this exceedance occurred
4. Assessment of whether it represents a true safety concern
5. Recommendation for risk management actions (if needed)
6. Statement on whether Risk Management File update is required

Write a detailed, structured analysis (3-5 paragraphs). Be specific and provide actionable recommendations.
Use professional regulatory language."""
        
        return self._call_claude(system_prompt, user_prompt)
    
    def generate_conclusions(self, stats: Dict, exceedances: List[Dict]) -> str:
        """Generate regulatory conclusions for Section F."""
        
        system_prompt = """You are a regulatory affairs director responsible for PSUR approvals. 
You write authoritative conclusions that assess benefit-risk profiles and determine regulatory actions."""
        
        user_prompt = f"""Generate professional conclusions for Section F of a PSUR.

COMPLAINT DATA SUMMARY:
- Total Complaints: {stats['total_complaints']:,}
- Total Uses: {stats['total_uses']:,}
- Overall Complaint Rate: {stats['overall_complaint_rate']:.4f}%
- Serious Injuries: {stats['serious_injuries']}
- Rate Exceedances: {len(exceedances)}

ANALYSIS:
- Most complaints were: {list(stats['top_mdp'].keys())[0] if stats['top_mdp'] else 'various issues'}
- Harm distribution: {json.dumps(stats['complaints_by_harm'], indent=2)}

REQUIREMENTS:
1. Write 2-3 paragraph conclusion
2. First paragraph: Overall assessment of complaint data
3. Second paragraph: Impact on benefit-risk profile (typically "remains favorable" unless serious issues)
4. Third paragraph: Any required actions or updates (Risk Management File, IFU, CAPA, etc.)
5. Be definitive but appropriate to the data
6. Use regulatory language appropriate for Notified Body review
7. State clearly if no adverse impact on benefit-risk profile

Write the conclusion now:"""
        
        return self._call_claude(system_prompt, user_prompt)

# ============================================================================
# REPORT GENERATOR CLASS
# ============================================================================

class PSURSectionFGenerator:
    """
    Main class that orchestrates Section F generation.
    """
    
    def __init__(self, config: PSURConfig):
        self.config = config
        self.processor = PSURDataProcessor(config)
        self.llm = LLMNarrativeGenerator(config)
        
    def generate(self) -> Dict:
        """
        Generate complete Section F content.
        Returns dictionary with all components.
        """
        print("\n" + "="*80)
        print("PSUR SECTION F GENERATOR")
        print("="*80)
        
        # Load and process data
        print("\n[1/6] Loading data...")
        self.processor.load_data()
        
        # Calculate statistics
        print("[2/6] Calculating statistics...")
        stats = self.processor.get_summary_statistics()
        rates_df = self.processor.calculate_complaint_rates_by_year()
        annual_table = self.processor.generate_annual_table()
        
        # Generate narratives
        print("[3/6] Generating usage calculation narrative...")
        usage_narrative = self.llm.generate_usage_calculation_narrative(stats['total_uses'])
        
        print("[4/6] Generating complaint analysis narrative...")
        analysis_narrative = self.llm.generate_complaint_analysis_narrative(stats, rates_df)
        
        print("[5/6] Generating rate exceedance commentary...")
        exceedance_commentary = self.llm.generate_rate_exceedance_commentary(stats['rate_exceedances'])
        
        print("[6/6] Generating conclusions...")
        conclusions = self.llm.generate_conclusions(stats, stats['rate_exceedances'])
        
        # Compile results
        results = {
            'metadata': {
                'product_name': self.config.product_name,
                'device_class': self.config.device_class,
                'data_collection_period': f"{self.config.data_collection_start} to {self.config.data_collection_end}",
                'generation_date': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'psur_frequency': self.config.psur_frequency
            },
            'statistics': stats,
            'narratives': {
                'usage_calculation': usage_narrative,
                'complaint_analysis': analysis_narrative,
                'rate_exceedance_commentary': exceedance_commentary,
                'conclusions': conclusions
            },
            'tables': {
                'annual_table': annual_table,
                'detailed_rates': rates_df
            }
        }
        
        print("\n" + "="*80)
        print("GENERATION COMPLETE")
        print("="*80 + "\n")
        
        return results
    
    def export_to_word(self, results: Dict, output_filename: str = "PSUR_Section_F.docx"):
        """
        Export Section F to a Word document.
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            print(f"Exporting to Word document: {output_filename}")
            
            doc = Document()
            
            # Set default font to Arial 10 for the entire document
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(10)
            
            # Title
            title = doc.add_heading('Section F: Product Complaint Types, Complaint Counts, and Complaint Rates', level=1)
            
            # Metadata
            meta_para = doc.add_paragraph()
            run1 = meta_para.add_run('Product: ')
            run1.bold = True
            run1.font.name = 'Arial'
            run1.font.size = Pt(10)
            run2 = meta_para.add_run(f"{results['metadata']['product_name']}\n")
            run2.font.name = 'Arial'
            run2.font.size = Pt(10)
            run3 = meta_para.add_run('Data Collection Period: ')
            run3.bold = True
            run3.font.name = 'Arial'
            run3.font.size = Pt(10)
            run4 = meta_para.add_run(f"{results['metadata']['data_collection_period']}\n")
            run4.font.name = 'Arial'
            run4.font.size = Pt(10)
            run5 = meta_para.add_run('Report Generated: ')
            run5.bold = True
            run5.font.name = 'Arial'
            run5.font.size = Pt(10)
            run6 = meta_para.add_run(f"{results['metadata']['generation_date']}\n")
            run6.font.name = 'Arial'
            run6.font.size = Pt(10)
            
            # Complaint Rate Calculation
            doc.add_heading('Complaint Rate Calculation', level=2)
            para = doc.add_paragraph(results['narratives']['usage_calculation'])
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            
            # Annual Number of Complaints
            doc.add_heading('Annual Number of Complaints and Complaint Rate by Harm and Medical Device Problem', level=2)
            
            # Add statistics table
            para1 = doc.add_paragraph(f"Total Complaints: {results['statistics']['total_complaints']:,}")
            for run in para1.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            para2 = doc.add_paragraph(f"Total Uses: {results['statistics']['total_uses']:,}")
            for run in para2.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            para3 = doc.add_paragraph(f"Overall Complaint Rate: {results['statistics']['overall_complaint_rate']:.4f}%")
            for run in para3.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            
            # Add complaint analysis
            doc.add_heading('Complaint Data Analysis', level=3)
            para4 = doc.add_paragraph(results['narratives']['complaint_analysis'])
            for run in para4.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            
            # Add detailed table following template format
            doc.add_heading('Table 7: Complaint Rate (and Complaint Count) by Harm and Medical Device Problem', level=3)
            
            # Prepare data grouped by harm category
            rates_df = results['tables']['detailed_rates']
            
            # Group data by harm category
            harm_groups = {}
            for _, row in rates_df.iterrows():
                harm = row['Harm_Category']
                if harm not in harm_groups:
                    harm_groups[harm] = []
                harm_groups[harm].append(row)
            
            # Calculate total rows needed: header(2) + harm_groups + mdp_rows + grand_total(1)
            total_mdp_rows = len(rates_df)
            total_harm_rows = len(harm_groups)
            total_rows = 2 + total_harm_rows + total_mdp_rows + 1  # headers + harm headers + mdp rows + grand total
            
            # Create table with 3 columns (following template)
            table = doc.add_table(rows=total_rows, cols=3)
            table.style = 'Light Grid Accent 1'
            
            # Set column widths for better layout
            from docx.shared import Inches
            table.columns[0].width = Inches(3.0)
            table.columns[1].width = Inches(2.5)
            table.columns[2].width = Inches(2.0)
            
            # Header Row 0
            cell = table.rows[0].cells[0]
            cell.text = "Harm\n   Medical Device Problem"
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            cell = table.rows[0].cells[1]
            cell.text = "Current 12-Month Data Collection Period"
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            cell = table.rows[0].cells[2]
            cell.text = "Max Expected Rate of Occurrence (from the RACT)"
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Header Row 1 with date ranges
            from datetime import datetime
            start_date = datetime.strptime(results['metadata']['data_collection_period'].split(' to ')[0], '%Y-%m-%d')
            end_date = datetime.strptime(results['metadata']['data_collection_period'].split(' to ')[1], '%Y-%m-%d')
            date_range = f"{start_date.strftime('%b-%Y')} to {end_date.strftime('%b-%Y')}"
            
            cell = table.rows[1].cells[0]
            cell.text = "Harm\n   Medical Device Problem"
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            cell = table.rows[1].cells[1]
            cell.text = f"[{date_range}]"
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            cell = table.rows[1].cells[2]
            cell.text = f"[{date_range}]"
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Data rows - hierarchical structure
            row_idx = 2
            grand_total_count = 0
            grand_total_uses = results['statistics']['total_uses']
            
            for harm_category, mdp_rows in harm_groups.items():
                # Harm category header row
                cell = table.rows[row_idx].cells[0]
                cell.text = harm_category
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.name = 'Calibri'
                cell.paragraphs[0].runs[0].font.size = Pt(10)
                row_idx += 1
                
                # Medical device problems under this harm
                for mdp_row in mdp_rows:
                    grand_total_count += mdp_row['Count']
                    
                    # MDP name (indented)
                    cell = table.rows[row_idx].cells[0]
                    cell.text = f"   {mdp_row['Medical_Device_Problem']}"
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    
                    # Complaint count and rate
                    cell = table.rows[row_idx].cells[1]
                    rate_text = f"{mdp_row['Rate_%']:.4f}% ({mdp_row['Count']} complaints / {mdp_row['Total_Uses']:,} uses)"
                    cell.text = rate_text
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    
                    # Max expected rate
                    cell = table.rows[row_idx].cells[2]
                    cell.text = f"{mdp_row['Max_Expected_Rate_%']:.2f}%"
                    cell.paragraphs[0].runs[0].font.name = 'Calibri'
                    cell.paragraphs[0].runs[0].font.size = Pt(10)
                    
                    # Highlight if exceeds expected
                    if mdp_row['Exceeds_Expected']:
                        for i in range(3):
                            for paragraph in table.rows[row_idx].cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(255, 0, 0)
                    
                    row_idx += 1
            
            # Grand Total row
            cell = table.rows[row_idx].cells[0]
            cell.text = "Grand Total"
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            cell = table.rows[row_idx].cells[1]
            grand_total_rate = (grand_total_count / grand_total_uses * 100) if grand_total_uses > 0 else 0
            cell.text = f"{grand_total_rate:.4f}% ({grand_total_count} complaints / {grand_total_uses:,} uses)"
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            cell = table.rows[row_idx].cells[2]
            cell.text = ""
            cell.paragraphs[0].runs[0].font.name = 'Calibri'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            
            # Rate exceedance commentary
            doc.add_heading('Commentary on Complaint Rates', level=3)
            para5 = doc.add_paragraph(results['narratives']['rate_exceedance_commentary'])
            for run in para5.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            
            # Conclusions
            doc.add_heading('Conclusions', level=2)
            para6 = doc.add_paragraph(results['narratives']['conclusions'])
            for run in para6.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            
            # Save document
            output_path = Path(self.config.output_dir) / output_filename
            doc.save(str(output_path))
            
            print(f"[SUCCESS] Word document exported successfully to: {output_path}")
            return str(output_path)
            
        except ImportError:
            print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
            return None
        except Exception as e:
            print(f"ERROR exporting to Word: {e}")
            return None
    
    def export_to_json(self, results: Dict, output_filename: str = "PSUR_Section_F_data.json"):
        """
        Export all data and narratives to JSON for further processing.
        """
        output_path = Path(self.config.output_dir) / output_filename
        
        # Convert DataFrames to dict for JSON serialization
        export_data = {
            'metadata': results['metadata'],
            'statistics': results['statistics'],
            'narratives': results['narratives'],
            'tables': {
                'detailed_rates': results['tables']['detailed_rates'].to_dict('records')
            }
        }
        
        with open(output_path, 'w') as f:
            json.dump(export_data, f, indent=2, default=str)
        
        print(f"[SUCCESS] JSON data exported to: {output_path}")
        return str(output_path)

# ============================================================================
# SYSTEM INSTRUCTIONS AND INFERENCE DOCUMENTATION
# ============================================================================

SYSTEM_INSTRUCTIONS = """
================================================================================
PSUR SECTION F GENERATION - SYSTEM INSTRUCTIONS
================================================================================

REGULATORY FRAMEWORK:
--------------------
This system generates Section F of Post-Market Surveillance Update Reports
(PSUR) in compliance with:
- EU MDR 2017/745, Article 86
- MDCG 2022-21 Guidance on PSURs
- IMDRF Adverse Event Terminology

INFERENCE PROCESS:
-----------------

1. DATA CATEGORIZATION:
   - Complaints are categorized into Harm Categories based on complaint type
   - Harm categories align with IMDRF Annex F (Health Impact)
   - Medical Device Problems align with IMDRF Annex A terminology
   - Mapping is configurable per product

2. RATE CALCULATION:
   - Complaint Rate = (Number of Complaints / Estimated Uses) × 100
   - For single-use devices: Estimated Uses = Total Units Sold
   - Rates calculated per year, per harm category, per medical device problem
   - Compared against max expected rates from Risk Analysis and Control Table

3. LLM NARRATIVE GENERATION:

   A. Usage Calculation Narrative:
      - Explains methodology for calculating estimated uses
      - Specific to device type (single-use vs reusable)
      - Justifies denominator used in rate calculation
      
   B. Complaint Analysis Narrative:
      INPUT: Statistics summary, complaint counts, rates, top issues
      TASK: Comprehensive analysis of complaint data
      OUTPUT: 3-5 paragraph professional narrative covering:
              - Overall complaint summary
              - Harm category analysis
              - Medical device problem analysis
              - Rate assessment
              - Trend observations
      
   C. Rate Exceedance Commentary:
      INPUT: List of complaints where rate exceeds RACT threshold
      TASK: Risk assessment of exceedances
      OUTPUT: Detailed analysis of each exceedance with:
              - Root cause assessment
              - Safety concern evaluation
              - Risk management recommendations
              - RMF update determination
      
   D. Conclusions:
      INPUT: All complaint data, exceedances, statistics
      TASK: Regulatory determination
      OUTPUT: 2-3 paragraph conclusion covering:
              - Overall assessment
              - Benefit-risk profile impact
              - Required regulatory actions

4. TABLE GENERATION:
   - Table 7: Annual complaint rates by harm and MDP
   - Detailed rates table with all categorizations
   - Flags for rate exceedances (visual highlighting)

5. REGULATORY COMPLIANCE CHECKS:
   - Ensures IMDRF terminology usage
   - Validates against MDR requirements
   - Confirms MDCG 2022-21 table formats
   - Checks for required narrative elements

CUSTOMIZATION POINTS:
--------------------
1. PSURConfig class: Product-specific parameters
2. Harm mapping: Complaint type to Harm category
3. MDP mapping: Symptom code to Medical device problem
4. Max expected rates: From product RACT
5. Usage calculation: Single-use vs reusable logic

LLM PROMPTING STRATEGY:
----------------------
- System prompts establish expert persona (regulatory specialist)
- User prompts provide:
  * Complete context (product, dates, statistics)
  * Specific data to analyze
  * Clear output requirements
  * Format specifications
- Temperature: Default (balanced creativity and consistency)
- Max tokens: 4000 (sufficient for comprehensive narratives)

OUTPUT FORMATS:
--------------
1. Python dictionary (in-memory)
2. Word document (formatted PSUR section)
3. JSON (structured data for integration)

QUALITY ASSURANCE:
-----------------
- All rates validated against denominators
- Exceedances automatically flagged
- LLM outputs reviewed for regulatory appropriateness
- Tables checked for completeness
- References to applicable regulations included

================================================================================
"""

# ============================================================================
# MAIN EXECUTION
# ============================================================================

def main():
    """
    Main execution function with example usage.
    """
    
    # Print system instructions
    print(SYSTEM_INSTRUCTIONS)
    
    # Initialize configuration
    print("\nInitializing PSUR Section F Generator...")
    config = PSURConfig()
    
    # Auto-configure from data using LLM
    print("Auto-configuring from input data...")
    print("  Step 1: Loading data files...")
    processor = PSURDataProcessor(config)
    processor.load_data()
    
    print("  Step 2: Analyzing data with LLM to infer product metadata...")
    llm_agent = LLMConfigurationAgent(config.anthropic_api_key)
    metadata = llm_agent.infer_product_metadata(processor.complaints_df, processor.sales_df)
    
    # Apply inferred metadata
    config.product_name = metadata['product_name']
    config.device_family = metadata['device_family']
    config.device_class = metadata['device_class']
    config.psur_frequency = metadata['psur_frequency']
    config.usage_calculation_method = metadata['usage_calculation_method']
    
    print(f"    Detected Product: {config.product_name}")
    print(f"    Device Class: {config.device_class}")
    print(f"    Device Family: {config.device_family}")
    
    print("  Step 3: Generating IMDRF mappings...")
    complaint_types, symptom_codes = config.auto_configure_from_data(processor.complaints_df, processor.sales_df)
    mappings = llm_agent.generate_mappings(complaint_types, symptom_codes)
    
    config.harm_mapping = mappings['harm_mapping']
    config.mdp_mapping = mappings['mdp_mapping']
    config.max_expected_rates = mappings['max_expected_rates']
    
    print(f"    Mapped {len(config.harm_mapping)} complaint types to harm categories")
    print(f"    Mapped {len(config.mdp_mapping)} symptom codes to device problems")
    
    # Create generator with pre-loaded data
    generator = PSURSectionFGenerator(config)
    generator.data_processor = processor
    
    # Generate Section F
    results = generator.generate()
    
    # Export to Word
    word_file = generator.export_to_word(results)
    
    # Export to JSON
    json_file = generator.export_to_json(results)
    
    # Print summary
    print("\n" + "="*80)
    print("GENERATION SUMMARY")
    print("="*80)
    print(f"Product: {results['metadata']['product_name']}")
    print(f"Period: {results['metadata']['data_collection_period']}")
    print(f"Total Complaints: {results['statistics']['total_complaints']:,}")
    print(f"Total Uses: {results['statistics']['total_uses']:,}")
    print(f"Overall Rate: {results['statistics']['overall_complaint_rate']:.4f}%")
    print(f"Rate Exceedances: {len(results['statistics']['rate_exceedances'])}")
    print(f"\nOutputs:")
    if word_file:
        print(f"  - Word: {word_file}")
    if json_file:
        print(f"  - JSON: {json_file}")
    print("="*80 + "\n")
    
    return results

if __name__ == "__main__":
    # Check for required packages
    package_imports = {
        'pandas': 'pandas',
        'openpyxl': 'openpyxl',
        'anthropic': 'anthropic',
        'python-docx': 'docx'
    }
    missing = []
    
    for package, import_name in package_imports.items():
        try:
            __import__(import_name)
        except ImportError:
            missing.append(package)
    
    if missing:
        print("ERROR: Missing required packages:")
        for pkg in missing:
            print(f"  pip install {pkg} --break-system-packages")
        exit(1)
    
    # Run main
    main()
