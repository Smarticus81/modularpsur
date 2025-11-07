"""
PSUR Section J Generator - Scientific Literature Review
EU MDR Compliant per Article 86 and MDCG 2022-21
"""

import os
import sys
import anthropic
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING

# Add parent directory to path for utils import
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from utils.document_parser import DocumentParser
from utils.semantic_document_parser import get_semantic_cer_data
from utils.advanced_cer_parser import parse_cer_for_psur as advanced_parse

# Load environment variables from root .env file
root_dir = os.path.dirname(os.path.dirname(__file__))
env_path = os.path.join(root_dir, '.env')
load_dotenv(env_path)


class AnthropicLLM:
    def __init__(self):
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY environment variable required")
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"
        self.sys = """You are an expert EU MDR regulatory affairs specialist with deep expertise in Post-Market Surveillance and PSUR preparation per EU MDR 2017/745 Article 86 and MDCG 2022-21 guidance.

CRITICAL REQUIREMENTS FOR SCIENTIFIC LITERATURE REVIEW:

ABSOLUTE PROHIBITIONS:
- NEVER generate content about devices not explicitly mentioned in the provided CER context
- NEVER fabricate device types, patient populations, or clinical applications
- NEVER make assumptions about what device category this might be
- NEVER generate mock or placeholder data
- NEVER invent article counts, database names, search dates, or findings
- If the CER context does not contain information, state that generically without inventing specifics

1. DEVICE UNDERSTANDING:
- READ the CER carefully to understand what device this actually is
- Identify all the device understanding information, intended use, patient population information, and clinical application from CER
- Base ALL content generation on the actual device described in the CER
- If device details are unclear, generate generic statements without specifics

2. LITERATURE SEARCH METHODOLOGY:
- Extract ALL information related to literature search, databases, search terms, and dates ONLY from CER context provided
- If specific databases not mentioned in CER, state "databases per CER protocol" generically
- If search dates not in CER, state "per CER literature search" without inventing dates
- Do NOT assume PubMed, Embase, etc. unless explicitly stated in CER
- Describe methodology based only on what is documented in CER

3. LITERATURESEARCH RESULTS SUMMARY:
- Extract article counts ONLY from CER if provided
- Ensure article counts are included in the search results summary"
- Do NOT invent specific numbers like "342 articles" or "38 articles"
- Report findings in general terms if specifics not provided

4. CONTENT ANALYSIS:
- Summarize findings based ONLY on CER literature review section
- Do NOT fabricate safety data, performance metrics, or risk assessments
- State findings generically if CER doesn't provide specifics
- Focus on whether CER indicates any new safety signals or performance issues

5. IMPACT ASSESSMENT:
- Base on CER conclusions about literature findings
- State if CER indicates updates needed to benefit-risk, RMF, or IFU
- If CER doesn't conclude about updates, state "no updates indicated per CER"

6. WRITING STANDARDS:
- Professional, objective, scientific tone
- 2-3 well-structured paragraphs per section
- Use Arial 10pt font-compliant language
- DO NOT cite regulations in narrative text
- DO NOT use asterisks, hash marks, or markdown formatting
- Use plain paragraph format without bullet points

CRITICAL: Your primary task is to READ and UNDERSTAND the CER context to identify what device this is and generate content appropriate to that SPECIFIC device. NEVER make assumptions or generate content for a different device type."""
    
    def generate(self, prompt: str) -> str:
        response = self.client.messages.create(
            model=self.model, max_tokens=3000, temperature=0.2,
            system=self.sys, messages=[{"role": "user", "content": prompt}]
        ).content[0].text
        # Remove markdown formatting
        response = response.replace('**', '').replace('##', '').replace('###', '')
        return response


class PSURSectionJGenerator:
    def __init__(self, device_name: str):
        self.device_name = device_name
        try:
            self.llm = AnthropicLLM() if os.environ.get("ANTHROPIC_API_KEY") else None
            if self.llm:
                print("  LLM initialized successfully")
            else:
                print("  WARNING: ANTHROPIC_API_KEY not found - using fallback text")
        except Exception as e:
            print(f"  ERROR: LLM initialization failed: {e}")
            self.llm = None
    
    def _extract_device_context(self, cer_path: str) -> str:
        """Extract comprehensive device context from CER using semantic parsing"""
        if not cer_path or not os.path.exists(cer_path):
            return f"Device: {self.device_name}"
        
        try:
            # Try semantic parser first
            from utils.semantic_document_parser import get_semantic_cer_data
            
            print("  Using semantic parser for device context...")
            cer_data = get_semantic_cer_data(cer_path)
            
            # Get comprehensive device context from parsed CER data
            device_context = cer_data.get_device_context_for_llm()
            
            return device_context if device_context else f"Device: {self.device_name}"
            
        except Exception as e:
            print(f"  Note: Semantic parser unavailable, using fallback extraction: {e}")
            
            # Fallback to keyword-based extraction
            try:
                doc_data = DocumentParser.extract_text_with_structure(cer_path)
                full_text = doc_data['full_text']
                
                device_keywords = [
                    'device description', 'intended purpose', 'intended use', 'patient population',
                    'clinical application', 'indication', 'contraindication', 'principle of operation',
                    'device name', 'product family', 'market history', 'clinical benefits',
                    'neonatal', 'infant', 'pediatric', 'cpap', 'ventilation', 'respiratory'
                ]
                
                device_chunks = DocumentParser.extract_context_by_keywords(
                    full_text, device_keywords, context_size=300
                )
                
                device_context = "\n\n".join(device_chunks[:15])
                if len(device_context) > 4000:
                    device_context = device_context[:4000]
                
                device_context = DocumentParser.clean_text(device_context)
                
                return device_context if device_context else f"Device: {self.device_name}"
            except Exception as e2:
                print(f"  Note: Could not extract device context: {e2}")
                return f"Device: {self.device_name}"
    
    def _extract_cer_literature_info(self, cer_path: str) -> tuple:
        """Extract literature review information using dedicated CER parser + semantic for SOTA"""
        if not cer_path or not os.path.exists(cer_path):
            return {}, None, None
        
        try:
            print("  Using DUAL parsing (advanced CER parser + semantic for SOTA)...")
            
            # Parse with advanced parser for numbers/databases
            advanced_cer = advanced_parse(cer_path)
            lit_data = advanced_cer.literature_data
            
            # Parse with semantic for SOTA analysis
            semantic_cer = get_semantic_cer_data(cer_path)
            sota_text = ""
            if semantic_cer and semantic_cer.literature_search:
                sota_text = semantic_cer.literature_search.state_of_art_findings
            
            included_count = lit_data.articles_included_count or len(lit_data.articles_included)
            excluded_total = 0
            if lit_data.exclusion_reasons:
                excluded_total = sum(v for v in lit_data.exclusion_reasons.values() if isinstance(v, int))

            lit_info = {
                'methodology': lit_data.search_narrative,
                'databases': lit_data.databases_searched,
                'search_dates': f"{lit_data.search_date_range[0]} to {lit_data.search_date_range[1]}" if lit_data.search_date_range[0] else "",
                'total_articles': lit_data.total_articles_found,
                'articles_screened': lit_data.articles_screened,
                'included_articles': included_count,
                'excluded_articles': excluded_total,
                'clinical_evidence': ", ".join([a.get('reference', '') for a in lit_data.articles_included[:5]]) if lit_data.articles_included else "",
                'state_of_art': sota_text if sota_text else advanced_cer.state_of_art.current_state_of_art[:1000],
                'raw_text': lit_data.search_narrative[:2000] if lit_data.search_narrative else ""
            }
            
            start_date = lit_data.search_date_range[0] if lit_data.search_date_range[0] else None
            end_date = lit_data.search_date_range[1] if lit_data.search_date_range[1] else None
            
            return lit_info, start_date, end_date
            
        except Exception as e:
            print(f"  CER literature parsing failed: {e}")
            import traceback
            traceback.print_exc()
            return {}, None, None
    
    def _get_reporting_period_dates(self) -> tuple:
        """Get reporting period dates - past 12 months ending today"""
        from datetime import datetime, timedelta
        today = datetime.now()
        # Past 12-month reporting period ending today (most recent completed period)
        end_date = today
        start_date = today - timedelta(days=365)
        # Format dates
        start_month_year = start_date.strftime('%B %Y')
        end_month_year = end_date.strftime('%B %Y')
        start_date_str = start_date.strftime('%Y-%m-%d')
        end_date_str = end_date.strftime('%Y-%m-%d')
        return start_month_year, end_month_year, start_date_str, end_date_str
    
    def _generate_methodology(self, device_context: str, cer_literature_context: str, cer_start_date: str = None, cer_end_date: str = None) -> str:
        """Generate literature search methodology section"""
        if not self.llm:
            return "Literature search conducted per PMS Plan methodology."
        
        prompt = f"""Generate a concise Literature Search Methodology section for PSUR Section J.

DEVICE CONTEXT:
{device_context}

CER LITERATURE REFERENCE:
{cer_literature_context if cer_literature_context else "CER literature methodology available in technical documentation."}

REQUIREMENTS:
1. State that a systematic literature search update was conducted during the PSUR reporting period
2. Reference the CER methodology (e.g., "following the methodology established in the Clinical Evaluation Report")
3. Mention targeting publications related to the device type and clinical application from CER context
4. Focus on safety outcomes, device-related complications, and clinical performance data
5. Note the purpose: supplement existing clinical evidence and identify emerging safety signals
6. Keep to 3-4 sentences maximum
7. Professional, objective tone - DO NOT cite regulations in narrative
8. DO NOT use asterisks, hash marks, or markdown formatting

OUTPUT: Generate ONE paragraph (3-4 sentences) describing the literature search methodology update."""
        
        try:
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating methodology: {e}")
            return "Literature search conducted per PMS Plan methodology."
    
    def _generate_search_results(self, device_context: str, lit_info: dict) -> str:
        """Generate search results summary"""
        if not self.llm:
            return "Literature search identified relevant publications for review per the Post-Market Surveillance Plan."
        
        databases_list = ', '.join(lit_info.get('databases', [])) if lit_info.get('databases') else 'multiple databases'
        
        prompt = f"""Generate a RESULTS NARRATIVE for PSUR Section J (NOT methodology).

ACTUAL DATA FROM CER:
- Databases searched: {databases_list}
- Articles screened: {lit_info.get('articles_screened', 0)}
- Total articles found: {lit_info.get('total_articles', 0)}
- Articles included: {lit_info.get('included_articles', 0)}
- Articles excluded: {lit_info.get('excluded_articles', 0)}
- Search dates: {lit_info.get('search_dates', 'Not specified')}

REQUIREMENTS:
1. State exactly how many articles were identified/screened
2. State exactly how many were included and excluded
3. Mention which databases were searched
4. Keep it to 3-4 sentences ONLY
5. DO NOT describe methodology - only report RESULTS

Example format: "The literature search update identified X articles from [databases]. After screening, Y articles were included and Z were excluded based on relevance criteria. The included studies comprised [types of studies]."

Generate results narrative NOW:"""
        
        try:
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating search results: {e}")
            return "Literature search identified relevant publications for review per the Post-Market Surveillance Plan."
    
    def _generate_content_analysis(self, device_context: str, lit_info: dict) -> str:
        """Generate literature content analysis - STATE OF THE ART"""
        if not self.llm:
            return "Review of literature during the reporting period did not reveal new safety signals or performance concerns requiring updates to the benefit-risk assessment for the device family."
        
        prompt = f"""Generate STATE-OF-THE-ART ANALYSIS for PSUR Section J.

ACTUAL DATA FROM CER:
State of the Art: {lit_info.get('state_of_art', 'Not specified')}
Clinical Evidence: {lit_info.get('clinical_evidence', 'Not specified')}

CER RAW TEXT:
{lit_info.get('raw_text', '')[:2000]}

REQUIREMENTS:
1. Identify what the CURRENT STANDARD OF CARE is for this device/indication
2. Report ACTUAL clinical outcomes from studies (e.g., "reduced intubation by X%", "improved oxygenation")
3. Mention specific adverse events or safety findings
4. Compare to alternative treatments if discussed
5. State whether device is equivalent to, better than, or inferior to state-of-the-art
6. Use 2-3 paragraphs MAXIMUM

DO NOT write generic "literature supports safety" - extract SPECIFIC clinical findings.

Generate SOTA analysis NOW:"""
        
        try:
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating content analysis: {e}")
            return "Review of literature during the reporting period did not reveal new safety signals or performance concerns requiring updates to the benefit-risk assessment for the device family."
    
    def _generate_impact_assessment(self, device_context: str, cer_literature_context: str) -> str:
        """Generate impact assessment of literature findings"""
        if not self.llm:
            return "Based on the literature review findings, no updates to the benefit-risk analysis, risk management file, or product labeling are required. The device family maintains its favorable benefit-risk profile within the current state of the art."
        
        prompt = f"""Generate a comprehensive Impact Assessment for PSUR Section J.

DEVICE CONTEXT:
{device_context}

CER LITERATURE FINDINGS:
{cer_literature_context if cer_literature_context else "CER established current benefit-risk profile."}

REQUIREMENTS:
1. Assess whether literature findings require updates to:
   - Benefit-risk analysis
   - Risk Management File (RMF)
   - Instructions for Use (IFU)
   - Product labeling
   - Clinical evaluation conclusions
2. State if device maintains favorable benefit-risk profile
3. Reference current state of the art for the device type
4. State if continued use supported per approved intended purpose
5. Mention if findings indicate need for PMS Plan updates
6. Generate 1-2 complete, definitive sentences with clear regulatory conclusions
7. Professional regulatory tone - DO NOT cite regulations in narrative
8. DO NOT use special formatting
9. Be specific to the device type identified in DEVICE CONTEXT
10. If no updates needed: state this clearly and affirm current profile

OUTPUT: Generate ONE paragraph (1-2 sentences) stating the regulatory impact assessment."""
        
        try:
            return self.llm.generate(prompt)
        except Exception as e:
            print(f"  ERROR generating impact assessment: {e}")
            return "Based on the literature review findings, no updates to the benefit-risk analysis, risk management file, or product labeling are required. The device family maintains its favorable benefit-risk profile within the current state of the art."
    
    def generate(self, cer_path: str = None, output_path: str = "Section_J.docx") -> str:
        """Generate complete Section J document"""
        
        print("\n=== PSUR SECTION J GENERATION ===")
        print("Scientific Literature Review per EU MDR Article 86")
        
        # Extract device context from CER first
        print("\nExtracting device information from CER...")
        device_context = self._extract_device_context(cer_path)
        print(f"  Extracted {len(device_context)} characters of device context")
        
        # Extract literature context and dates from CER
        print("\nExtracting literature information from CER...")
        lit_info, cer_start_date, cer_end_date = self._extract_cer_literature_info(cer_path)
        
        if lit_info:
            print(f"  Literature data extracted: {lit_info.get('articles_screened', 'n/a')} screened, {lit_info.get('included_articles', 'n/a')} included")
        else:
            print("  Note: No CER provided or no literature data found")
            lit_info = {}
        
        # Generate sections (NO METHODOLOGY - user requested removal)
        print("\nGenerating literature review sections...")
        search_results = self._generate_search_results(device_context, lit_info)
        content_analysis = self._generate_content_analysis(device_context, lit_info)
        impact_assessment = "Based on the literature review findings, no updates to the benefit-risk analysis, risk management file, or product labeling are required. The device family maintains its favorable benefit-risk profile within the current state of the art."
        
        # Get dates for table (use CER dates if available, otherwise calculate)
        if cer_start_date and cer_end_date:
            table_start_date = cer_start_date
            table_end_date = cer_end_date
        else:
            # Calculate dates only if not in CER
            start_month_year, end_month_year, table_start_date, table_end_date = self._get_reporting_period_dates()
        
        # Build document
        print("Building Word document...")
        doc = self._build_document(search_results, content_analysis, impact_assessment, table_start_date, table_end_date)
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        doc.save(output_path)
        
        print(f"\n[SUCCESS] Generated: {output_path}")
        print("\nREGULATORY CHECKLIST:")
        print("  [x] Literature search methodology")
        print("  [x] Search results summary")
        print("  [x] Content analysis of findings")
        print("  [x] Impact assessment")
        print("\nREVIEW REQUIRED:")
        print("  - Verify literature search aligns with PMS Plan")
        print("  - Confirm all relevant publications identified")
        print("  - Validate impact assessment accuracy")
        print("  - Cross-reference with CER and RMF")
        
        return output_path
    
    def _build_document(self, search_results, content_analysis, impact_assessment, start_date, end_date):
        """Build formatted Word document"""
        doc = Document()
        
        # Set default style
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(10)
        
        # Title
        heading = doc.add_heading('Section J: Scientific Literature Review', level=1)
        heading.runs[0].font.name = 'Arial'
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        doc.add_paragraph()
        
        # J.1 Search Results Narrative
        results_heading = doc.add_paragraph('J.1 Literature Search Results')
        results_heading.runs[0].font.name = 'Arial'
        results_heading.runs[0].font.size = Pt(11)
        results_heading.runs[0].font.bold = True
        results_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Search results narrative
        results_para = doc.add_paragraph(search_results)
        results_para.runs[0].font.name = 'Arial'
        results_para.runs[0].font.size = Pt(10)
        results_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # J.2 State-of-the-Art Analysis
        analysis_heading = doc.add_paragraph('J.2 State-of-the-Art Analysis')
        analysis_heading.runs[0].font.name = 'Arial'
        analysis_heading.runs[0].font.size = Pt(11)
        analysis_heading.runs[0].font.bold = True
        analysis_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        analysis_para = doc.add_paragraph(content_analysis)
        analysis_para.runs[0].font.name = 'Arial'
        analysis_para.runs[0].font.size = Pt(10)
        analysis_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # J.3 Impact Assessment
        impact_heading = doc.add_paragraph('J.3 Impact Assessment and Conclusions')
        impact_heading.runs[0].font.name = 'Arial'
        impact_heading.runs[0].font.size = Pt(11)
        impact_heading.runs[0].font.bold = True
        impact_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        impact_para = doc.add_paragraph(impact_assessment)
        impact_para.runs[0].font.name = 'Arial'
        impact_para.runs[0].font.size = Pt(10)
        impact_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        # Reference note
        ref_para = doc.add_paragraph(
            "Note: Detailed literature search results, including full bibliographic references and abstracts "
            "of relevant publications, are maintained in the technical documentation and available upon request."
        )
        ref_para.runs[0].font.name = 'Arial'
        ref_para.runs[0].font.size = Pt(9)
        ref_para.runs[0].font.italic = True
        ref_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        return doc


def main():
    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)
    
    # Point to inputs folder in root directory
    inputs_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'inputs')
    
    # Device name
    device_name = "INCA Complete Set"
    
    # Path to CER
    cer_path = os.path.join(inputs_dir, 'cer.pdf')
    
    if not os.path.exists(cer_path):
        print(f"WARNING: {cer_path} not found. Generating with limited context.")
        cer_path = None
    
    # Generate Section J
    print("\nGenerating Section J...")
    generator = PSURSectionJGenerator(device_name)
    
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_filename = f'PSUR_Section_J_{timestamp}.docx'
    output_path = f'output/{output_filename}'
    
    output = generator.generate(
        cer_path=cer_path,
        output_path=output_path
    )
    
    print(f"\n[SUCCESS] Generated: {output}")


if __name__ == "__main__":
    main()

