"""
PSUR Section C Generator - EU MDR Compliant
Combines INCA data processing with improved table layout
"""

import os
import sys
# Add parent directory to Python path for utils imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pandas as pd
from dataclasses import dataclass
from typing import Dict, Optional
import anthropic
from dotenv import load_dotenv

# Load environment variables from root .env file
root_dir = os.path.dirname(os.path.dirname(__file__))
env_path = os.path.join(root_dir, '.env')
load_dotenv(env_path)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING

try:
    import matplotlib.pyplot as plt
except ImportError:
    os.system("pip install matplotlib --break-system-packages -q")
    import matplotlib.pyplot as plt


# Country to PSUR template region mapping (EU MDR / MDCG 2022-21)
COUNTRY_TO_REGION = {
    # EEA+TR+XI (European Economic Area + Turkey + Northern Ireland)
    'Bulgaria': 'EEA+TR+XI', 'Greece': 'EEA+TR+XI', 'Ireland': 'EEA+TR+XI',
    'Italy': 'EEA+TR+XI', 'Poland': 'EEA+TR+XI', 'Turkey': 'EEA+TR+XI',
    'Austria': 'EEA+TR+XI', 'Belgium': 'EEA+TR+XI', 'Croatia': 'EEA+TR+XI',
    'Cyprus': 'EEA+TR+XI', 'Czech Republic': 'EEA+TR+XI', 'Denmark': 'EEA+TR+XI',
    'Estonia': 'EEA+TR+XI', 'Finland': 'EEA+TR+XI', 'France': 'EEA+TR+XI',
    'Germany': 'EEA+TR+XI', 'Hungary': 'EEA+TR+XI', 'Iceland': 'EEA+TR+XI',
    'Latvia': 'EEA+TR+XI', 'Liechtenstein': 'EEA+TR+XI', 'Lithuania': 'EEA+TR+XI',
    'Luxembourg': 'EEA+TR+XI', 'Malta': 'EEA+TR+XI', 'Netherlands': 'EEA+TR+XI',
    'Norway': 'EEA+TR+XI', 'Portugal': 'EEA+TR+XI', 'Romania': 'EEA+TR+XI',
    'Slovakia': 'EEA+TR+XI', 'Slovenia': 'EEA+TR+XI', 'Spain': 'EEA+TR+XI',
    'Sweden': 'EEA+TR+XI',
    # UK (Great Britain - separate from EEA post-Brexit)
    'United Kingdom': 'UK', 'England': 'UK', 'Scotland': 'UK', 'Wales': 'UK',
    # Canada
    'Canada': 'Canada',
    # USA
    'United States of America': 'USA', 'United States': 'USA', 'USA': 'USA',
    # Brazil
    'Brazil': 'Brazil',
    # China
    'China': 'China',
    # Japan
    'Japan': 'Japan',
    # Australia
    'Australia': 'Australia', 'New Zealand': 'Australia',
    # Other countries
    'Mexico': 'Other', 'Peru': 'Other', 'Puerto Rico': 'Other', 'Nicaragua': 'Other',
    'Hong Kong': 'Other', 'Malaysia': 'Other', 'Taiwan': 'Other', 'India': 'Other',
    'Iran': 'Other', 'Qatar': 'Other', 'Kuwait': 'Other', 'Saudi Arabia': 'Other',
    'United Arab Emirates': 'Other', 'South Africa': 'Other', 'Argentina': 'Other',
    'Chile': 'Other', 'Colombia': 'Other', 'Singapore': 'Other', 'South Korea': 'Other',
    'Thailand': 'Other', 'Vietnam': 'Other', 'Philippines': 'Other', 'Indonesia': 'Other'
}


@dataclass
class DeviceInfo:
    name: str
    basic_udi_di: str
    device_class: str
    implantable: bool
    single_use: bool


class SalesCriteria:
    PLACED_ON_MARKET = "placed_on_market"
    UNITS_DISTRIBUTED = "units_distributed"
    EPISODES_OF_USE = "episodes_of_use"
    ACTIVE_INSTALLED_BASE = "active_installed_base"
    UNITS_IMPLANTED = "units_implanted"
    OTHER = "other"


class AnthropicLLM:
    def __init__(self):
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY environment variable required")
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"
        self.sys = """You are an expert EU MDR regulatory affairs specialist with deep expertise in Post-Market Surveillance and PSUR preparation per EU MDR 2017/745 Article 86(1)(c) and MDCG 2022-21 guidance.

CRITICAL REQUIREMENTS:

1. SALES DATA TABLE ACCURACY (MDCG 2022-21 Annex II):
- Verify all numerical data matches source sales records exactly
- Ensure regional classifications are correct (EEA+TR+XI, Australia, Brazil, Canada, China, Japan, UK, USA, Other)
- Calculate Global Total as sum of all regional sales
- Calculate percentage of global sales for each region: (Regional Sales / Global Total) × 100
- Verify period labels match format: [mmm-yyyy to mmm-yyyy] (e.g., Jan-2023 to Dec-2023)
- Ensure all cells contain data or are explicitly marked as empty (no null values)
- Cross-check totals: sum of all regional periods must equal grand total
- Validate that table structure matches FormQAR-054 exactly

2. SALES METHODOLOGY (MDCG 2022-21 Section 5.3):
- Clearly state which sales criteria is used (devices placed on market, units distributed, episodes of use, active installed base, units implanted)
- Provide explicit justification for why this methodology is appropriate for the device type
- Reference data collection methods and verification against QMS records
- State any limitations or assumptions
- Ensure consistency with device characteristics (single-use vs reusable, implantable vs non-implantable)
- Confirm methodology is consistent with previous PSURs

3. SALES DATA ANALYSIS (MDCG 2022-21 Annex III):
- Calculate and report total units across all periods and regions with exact figures
- Identify growth trends: calculate percentage change = ((Current Period - First Period) / First Period) × 100
- Analyze period-over-period changes and identify significant fluctuations (>20% change)
- Compare current period to historical average
- Regional analysis: identify top 3 markets by volume and percentage of total
- Assess if trends align with market expectations and device lifecycle stage
- State whether sales patterns are stable, increasing, or decreasing with supporting data
- Flag any anomalies requiring investigation (sudden drops, unexpected spikes)
- Compare year-over-year if multi-year data available

4. POPULATION EXPOSURE ESTIMATION (MDR Article 86(1)(c)):
- Calculate estimated patient exposure from sales data with explicit formula
- For single-use devices: Patient Exposure = Total Units Sold × 1.0
- For reusable devices: Patient Exposure = Total Units Sold × Average Uses Per Device
- For implantable devices: Patient Exposure = Total Units Sold ÷ Devices Per Patient
- Clearly state all assumptions (e.g., "Assuming average 50 uses per device over 5-year lifecycle")
- Justify patient exposure estimates with device-specific usage patterns from post-market surveillance
- Account for installed base vs. new sales if relevant
- Reference IFU specifications for usage frequency

5. DEMOGRAPHICS AND CHARACTERISTICS:
- Age range of patient population (specific ranges, not vague)
- Gender distribution (percentages, not general statements)
- Primary indications for use (from IFU)
- Relevant comorbidities with prevalence estimates
- Geographic distribution (align with sales data regions)
- Usage settings (hospital, outpatient, home use) with percentages if known
- Special populations (pediatric, elderly >65, pregnant, immunocompromised) with specific mention if device is used in these groups

6. DATA QUALITY VALIDATION:
- Verify all calculations are correct (totals, percentages, growth rates)
- Check for data consistency across tables and narrative
- Ensure units of measurement are clearly stated
- Confirm all figures have appropriate significant figures (whole numbers for units, 1 decimal for percentages)
- Cross-reference sales data with complaint rates to ensure denominators are consistent
- Flag any discrepancies between data sources

7. WRITING STANDARDS:
- Professional, objective, scientific tone
- No marketing or promotional language  
- Evidence-based statements with specific numerical support
- Precise numerical data: "7,280 units" not "approximately 7,000 units"
- Clear, defensible justifications citing MDCG requirements
- Regulatory terminology per MDR/MDCG guidance
- 2-3 well-structured paragraphs per section
- Each paragraph structure: opening statement → supporting analysis with specific data → conclusion
- Use Arial 10pt font-compliant language

8. ANALYTICAL DEPTH:
- Don't just report numbers - interpret what they mean for benefit-risk profile
- Contextualize data within device lifecycle stage (launch, growth, maturity, decline)
- Compare to previous PSUR periods with specific percentage changes
- Identify patterns (seasonality, regional variations, trend changes)
- Assess deviations from PMS Plan projections
- Determine if findings impact PMS strategy or risk management
- Note any areas requiring further investigation or corrective action

9. REGULATORY COMPLIANCE CHECKLIST:
- All statements must be verifiable against source documents
- Every claim must be supported by data or reference
- Assumptions must be clearly stated and justified
- Limitations must be acknowledged
- Consistency with Clinical Evaluation Report demographics
- Alignment with PMS Plan objectives
- Cross-reference with complaint and incident data for rate calculations

OUTPUT FORMAT:
- Use clear topic sentences with numerical anchors
- Support every claim with specific calculated data points
- Provide formulas for calculations: "Growth rate = (1,650 - 1,250) / 1,250 × 100 = 32%"
- Include comparison statements: "Current period sales of 1,650 units represent a 32% increase from initial period"
- End each section with conclusive summary statement
- DO NOT cite regulations, standards, or MDCG guidance in the narrative text
- DO NOT use asterisks, hash marks, or markdown formatting
- Make only data-driven statements, avoid broad assumptions unless rooted in actual data or regulatory requirements
- Use professional paragraph format without bullet points or special characters

CRITICAL: This content will be reviewed by Notified Bodies and Competent Authorities. Every numerical value, calculation, and statement must be accurate, traceable to source data, and defensible under regulatory scrutiny."""
    
    def generate(self, prompt: str) -> str:
        return self.client.messages.create(
            model=self.model, max_tokens=2500, temperature=0.2,
            system=self.sys, messages=[{"role": "user", "content": prompt}]
        ).content[0].text


class PSURSectionCGenerator:
    def __init__(self, device: DeviceInfo):
        self.device = device
        self.llm = AnthropicLLM() if os.environ.get("ANTHROPIC_API_KEY") else None
        self.validation_log = []
    
    def _validate_data(self, df: pd.DataFrame) -> bool:
        """Validate sales data per MDCG 2022-21 data quality requirements"""
        validation_passed = True
        
        if df['units_sold'].isnull().any():
            self.validation_log.append("ERROR: Null values found in units_sold column")
            validation_passed = False
        
        if (df['units_sold'] < 0).any():
            negative_count = (df['units_sold'] < 0).sum()
            negative_total = df[df['units_sold'] < 0]['units_sold'].sum()
            self.validation_log.append(f"WARNING: {negative_count} negative values found (total: {negative_total:,}), likely returns/adjustments")
        
        if 'region' in df.columns:
            valid_regions = ['EEA+TR+XI', 'Australia', 'Brazil', 'Canada', 'China', 'Japan', 'UK', 'USA', 'Other', 'Global']
            invalid_regions = set(df['region'].unique()) - set(valid_regions)
            if invalid_regions:
                self.validation_log.append(f"WARNING: Non-standard regions found: {invalid_regions}")
        
        if 'period' in df.columns and df['period'].isnull().any():
            self.validation_log.append("ERROR: Null values found in period column")
            validation_passed = False
        
        period_totals = df.groupby('period')['units_sold'].sum()
        for period in period_totals.index:
            if period_totals[period] == 0:
                self.validation_log.append(f"WARNING: Zero sales in period {period}")
        
        if len(df) < 4:
            self.validation_log.append(f"WARNING: Only {len(df)} data points. MDCG 2022-21 recommends 4 periods minimum for annual PSUR")
        
        if validation_passed:
            self.validation_log.append("VALIDATION PASSED: All data quality checks satisfied")
        
        return validation_passed
    
    def generate(self, sales_data_path: str, sales_criteria: str, first_market_date: str = None,
                 market_history: str = None, usage_factor: float = 1.0, demographics: Dict = None,
                 output_path: str = "Section_C.docx", cer_path: str = None, previous_psur_path: str = None) -> str:
        
        print("\n=== PSUR SECTION C GENERATION ===")
        print("Per EU MDR 2017/745 Article 86(1)(c) and MDCG 2022-21")
        
        df = self._load_data(sales_data_path)
        print(f"\nLoaded {len(df)} records from {sales_data_path}")
        
        print("\nValidating data quality...")
        validation_passed = self._validate_data(df)
        for log_entry in self.validation_log:
            print(f"  {log_entry}")
        
        if not validation_passed:
            print("\n[WARNING] Data validation detected issues. Review recommended before submission.")
        
        print("\nGenerating regulatory content...")
        methodology = self._generate_methodology(sales_criteria)
        
        # Generate market history from documents or use provided text
        if market_history:
            market_history_text = market_history
        else:
            print("Analyzing CER and previous PSUR for market history...")
            market_history_text = self._generate_market_history(cer_path, previous_psur_path)
        
        analysis = self._generate_analysis(df)
        
        # Generate population with device context from CER
        print("Extracting device and patient information from CER...")
        population, characteristics = self._generate_population(df, usage_factor, demographics, cer_path)
        
        print("Creating sales trend visualization...")
        chart_path = self._create_chart(df, output_dir=os.path.dirname(output_path) or '.')
        
        print("Building Word document with formatted tables...")
        doc = self._build_document(df, methodology, market_history_text, analysis, population, characteristics, chart_path, sales_criteria, demographics)
        doc.save(output_path)
        
        print(f"\n[SUCCESS] Generated: {output_path}")
        print("\nREGULATORY CHECKLIST:")
        print("  [x] Sales data table with regional breakdown")
        print("  [x] Methodology justification per MDCG 2022-21 Section 5.3")
        print("  [x] Comprehensive sales analysis with calculations")
        print("  [x] Population exposure estimation with explicit formulas")
        print("  [x] Trend visualization")
        print("\nREVIEW REQUIRED:")
        print("  - Verify all numerical data against source QMS records")
        print("  - Confirm regional classifications are accurate")
        print("  - Validate patient exposure assumptions with IFU")
        print("  - Cross-reference with complaint rates for consistency")
        
        return output_path
    
    def _load_data(self, path: str) -> pd.DataFrame:
        ext = path.split('.')[-1].lower()
        if ext == 'csv':
            df = pd.read_csv(path)
        elif ext in ['xlsx', 'xls']:
            df = pd.read_excel(path)
        else:
            df = pd.read_json(path)
        
        df.columns = df.columns.str.lower().str.strip().str.replace(' ', '_')
        if 'units' in df.columns:
            df['units_sold'] = df['units']
        if 'year' in df.columns and 'period' not in df.columns:
            df['period'] = df['year']
        return df
    
    def _generate_methodology(self, criteria: str) -> str:
        """Generate brief methodology justification (checkboxes answer most of it)"""
        if not self.llm:
            return "Sales data extracted from QMS distribution records."
        
        criteria_desc = {
            SalesCriteria.PLACED_ON_MARKET: "devices placed on the market",
            SalesCriteria.UNITS_DISTRIBUTED: "units distributed within each time period",
            SalesCriteria.EPISODES_OF_USE: "episodes of use for reusable devices",
            SalesCriteria.ACTIVE_INSTALLED_BASE: "active installed base",
            SalesCriteria.UNITS_IMPLANTED: "units implanted"
        }.get(criteria, criteria)
        
        device_type = "single-use" if self.device.single_use else "reusable"
        
        prompt = f"""Generate a brief 2-3 sentence methodology justification for PSUR sales data.

Device Type: {device_type}
Sales Criteria: {criteria_desc}

Requirements:
1. Briefly justify why {criteria_desc} is appropriate for this {device_type} device
2. Reference data source (e.g., "Sales data extracted from QMS distribution records")
3. Keep it SHORT - only 2-3 sentences maximum
4. DO NOT cite regulations or guidance documents
5. DO NOT use asterisks or markdown formatting

Generate 2-3 sentences only."""
        
        return self.llm.generate(prompt)
    
    def _generate_market_history(self, cer_path: str = None, previous_psur_path: str = None) -> str:
        """Generate market history by analyzing CER and previous PSUR documents"""
        if not self.llm:
            return "Market history information not available."
        
        context = []
        
        # Helper function to extract relevant sections from document
        def extract_relevant_sections(doc_text: str, doc_type: str) -> str:
            """Extract sections likely to contain market history information"""
            relevant_keywords = [
                'market', 'CE mark', 'approval', 'first marketed', 'distribution', 
                'commercialization', 'launch', 'introduction', 'regulatory', 
                'clearance', 'certification', 'conformity assessment', 'notified body',
                'placed on market', 'market entry', 'available since', 'marketed since'
            ]
            
            # Split into paragraphs
            paragraphs = [p.strip() for p in doc_text.split('\n') if p.strip()]
            
            # Find paragraphs containing relevant keywords
            relevant_paras = []
            for para in paragraphs:
                para_lower = para.lower()
                if any(keyword in para_lower for keyword in relevant_keywords):
                    relevant_paras.append(para)
            
            # Return relevant sections (up to 3000 chars to stay within reasonable limits)
            relevant_text = "\n".join(relevant_paras)
            if len(relevant_text) > 3000:
                relevant_text = relevant_text[:3000] + "..."
            
            return relevant_text if relevant_text else ""
        
        # Try to read CER if provided
        if cer_path and os.path.exists(cer_path):
            try:
                # Use flexible document parser (handles PDF and Word)
                from utils.document_parser import DocumentParser
                doc_data = DocumentParser.extract_text_with_structure(cer_path)
                cer_text = doc_data['full_text']
                
                relevant_cer = extract_relevant_sections(cer_text, "CER")
                if relevant_cer:
                    context.append(f"CER - Market Information:\n{relevant_cer}")
            except Exception as e:
                print(f"  Note: Could not read CER document: {e}")
        
        # Try to read previous PSUR if provided
        if previous_psur_path and os.path.exists(previous_psur_path):
            try:
                from docx import Document as DocxDocument
                psur_doc = DocxDocument(previous_psur_path)
                psur_text = "\n".join([para.text for para in psur_doc.paragraphs if para.text.strip()])
                relevant_psur = extract_relevant_sections(psur_text, "PSUR")
                if relevant_psur:
                    context.append(f"Previous PSUR - Market Information:\n{relevant_psur}")
            except Exception as e:
                print(f"  Note: Could not read previous PSUR document: {e}")
        
        if not context:
            return f"The {self.device.name} has been marketed globally for distribution to healthcare facilities."
        
        prompt = f"""Extract and summarize the market history for {self.device.name} from the provided documents.

{chr(10).join(context)}

Requirements:
1. Identify when the device was first marketed (look for dates, CE mark, approval dates)
2. Identify which markets/regions it has been distributed to
3. Note any significant milestones or changes in distribution
4. Keep to 2-3 sentences summarizing the key market history facts
5. DO NOT cite regulations or guidance documents
6. DO NOT use asterisks or markdown formatting
7. Only include factual information found in the documents

Generate a concise 2-3 sentence market history summary."""
        
        return self.llm.generate(prompt)
    
    def _generate_analysis(self, df: pd.DataFrame) -> str:
        total = int(df['units_sold'].sum())
        if not self.llm:
            return f"During the reporting period, {total:,} units were distributed."
        
        periods = sorted(df['period'].unique())
        period_totals = []
        for period in periods:
            period_sum = int(df[df['period'] == period]['units_sold'].sum())
            period_totals.append(period_sum)
        
        first_period_sales = period_totals[0]
        last_period_sales = period_totals[-1]
        growth_rate = ((last_period_sales - first_period_sales) / first_period_sales * 100) if first_period_sales > 0 and len(period_totals) > 1 else 0
        
        historical_avg = sum(period_totals[:-1]) / len(period_totals[:-1]) if len(period_totals) > 1 else period_totals[0]
        current_vs_avg = ((last_period_sales - historical_avg) / historical_avg * 100) if historical_avg > 0 and len(period_totals) > 1 else 0
        
        period_changes = []
        for i in range(1, len(period_totals)):
            if period_totals[i-1] > 0:
                change_pct = ((period_totals[i] - period_totals[i-1]) / period_totals[i-1] * 100)
                period_changes.append(change_pct)
        
        significant_fluctuations = [chg for chg in period_changes if abs(chg) > 20]
        
        regional_data = ""
        if 'region' in df.columns:
            region_totals = df.groupby('region')['units_sold'].sum().sort_values(ascending=False)
            top_regions = region_totals.head(3)
            regional_data = "\n".join([f"- {region}: {int(units):,} units ({units/total*100:.1f}% of global total)" 
                                      for region, units in top_regions.items()])
        
        prompt = f"""Generate comprehensive sales data analysis.

EXACT SALES DATA:
- Total Units Across All Periods: {total:,}
- Number of Reporting Periods: {len(periods)}
- First Period Sales: {first_period_sales:,} units
- Most Recent Period Sales: {last_period_sales:,} units
- Overall Growth Rate: {growth_rate:.1f}% (calculation: ({last_period_sales:,} - {first_period_sales:,}) / {first_period_sales:,} × 100)
- Historical Average (preceding periods): {int(historical_avg):,} units
- Current Period vs Historical Average: {current_vs_avg:+.1f}%
- Period-over-Period Changes: {[f"{chg:+.1f}%" for chg in period_changes]}
- Significant Fluctuations (>20%): {len(significant_fluctuations)} instances
- Period Breakdown: {", ".join([f"{period_totals[i]:,} units" for i in range(len(period_totals))])}

{f"TOP REGIONAL MARKETS:{chr(10)}{regional_data}" if regional_data else ""}

Requirements:
1. Report total units with exact figures using comma separators
2. Calculate growth trends with explicit percentage formulas showing the calculation
3. Compare current period to historical average with specific data
4. Assess if trends are stable, increasing, or decreasing with supporting numerical evidence
5. If regional data provided, identify top 3 markets by volume and percentage
6. DO NOT cite regulations, standards, or guidance documents
7. DO NOT use asterisks, hash marks, or markdown formatting
8. DO NOT speculate on causes (supply chain, market access, competitive pressures, etc.)
9. DO NOT provide interpretations about lifecycle stage implications
10. DO NOT discuss benefit-risk assessment implications
11. DO NOT mention volatility, anomalies, or investigate causes
12. DESCRIBE the data only - report what happened, not why it happened

OUTPUT: 1 paragraph only
- Overall sales summary with total units, growth trend calculation, historical average comparison
- If regional data: include top markets by volume and percentage
- Pure description of numerical facts without interpretation

Use precise numerical data. Professional regulatory tone. Descriptive only - no speculation or interpretation."""
        
        return self.llm.generate(prompt)
    
    def _extract_device_info(self, cer_path: str = None) -> str:
        """Extract comprehensive device information from CER using semantic parsing"""
        if not cer_path or not os.path.exists(cer_path):
            return ""
        
        try:
            # Try semantic parser first
            from utils.semantic_document_parser import get_semantic_cer_data
            
            print("  Using semantic parser for device info...")
            cer_data = get_semantic_cer_data(cer_path)
            
            # Get comprehensive device context
            device_info = cer_data.get_device_context_for_llm()
            
            return device_info
            
        except Exception as e:
            print(f"  Note: Semantic parser unavailable, using fallback: {e}")
            
            # Fallback to keyword extraction
            try:
                # Use flexible document parser (handles PDF and Word)
                from utils.document_parser import DocumentParser
                doc_data = DocumentParser.extract_text_with_structure(cer_path)
                cer_text = doc_data['full_text']
                
                relevant_keywords = [
                    'intended use', 'indication', 'patient population', 'target population',
                    'device description', 'design', 'single use', 'sterile', 'disposable',
                    'procedure', 'clinical application', 'user', 'patient', 'age', 'gender',
                    'contraindication', 'population', 'obstetric', 'gynecological', 'surgical',
                    'healthcare', 'medical', 'treatment', 'therapy'
                ]
                
                paragraphs = [p.strip() for p in cer_text.split('\n') if p.strip()]
                relevant_paras = []
                
                for para in paragraphs:
                    para_lower = para.lower()
                    if any(keyword in para_lower for keyword in relevant_keywords):
                        relevant_paras.append(para)
                
                device_info = "\n".join(relevant_paras[:20])
                if len(device_info) > 4000:
                    device_info = device_info[:4000] + "..."
                
                return device_info
            except Exception as e2:
                print(f"  Note: Could not extract device info: {e2}")
                return ""
    
    def _generate_population(self, df: pd.DataFrame, factor: float, demo: Dict, cer_path: str = None) -> str:
        total = int(df['units_sold'].sum())
        patients = int(total * factor)
        if not self.llm:
            return f"Estimated {patients:,} patients exposed."
        
        # Extract device information from CER
        device_context = self._extract_device_info(cer_path)
        
        formula_explanation = ""
        if self.device.single_use:
            formula_explanation = f"Patient Exposure = Total Units Sold × 1.0 = {total:,} × 1.0 = {patients:,} patients"
            assumption = "Each device is used once per patient as per single-use designation in the IFU"
        elif self.device.implantable:
            devices_per_patient = 1.0 / factor if factor != 0 else 1.0
            formula_explanation = f"Patient Exposure = Total Units Sold ÷ Devices Per Patient = {total:,} ÷ {devices_per_patient:.1f} = {patients:,} patients"
            assumption = f"Assuming average {devices_per_patient:.1f} device(s) per patient based on typical implantation patterns"
        else:
            formula_explanation = f"Patient Exposure = Total Units Sold × Average Uses Per Device = {total:,} × {factor:.1f} = {patients:,} patients"
            assumption = f"Assuming average {int(1/factor) if factor < 1 else int(factor)} uses per device over expected lifecycle based on IFU specifications and post-market surveillance data"
        
        demo_str = "\n".join([f"- {k}: {v}" for k, v in demo.items()]) if demo else ""
        
        regional_distribution = ""
        if 'region' in df.columns:
            region_totals = df.groupby('region')['units_sold'].sum()
            regional_distribution = "\n".join([f"- {region}: {int(units):,} units" 
                                              for region, units in region_totals.items()])
        
        device_info_context = f"""
DEVICE CONTEXT FROM CER/IFU:
{device_context}

""" if device_context else ""
        
        prompt = f"""Generate population exposure estimation.

{device_info_context}

DEVICE INFORMATION:
- Device Name: {self.device.name}
- Class: {self.device.device_class}
- Single-Use: {self.device.single_use}
- Implantable: {self.device.implantable}

CALCULATION:
- Total Units Sold: {total:,}
- Usage Factor: {factor}
- Estimated Patient Exposure: {patients:,} patients
- Formula: {formula_explanation}
- Assumption: {assumption}

DEMOGRAPHICS PROVIDED:
{demo_str}

{f"GEOGRAPHIC DISTRIBUTION:{chr(10)}{regional_distribution}" if regional_distribution else ""}

CRITICAL REQUIREMENTS:
1. READ AND UNDERSTAND the device context provided above from the CER/IFU
2. Base your patient population description on the ACTUAL device type, intended use, and indications from the CER
3. State the estimated patient exposure with exact calculated figure: {patients:,} patients
4. Show the explicit calculation formula used: {formula_explanation}
5. Clearly state all assumptions: {assumption}
6. Describe the SPECIFIC patient population who would use THIS device based on its intended use
7. Account for device type (single-use vs reusable vs implantable) and clinical application
8. Integrate demographic information if provided, but ensure it aligns with device's intended use
9. Note any limitations in the estimate
10. DO NOT cite regulations, standards, or guidance documents
11. DO NOT use asterisks, hash marks, or markdown formatting
12. DO NOT make generic statements - be specific to THIS device and its clinical application
13. Ensure accuracy - if this is an obstetric/gynecological device, the population should reflect that
14. If this is a surgical device, describe the surgical patient population appropriately

OUTPUT FORMAT - TWO SEPARATE SECTIONS:

Section 1 - Patient Exposure (for "Estimated Size of the Patient Population Exposed to the Device"):
One paragraph with patient exposure calculation, explicit formula, and assumptions based on device's actual use.

---SECTION_BREAK---

Section 2 - Patient Characteristics (for "Characteristics of Patient Population Exposed to the Device"):
One paragraph describing the specific demographics and characteristics of THIS device's patient population based on CER context. Include:
- Age range (specific, not vague)
- Gender distribution (if relevant to device)
- Primary indications for use
- Clinical settings (hospital, outpatient, etc.)
- Any special populations (if applicable to THIS device)

Separate the two sections with ---SECTION_BREAK--- exactly as shown above.
Use precise numerical data. Professional regulatory tone. Plain paragraph format only. ACCURACY IS CRITICAL."""
        
        response = self.llm.generate(prompt)
        
        # Split the response into population and characteristics
        if "---SECTION_BREAK---" in response:
            parts = response.split("---SECTION_BREAK---")
            population_text = parts[0].strip()
            characteristics_text = parts[1].strip() if len(parts) > 1 else ""
        else:
            # Fallback if AI doesn't follow format
            population_text = response
            characteristics_text = "Patient characteristics are described in the population exposure section above."
        
        return population_text, characteristics_text
    
    def _create_chart(self, df: pd.DataFrame, output_dir: str = '.') -> str:
        """Create simple annual global sales bar chart"""
        
        # Aggregate to annual global totals for clean visualization
        df['year'] = pd.to_datetime(df['period']).dt.year
        annual_totals = df.groupby('year')['units_sold'].sum().reset_index()
        annual_totals = annual_totals.sort_values('year')
        
        # Create clean, professional bar chart
        plt.figure(figsize=(10, 6))
        bars = plt.bar(annual_totals['year'], annual_totals['units_sold'], 
                      color='#2E75B6', width=0.6, edgecolor='white', linewidth=0.5)
        
        # Add value labels on top of each bar
        for idx, row in annual_totals.iterrows():
            plt.text(row['year'], row['units_sold'], f"{int(row['units_sold']):,}", 
                    ha='center', va='bottom', fontsize=10, fontfamily='Arial', fontweight='bold')
        
        plt.xlabel('Year', fontsize=12, fontfamily='Arial', fontweight='bold')
        plt.ylabel('Global Units Sold', fontsize=12, fontfamily='Arial', fontweight='bold')
        plt.title('Global Sales Trend - Annual Total Units', fontsize=14, 
                 fontweight='bold', fontfamily='Arial', pad=20)
        plt.grid(True, axis='y', alpha=0.3, linestyle='--')
        plt.xticks(annual_totals['year'], fontsize=11, fontfamily='Arial')
        plt.yticks(fontsize=11, fontfamily='Arial')
        
        # Format y-axis with commas
        ax = plt.gca()
        ax.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'{int(x):,}'))
        
        # Add some padding at the top for labels
        y_max = annual_totals['units_sold'].max()
        plt.ylim(0, y_max * 1.15)
        
        plt.tight_layout()
        
        chart_path = os.path.join(output_dir, 'sales_trend.png')
        plt.savefig(chart_path, dpi=300, bbox_inches='tight')
        plt.close()
        return chart_path
    
    def _build_document(self, df, methodology_text, market_history_text, analysis_text, 
                       population_text, characteristics_text, chart_path, criteria, demographics=None):
        doc = Document()
        
        # Set default style
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Pt(10)
        
        # Title
        heading = doc.add_heading('Section C: Volume Of Sales and Population Exposure', level=1)
        heading.runs[0].font.name = 'Arial'
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.bold = True
        heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Sales Methodology
        method_heading = doc.add_paragraph('Sales Methodology')
        method_heading.runs[0].font.name = 'Arial'
        method_heading.runs[0].font.size = Pt(10)
        method_heading.runs[0].font.bold = True
        method_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        # Criteria checkboxes
        criteria_heading = doc.add_paragraph('Criteria Used for Sales Data:')
        criteria_heading.runs[0].font.name = 'Arial'
        criteria_heading.runs[0].font.size = Pt(10)
        criteria_heading.runs[0].font.bold = True
        
        criteria_options = [
            ('Devices placed on the market or put into service', SalesCriteria.PLACED_ON_MARKET),
            ('Units distributed from the date of declaration of conformity or EC/EU mark approval to the end date of each time period', 'conformity'),
            ('Units distributed within each time period', SalesCriteria.UNITS_DISTRIBUTED),
            ('Episodes of use (for reusable devices)', SalesCriteria.EPISODES_OF_USE),
            ('Active installed base', SalesCriteria.ACTIVE_INSTALLED_BASE),
            ('Units implanted', SalesCriteria.UNITS_IMPLANTED),
            ('Other: [Specify with rationale]', SalesCriteria.OTHER)
        ]
        
        for label, c in criteria_options:
            checkbox = "☑" if c == criteria else "☐"
            p = doc.add_paragraph(f'{checkbox} {label}')
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(10)
        
        # Market History
        method_para = doc.add_paragraph(methodology_text)
        method_para.runs[0].font.name = 'Arial'
        method_para.runs[0].font.size = Pt(10)
        method_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        market_heading = doc.add_paragraph('Market History:')
        market_heading.runs[0].font.name = 'Arial'
        market_heading.runs[0].font.size = Pt(10)
        market_heading.runs[0].font.bold = True
        
        market_para = doc.add_paragraph(market_history_text)
        market_para.runs[0].font.name = 'Arial'
        market_para.runs[0].font.size = Pt(10)
        market_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        # Table 1
        table_heading = doc.add_paragraph('Table 1. Annual Number of Devices Sold by Region (MDCG 2022-21 FormQAR-054)')
        table_heading.runs[0].font.name = 'Arial'
        table_heading.runs[0].font.size = Pt(10)
        table_heading.runs[0].font.bold = True
        
        # Create comprehensive table per MDCG 2022-21 Annex II
        regions = ['EEA+TR+XI', 'Australia', 'Brazil', 'Canada', 'China', 'Japan', 'UK', 'USA', 'Other', 'Global Total']
        
        # Aggregate to ANNUAL 12-month periods (MDCG requirement)
        df['year'] = pd.to_datetime(df['period']).dt.year
        available_years = sorted(df['year'].unique())
        periods_years = available_years[-4:] if len(available_years) >= 4 else available_years
        
        table = doc.add_table(rows=len(regions)+3, cols=len(periods_years)+2)
        table.style = 'Light Grid Accent 1'
        
        # Headers row 1
        table.cell(0, 0).text = 'Region'
        for i in range(len(periods_years)-1):
            table.cell(0, i+1).text = 'Preceding 12-Month Periods'
        table.cell(0, len(periods_years)).text = 'Current Data Collection Period'
        table.cell(0, len(periods_years)+1).text = 'Current Data Collection Period'
        
        # Headers row 2
        for i in range(len(periods_years)):
            table.cell(1, i+1).text = 'Number of Devices Sold During Each 12-Month Period'
        table.cell(1, len(periods_years)+1).text = '12-Month Percent of\nGlobal Sales'
        
        # Headers row 3 - Annual period labels [Jan-YYYY to Dec-YYYY]
        for i, year in enumerate(periods_years):
            table.cell(2, i+1).text = f'[Jan-{year} to Dec-{year}]'
        table.cell(2, len(periods_years)+1).text = '12-Month Percent of\nGlobal Sales'
        
        # Calculate regional data with validation - ANNUAL AGGREGATION
        regional_data = {}
        total_by_year = {}
        has_regional_breakdown = 'region' in df.columns and len(df['region'].unique()) > 1
        
        for i, year in enumerate(periods_years):
            year_df = df[df['year'] == year]
            
            if has_regional_breakdown:
                for region in regions[:-1]:
                    region_df = year_df[year_df['region'] == region]
                    regional_data[(region, i)] = int(region_df['units_sold'].sum())
            
            total_by_year[i] = int(year_df['units_sold'].sum())
        
        # Populate data rows
        for r_idx, region in enumerate(regions):
            row_idx = r_idx + 3
            table.cell(row_idx, 0).text = region
            
            if region == 'Global Total':
                for y_idx in range(len(periods_years)):
                    total_units = total_by_year.get(y_idx, 0)
                    table.cell(row_idx, y_idx+1).text = f'{total_units:,}'
                
                table.cell(row_idx, len(periods_years)+1).text = '100.0%'
            else:
                if has_regional_breakdown:
                    for y_idx in range(len(periods_years)):
                        units = regional_data.get((region, y_idx), 0)
                        table.cell(row_idx, y_idx+1).text = f'{units:,}'
                    
                    current_year_units = regional_data.get((region, len(periods_years)-1), 0)
                    current_year_total = total_by_year.get(len(periods_years)-1, 0)
                    if current_year_total > 0 and current_year_units > 0:
                        percentage = (current_year_units / current_year_total) * 100
                        table.cell(row_idx, len(periods_years)+1).text = f'{percentage:.1f}%'
                    else:
                        table.cell(row_idx, len(periods_years)+1).text = '0.0%'
                else:
                    for y_idx in range(len(periods_years)):
                        table.cell(row_idx, y_idx+1).text = '0'
                    table.cell(row_idx, len(periods_years)+1).text = '0.0%'
        
        # Format table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
        
        for i in range(3):
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
        
        for cell in table.columns[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        doc.add_paragraph()
        
        # Analysis
        analysis_heading = doc.add_paragraph('SALES DATA ANALYSIS')
        analysis_heading.runs[0].font.name = 'Arial'
        analysis_heading.runs[0].font.size = Pt(10)
        analysis_heading.runs[0].font.bold = True
        analysis_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        analysis_para = doc.add_paragraph(analysis_text)
        analysis_para.runs[0].font.name = 'Arial'
        analysis_para.runs[0].font.size = Pt(10)
        analysis_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        doc.add_picture(chart_path, width=Inches(6.0))
        doc.add_paragraph()
        
        # Population
        pop_heading = doc.add_paragraph('Size and Characteristics of Population Using Device')
        pop_heading.runs[0].font.name = 'Arial'
        pop_heading.runs[0].font.size = Pt(10)
        pop_heading.runs[0].font.bold = True
        pop_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)
        
        usage_heading = doc.add_paragraph('Usage Frequency')
        usage_heading.runs[0].font.name = 'Arial'
        usage_heading.runs[0].font.size = Pt(10)
        usage_heading.runs[0].font.bold = True
        
        single_para = doc.add_paragraph(
            f'Single Use per Patient:   {"☑ Yes   ☐ No" if self.device.single_use else "☐ Yes   ☑ No"}'
        )
        single_para.runs[0].font.name = 'Arial'
        single_para.runs[0].font.size = Pt(10)
        
        multi_para = doc.add_paragraph(
            f'Multiple Uses per Patient:   {"☐ Yes   ☑ No" if self.device.single_use else "☑ Yes   ☐ No"}'
        )
        multi_para.runs[0].font.name = 'Arial'
        multi_para.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()
        
        pop_est_heading = doc.add_paragraph('Estimated Size of the Patient Population Exposed to the Device')
        pop_est_heading.runs[0].font.name = 'Arial'
        pop_est_heading.runs[0].font.size = Pt(10)
        pop_est_heading.runs[0].font.bold = True
        
        pop_para = doc.add_paragraph(population_text)
        pop_para.runs[0].font.name = 'Arial'
        pop_para.runs[0].font.size = Pt(10)
        pop_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        doc.add_paragraph()
        
        pop_char_heading = doc.add_paragraph('Characteristics of Patient Population Exposed to the Device')
        pop_char_heading.runs[0].font.name = 'Arial'
        pop_char_heading.runs[0].font.size = Pt(10)
        pop_char_heading.runs[0].font.bold = True
        
        # Use the characteristics text that was generated along with population
        char_para = doc.add_paragraph(characteristics_text)
        char_para.runs[0].font.name = 'Arial'
        char_para.runs[0].font.size = Pt(10)
        char_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        return doc


def main():
    # Ensure output directory exists
    os.makedirs('output', exist_ok=True)
    
    # Point to inputs folder in root directory
    inputs_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'inputs')
    
    # Device information
    device = DeviceInfo(
        name="INCA Complete Set",
        basic_udi_di="[Basic UDI-DI to be provided]",
        device_class="IIa",
        implantable=False,
        single_use=True
    )
    
    # Use actual sales data from inputs folder
    sales_file = os.path.join(inputs_dir, '33_sales.xlsx')
    
    if not os.path.exists(sales_file):
        print(f"ERROR: {sales_file} not found.")
        print(f"  Looking in: {inputs_dir}")
        return
    
    # Load and process sales data with regional mapping
    print(f"Loading sales data from {sales_file}...")
    sales_df = pd.read_excel(sales_file)
    
    print(f"  Detected columns: {list(sales_df.columns)[:10]}...")
    
    # Flexible column detection - handle different naming conventions
    month_col = None
    year_col = None
    quantity_col = None
    country_col = None
    region_col = None
    
    for col in sales_df.columns:
        col_lower = col.lower()
        if 'month' in col_lower and month_col is None:
            month_col = col
        if 'year' in col_lower and year_col is None:
            year_col = col
        if 'quantity' in col_lower or 'units' in col_lower and quantity_col is None:
            quantity_col = col
        if 'country' in col_lower and 'shipping' in col_lower and country_col is None:
            country_col = col
        elif 'country' in col_lower and country_col is None:
            country_col = col
        if 'region' in col_lower and region_col is None:
            region_col = col
    
    print(f"  Using columns: Month={month_col}, Year={year_col}, Quantity={quantity_col}, Country={country_col}")
    
    if month_col and year_col and quantity_col:
        # Apply country to region mapping
        if country_col:
            sales_df['PSUR_Region'] = sales_df[country_col].map(COUNTRY_TO_REGION)
            
            # Handle unmapped countries
            unmapped = sales_df[sales_df['PSUR_Region'].isna()][country_col].unique()
            if len(unmapped) > 0:
                print(f"  WARNING: {len(unmapped)} unmapped countries assigned to 'Other'")
                sales_df['PSUR_Region'] = sales_df['PSUR_Region'].fillna('Other')
        elif region_col:
            # Use existing region column if available
            print(f"  Using existing region column: {region_col}")
            sales_df['PSUR_Region'] = sales_df[region_col]
        else:
            print("  WARNING: No country/region column found, using global data only")
            sales_df['PSUR_Region'] = 'Global'
        
        # Create period column
        month_map = {
            'January': 1, 'February': 2, 'March': 3, 'April': 4,
            'May': 5, 'June': 6, 'July': 7, 'August': 8,
            'September': 9, 'October': 10, 'November': 11, 'December': 12,
            'Jan': 1, 'Feb': 2, 'Mar': 3, 'Apr': 4, 'May': 5, 'Jun': 6,
            'Jul': 7, 'Aug': 8, 'Sep': 9, 'Oct': 10, 'Nov': 11, 'Dec': 12
        }
        sales_df['month_num'] = sales_df[month_col].map(month_map)
        
        # Handle numeric months (1-12)
        if sales_df['month_num'].isna().any():
            # Try converting to int if it's numeric
            try:
                sales_df['month_num'] = sales_df['month_num'].fillna(pd.to_numeric(sales_df[month_col], errors='coerce'))
            except:
                pass
        
        sales_df['period'] = pd.to_datetime(
            sales_df[[year_col, 'month_num']].rename(
                columns={year_col: 'year', 'month_num': 'month'}
            ).assign(day=1)
        )
        
        # Aggregate by region and period
        processed_sales = sales_df.groupby(['PSUR_Region', 'period'])[quantity_col].sum().reset_index()
        processed_sales.columns = ['region', 'period', 'units_sold']
        
        # Also create annual summary for verification
        annual_by_region = sales_df.groupby(['PSUR_Region', year_col])[quantity_col].sum().reset_index()
        annual_by_region.columns = ['region', 'year', 'units_sold']
        
        processed_sales.to_csv('output/sales_processed.csv', index=False)
        annual_by_region.to_csv('output/sales_annual_by_region.csv', index=False)
        
        total_units = processed_sales['units_sold'].sum()
        num_regions = len(processed_sales['region'].unique())
        num_periods = len(processed_sales['period'].unique())
        
        print(f"  Processed {num_periods} periods across {num_regions} regions")
        print(f"  Total units: {total_units:,}")
        print(f"  Regions: {sorted(processed_sales['region'].unique())}")
    else:
        print(f"ERROR: Required columns not found in {sales_file}")
        print(f"  Need: Month column, Year column, Quantity/Units column")
        print(f"  Found: {list(sales_df.columns)}")
        return
    
    # Demographics
    demographics = {
        'Age Range': 'All adult ages',
        'Gender': 'All genders',
        'Primary Indication': 'Obstetric and gynecological procedures'
    }
    
    # Generate Section C
    print("\nGenerating Section C...")
    generator = PSURSectionCGenerator(device)
    
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_filename = f'PSUR_Section_C_{timestamp}.docx'
    output_path = f'output/{output_filename}'
    
    # Point to CER and previous PSUR in inputs folder
    cer_path = os.path.join(inputs_dir, 'cer.pdf')
    previous_psur_path = os.path.join(inputs_dir, 'Previous_psur.docx')
    
    output = generator.generate(
        sales_data_path='output/sales_processed.csv',
        sales_criteria=SalesCriteria.UNITS_DISTRIBUTED,
        first_market_date='2020-04-01',
        usage_factor=1.0,
        demographics=demographics,
        output_path=output_path,
        cer_path=cer_path,  # Will use if exists
        previous_psur_path=previous_psur_path  # Will use if exists
    )
    
    print(f"\n[SUCCESS] Generated: {output}")


if __name__ == "__main__":
    main()

